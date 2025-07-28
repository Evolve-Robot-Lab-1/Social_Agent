from flask import Flask, render_template, request, jsonify, url_for, send_file, redirect
from io import BytesIO
import base64
import requests
import google.generativeai as genai
import groq
from dotenv import load_dotenv
import json
import os
import re
from PIL import Image, ImageDraw
import io
import random
import uuid
import time
import openai
import math
from werkzeug.utils import secure_filename
import copy
import traceback
import sys

# Add new imports for PDF support
try:
    from pypdf import PdfReader
except ImportError:
    print("pypdf not installed. PDF support will be limited.")
    print("To install: pip install pypdf")

# Try to import pdf2image for enhanced PDF support
try:
    from pdf2image import convert_from_path
except ImportError:
    print("pdf2image not installed. PDF image extraction will be limited.")
    print("To install: pip install pdf2image")
    print("\nIMPORTANT: pdf2image requires poppler to be installed:")
    print("1. Download poppler from: https://github.com/oschwartz10612/poppler-windows/releases/")
    print("2. Extract to C:\\poppler")
    print("3. Add C:\\poppler\\Library\\bin to your PATH environment variable")
    print("   You can do this by running: $env:PATH += \";C:\\poppler\\Library\\bin\"")
    print("   For permanent PATH addition, add it through System Properties > Environment Variables")

# Try to verify poppler is working
try:
    import pdf2image
    # Try to get the poppler path
    poppler_path = pdf2image.poppler_path()
    if poppler_path:
        print(f"Poppler found at: {poppler_path}")
    else:
        print("Poppler path not set. PDF to image conversion may not work correctly.")
        print("You may need to specify the path manually in convert_from_path() calls.")
except Exception as e:
    print(f"Could not verify poppler installation: {str(e)}")
    print("PDF to image conversion may not work correctly.")

load_dotenv()

# Initialize Flask app
app = Flask(__name__)

# Configure upload and templates folders
UPLOAD_FOLDER = "uploads"
TEMPLATES_FOLDER = "static/templates"
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# Ensure directories exist
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(TEMPLATES_FOLDER, exist_ok=True)

# Load API Keys from Environment Variables
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
NEWS_API_KEY = os.getenv("NEWS_API_KEY")
UNSPLASH_API_KEY = os.getenv("UNSPLASH_API_KEY", "")  
HUGGINGFACE_API_KEY = os.getenv("HUGGINGFACE_API_KEY")

# Initialize Groq client
client = groq.Groq(api_key=GROQ_API_KEY)

# Template storage
TEMPLATES_FOLDER = "static/templates"
os.makedirs(TEMPLATES_FOLDER, exist_ok=True)

# Default tech blog template
tech_blog_template = {
    "id": "tech",
    "name": "Tech Blog",
    "title": "The Future of Artificial Intelligence and Machine Learning",
    "author": "Tech Insights Team",
    "date": "March 12, 2025",
    "content": [
        {
            "type": "text",
            "content": "Artificial intelligence continues to evolve at a remarkable pace, transforming industries and reshaping our daily lives. The convergence of improved algorithms, massive datasets, and unprecedented computing power has propelled AI from research laboratories into our everyday experiences, from voice assistants to autonomous vehicles, from fraud detection systems to medical diagnostics."
        },
        {
            "type": "image",
            "url": "https://images.unsplash.com/photo-1677442136019-21780ecad995",
            "alt": "AI Technology Visualization"
        },
        {
            "type": "heading",
            "content": "Recent Breakthroughs in AI"
        },
        {
            "type": "text",
            "content": "Recent breakthroughs in neural networks and machine learning have accelerated progress in fields like natural language processing and computer vision. Large language models now generate text virtually indistinguishable from human writing, while diffusion models create photorealistic images from text prompts with astonishing fidelity. These capabilities are not merely impressive technical achievements—they're fundamentally changing how creative work is produced and consumed."
        },
        {
            "type": "image",
            "url": "https://images.unsplash.com/photo-1620712943543-bcc4688e7485",
            "alt": "Neural network visualization"
        },
        {
            "type": "heading",
            "content": "Enterprise AI Adoption"
        },
        {
            "type": "text",
            "content": "As AI technologies mature, we're witnessing unprecedented enterprise adoption. Organizations are moving beyond experimentation to deeply integrate AI throughout their operations. Process automation, intelligent document processing, and predictive analytics are becoming standard components of the business toolkit. According to recent industry surveys, over 65% of enterprises have increased their AI investments substantially in the past year, with the greatest growth in retail, financial services, and healthcare sectors."
        },
        {
            "type": "image",
            "url": "https://images.unsplash.com/photo-1661956602868-6ae368943878",
            "alt": "Enterprise AI implementation"
        },
        {
            "type": "heading",
            "content": "Ethical Considerations and Responsible AI"
        },
        {
            "type": "text",
            "content": "As AI systems become more powerful, the ethical considerations surrounding their deployment grow increasingly complex. Issues of bias, fairness, transparency, and accountability have moved from academic discussions to board rooms. Leading organizations are developing comprehensive responsible AI frameworks that address these concerns while enabling innovation. Regulatory frameworks are beginning to emerge globally, with the EU's AI Act representing the most comprehensive attempt to classify AI applications according to risk and establish appropriate governance mechanisms."
        },
        {
            "type": "image",
            "url": "https://images.unsplash.com/photo-1668001027022-c9e377de649c",
            "alt": "AI ethics concept image"
        },
        {
            "type": "heading",
            "content": "The Future of Human-AI Collaboration"
        },
        {
            "type": "text",
            "content": "Perhaps the most promising direction for AI development is not autonomy but augmentation—systems designed to enhance human capabilities rather than replace them. We're seeing the emergence of 'centaur' models across industries, where humans and AI systems collaborate to achieve outcomes superior to what either could accomplish alone. In medicine, AI-enhanced diagnostic systems empower doctors to detect conditions earlier and with greater accuracy. In engineering, generative design tools suggest novel solutions that might never occur to human designers. These collaborative approaches address many of the concerns about displacement while maximizing the unique strengths of both human and artificial intelligence."
        },
        {
            "type": "image",
            "url": "https://images.unsplash.com/photo-1531746790731-6c087fecd65a",
            "alt": "Human and AI collaboration concept"
        }
    ],
    "image": "https://images.unsplash.com/photo-1485827404703-89b55fcc595e"
}

# Finance blog template with expanded content
finance_blog_template = {
    "id": "finance",
    "name": "Finance Blog",
    "title": "Economic Update: Market Expectations and Bond Trends for 2025",
    "author": "Financial Analysis Team",
    "date": "February 27, 2025",
    "content": [
        {
            "type": "text",
            "content": "As we navigate through 2025, the global economic landscape presents a complex interplay of market dynamics, technological disruption, and evolving investment strategies. Traditional financial models are being challenged by new paradigms, while the intersection of artificial intelligence and finance continues to reshape how we think about money, investments, and economic growth. Financial institutions worldwide are adapting to these changes, with many embracing digital transformation as a core strategic initiative rather than simply a technological upgrade."
        },
        {
            "type": "image",
            "url": "https://images.unsplash.com/photo-1611974789855-9c2a0a7236a3",
            "alt": "Modern financial district with skyscrapers"
        },
        {
            "type": "heading",
            "content": "Global Market Trends and Economic Indicators"
        },
        {
            "type": "text",
            "content": "Recent data from bond markets suggests elevated inflation expectations for the years ahead. The gap between standard U.S. Treasury securities and inflation-protected bonds, known as the breakeven rate, has shown significant movement. Current bond prices imply a 2.7% annual Consumer Price Index inflation for the next five years. This divergence reflects growing investor concern about inflationary pressures despite central bank assertions that recent price increases are transitory."
        },
        {
            "type": "image",
            "url": "https://images.unsplash.com/photo-1642543492481-44e81e3914a6",
            "alt": "Digital financial data visualization"
        },
        {
            "type": "text",
            "content": "The yield curve, another critical market indicator, has undergone meaningful shifts. The flattening of the curve in certain segments suggests changing market expectations about future economic growth and monetary policy. Specifically, the 10-year Treasury yield stabilizing around 4.2% indicates a market consensus forming around medium-term growth and inflation prospects. Portfolio managers are increasingly allocating assets based on these yield curve signals, with many rotating into sectors that historically perform well during periods of higher inflation."
        },
        {
            "type": "split-section",
            "layout": "image-right",
            "image": {
                "url": "https://images.unsplash.com/photo-1563986768609-322da13575f3",
                "alt": "Market analyst examining financial charts"
            },
            "content": {
                "type": "text",
                "content": "Market analysts point to several key factors driving these trends: the continued digital transformation of financial services, evolving regulatory frameworks, and changing consumer preferences. The rise of decentralized finance (DeFi) and digital assets has introduced new variables into traditional market analysis. Central banks worldwide are accelerating research into Central Bank Digital Currencies (CBDCs), with pilot programs already underway in China, Sweden, and the Bahamas, potentially reshaping the future of monetary policy implementation."
            }
        },
        {
            "type": "heading",
            "content": "Technological Innovation in Financial Services"
        },
        {
            "type": "text",
            "content": "The financial technology sector continues to evolve rapidly, with artificial intelligence and blockchain technology leading the transformation. Traditional banks are investing heavily in digital infrastructure, while fintech startups are challenging established business models with innovative solutions. This technological revolution is not just changing how we conduct transactions, but also how we think about value and financial relationships. The integration of machine learning algorithms into credit decisioning has improved access to financial services for previously underserved populations, while simultaneously reducing default rates for lenders through more nuanced risk assessment models."
        },
        {
            "type": "image",
            "url": "https://images.unsplash.com/photo-1640340434855-6084b1f4901c",
            "alt": "Advanced trading terminal with multiple screens"
        },
        {
            "type": "text",
            "content": "Blockchain technology continues to mature beyond its cryptocurrency origins, finding applications in trade finance, securities settlement, and identity verification. Enterprise blockchain solutions are being deployed by major financial institutions to reduce friction in cross-border payments and streamline back-office operations. Smart contracts are automating complex financial agreements, reducing overhead costs and minimizing the potential for disputes. These innovations are particularly transformative for emerging markets, where traditional financial infrastructure limitations can be leapfrogged through distributed ledger technologies."
        },
        {
            "type": "grid-section",
            "columns": 2,
            "items": [
                {
                    "type": "stat-card",
                    "title": "Digital Payments Growth",
                    "value": "127%",
                    "trend": "up",
                    "description": "Year-over-year increase in digital transaction volume"
                },
                {
                    "type": "stat-card",
                    "title": "AI Adoption in Finance",
                    "value": "85%",
                    "trend": "up",
                    "description": "Financial institutions implementing AI solutions"
                },
                {
                    "type": "stat-card",
                    "title": "Sustainable Investment",
                    "value": "$31.7T",
                    "trend": "up",
                    "description": "Global ESG assets under management"
                },
                {
                    "type": "stat-card",
                    "title": "Fintech Market Share",
                    "value": "23%",
                    "trend": "up",
                    "description": "Of global financial services revenue"
                }
            ]
        },
        {
            "type": "heading",
            "content": "Investment Strategies for the Digital Age"
        },
        {
            "type": "text",
            "content": "Modern investment strategies are evolving to incorporate both traditional wisdom and new technological capabilities. Quantitative analysis, powered by machine learning algorithms, is becoming increasingly sophisticated, enabling investors to identify patterns and opportunities that were previously invisible. The democratization of investment tools through mobile apps and robo-advisors has made sophisticated investment strategies accessible to retail investors. Factor-based investing has gained significant traction, with multi-factor models demonstrating resilience across various market conditions by balancing exposure to value, momentum, quality, and low volatility characteristics."
        },
        {
            "type": "image",
            "url": "https://images.unsplash.com/photo-1639322537228-f710d846310a",
            "alt": "Investment strategy planning session"
        },
        {
            "type": "text",
            "content": "Alternative data sources are becoming increasingly central to investment decision-making. Satellite imagery is being used to monitor agricultural yields and retail parking lot occupancy. Natural language processing algorithms analyze earnings call transcripts to detect sentiment shifts among corporate executives. Mobile payment data provides real-time insights into consumer spending patterns. These non-traditional data streams, combined with computational power to analyze them, are creating new forms of information asymmetry that sophisticated investors can leverage for alpha generation."
        },
        {
            "type": "heading",
            "content": "The Future of Sustainable Finance"
        },
        {
            "type": "text",
            "content": "Environmental, Social, and Governance (ESG) considerations are becoming central to investment decision-making. Asset managers are developing sophisticated frameworks to evaluate companies' sustainability practices, while regulators are implementing new disclosure requirements. This shift towards sustainable finance is driving innovation in financial products and creating new opportunities for both investors and entrepreneurs. The EU Taxonomy for sustainable activities has established clear criteria for what constitutes environmentally sustainable economic activities, bringing standardization to a previously fragmented market and facilitating the flow of capital towards genuinely sustainable projects."
        },
        {
            "type": "image",
            "url": "https://images.unsplash.com/photo-1623659873902-3b11d35d5c33",
            "alt": "Renewable energy investment project"
        },
        {
            "type": "text",
            "content": "Green bonds continue to see record issuance, with proceeds dedicated to environmental projects ranging from renewable energy infrastructure to biodiversity conservation. Sustainability-linked loans and bonds, which tie financing costs to the achievement of specific sustainability targets, are incentivizing corporations to accelerate their transition to more sustainable business models. Impact investing is moving from the periphery to the mainstream, with major asset managers launching dedicated funds focused on generating measurable social and environmental benefits alongside financial returns."
        },
        {
            "type": "heading",
            "content": "Regulatory Developments and Compliance Challenges"
        },
        {
            "type": "text",
            "content": "The regulatory landscape for financial services continues to evolve at a rapid pace. Global harmonization efforts, such as the Basel IV framework for banking regulation, aim to establish consistent standards across jurisdictions. However, regional variations in implementation create compliance complexity for multinational financial institutions. The EU's Digital Operational Resilience Act (DORA) and the UK's Operational Resilience Framework are elevating expectations regarding financial institutions' ability to withstand, absorb and recover from operational disruptions, with an emphasis on critical business services rather than just systems."
        },
        {
            "type": "image",
            "url": "https://images.unsplash.com/photo-1589829085413-56de8ae18c73",
            "alt": "Financial regulatory compliance meeting"
        },
        {
            "type": "text",
            "content": "Cryptocurrency regulation represents a significant frontier, with jurisdictions taking divergent approaches. Some are creating comprehensive regulatory frameworks to encourage innovation while protecting consumers, while others are adopting more restrictive stances. Anti-money laundering and counter-terrorism financing requirements continue to expand, with beneficial ownership registries becoming more common globally. Financial institutions are investing heavily in RegTech solutions to manage these compliance challenges efficiently, with artificial intelligence playing an increasingly important role in transaction monitoring and fraud detection."
        },
        {
            "type": "expert-panel",
            "title": "Expert Insights",
            "experts": [
                {
                    "name": "Dr. Sarah Chen",
                    "role": "Chief Investment Strategist",
                    "quote": "The convergence of AI and financial markets is creating unprecedented opportunities for sophisticated risk management and alpha generation.",
                    "image": "https://images.unsplash.com/photo-1573497019940-1c28c88b4f3e"
                },
                {
                    "name": "Michael Rodriguez",
                    "role": "Global Markets Director",
                    "quote": "Sustainable finance is no longer just a trend - it's becoming a fundamental factor in investment decision-making.",
                    "image": "https://images.unsplash.com/photo-1472099645785-5658abf4ff4e"
                }
            ]
        }
    ],
    "image": "https://images.unsplash.com/photo-1611974789855-9c2a0a7236a3"
}

# Travel blog template with significantly expanded content and more images
travel_blog_template = {
    "id": "travel",
    "name": "Travel Blog",
    "title": "Hidden Gems of Southeast Asia: A Journey Beyond the Tourist Trail",
    "author": "Travel Explorer Team",
    "date": "March 15, 2024",
    "content": [
        {
            "type": "text",
            "content": "Southeast Asia has long captivated travelers with its rich tapestry of cultures, stunning landscapes, and warm hospitality. While destinations like Bangkok, Bali, and Singapore draw millions of visitors annually, the region harbors countless hidden treasures waiting to be discovered by those willing to venture off the beaten path. These lesser-known destinations offer authentic experiences, unburdened by mass tourism, where travelers can forge genuine connections with local communities and immerse themselves in the region's incredible diversity."
        },
        {
            "type": "image",
            "url": "https://images.unsplash.com/photo-1552465011-b4e21bf6e79a",
            "alt": "Pristine beach with traditional long-tail boats"
        },
        {
            "type": "heading",
            "content": "Undiscovered Islands of Thailand"
        },
        {
            "type": "text",
            "content": "Beyond the well-known islands of Phuket and Koh Samui lies an archipelago of lesser-visited gems. Koh Yao Noi, situated in Phang Nga Bay, offers a glimpse of traditional Thai island life, with rubber plantations, fishing villages, and empty beaches. The island's western coast provides spectacular sunset views over the limestone karsts that dot the bay, while local communities welcome visitors with genuine warmth and hospitality. The pace of life here remains blissfully unhurried, with bicycles and motorbikes serving as the primary modes of transportation on the island's quiet roads."
        },
        {
            "type": "image",
            "url": "https://images.unsplash.com/photo-1552733407-5d5c46c3bb3b",
            "alt": "Aerial view of pristine island beaches"
        },
        {
            "type": "text",
            "content": "Further south, the Trang Islands remain one of Thailand's best-kept secrets. Koh Mook's Emerald Cave (Tham Morakot) requires swimming through a dark passageway that suddenly opens to reveal a hidden beach surrounded by towering limestone cliffs and lush vegetation – a true pirate's hideaway. Nearby Koh Kradan offers some of Thailand's most pristine beaches and spectacular snorkeling directly from shore, with vibrant coral reefs just a short swim from the powdery sand. Koh Ngai's eastern shore features shallow, crystal-clear waters perfect for families, while its western side offers sunset vistas that rival any in the world."
        },
        {
            "type": "image",
            "url": "https://images.unsplash.com/photo-1544644181-1484b3fdfc32",
            "alt": "Emerald Cave hidden beach surrounded by limestone cliffs"
        },
        {
            "type": "heading",
            "content": "Vietnam's Hidden Highlands"
        },
        {
            "type": "text",
            "content": "In the northern reaches of Vietnam, beyond the tourist hotspots of Hanoi and Halong Bay, the Ha Giang province offers an untouched landscape of dramatic mountain passes, terraced rice fields, and authentic cultural experiences. Home to various ethnic minority groups, this region provides opportunities for homestays, allowing travelers to immerse themselves in local traditions and daily life. The Dong Van Karst Plateau Geopark, recognized by UNESCO for its geological significance, features a otherworldly landscape of limestone formations that have been shaped over millions of years."
        },
        {
            "type": "split-section",
            "layout": "image-left",
            "image": {
                "url": "https://images.unsplash.com/photo-1528127269322-539801943592",
                "alt": "Traditional Vietnamese village in the mountains"
            },
            "content": {
                "type": "text",
                "content": "The winding roads of the Ha Giang Loop offer some of Asia's most spectacular motorcycle routes, taking riders through remote villages and breathtaking mountain vistas. Local markets burst with color and activity, where different ethnic groups gather to trade goods and maintain centuries-old traditions. The Sunday market in Bac Ha is particularly renowned, with Flower Hmong women in vibrant traditional dress selling textiles, medicinal herbs, and local produce. This high-altitude region also produces exceptional tea, with ancient tea forests containing trees over 400 years old that produce complex, nuanced flavors found nowhere else."
            }
        },
        {
            "type": "image",
            "url": "https://images.unsplash.com/photo-1518802170774-e31e2bb6b6eb",
            "alt": "Terraced rice fields in northern Vietnam highlands"
        },
        {
            "type": "text",
            "content": "Venturing into central Vietnam, the ancient Cham towers of Qui Nhon stand as testament to the once-powerful Champa civilization that ruled these lands for centuries. These brick structures, adorned with intricate carvings and sculptures, have weathered centuries of monsoons and wars yet remain remarkably intact. The nearby coastal city of Qui Nhon itself offers pristine beaches without the crowds of more famous Vietnamese coastal destinations, along with a thriving seafood culture where fishermen bring their daily catch directly to beachside restaurants."
        },
        {
            "type": "image",
            "url": "https://images.unsplash.com/photo-1557750255-c76072a7aad1",
            "alt": "Ancient Cham towers against blue sky"
        },
        {
            "type": "heading",
            "content": "Cambodia's Forgotten Temples"
        },
        {
            "type": "text",
            "content": "While Angkor Wat draws millions of visitors, Cambodia's Preah Vihear province houses equally impressive but far less visited temple complexes. Perched atop a cliff in the Dangrek Mountains, the Preah Vihear temple offers stunning views across the Cambodian plains and a peaceful atmosphere that's increasingly rare at more popular sites. Built between the 9th and 12th centuries, this Hindu temple dedicated to Shiva features some of the most impressive carved lintels and pediments of any Khmer structure, telling ancient stories of gods and kings through intricate stone carvings."
        },
        {
            "type": "image",
            "url": "https://images.unsplash.com/photo-1508009603885-50cf7c579365",
            "alt": "Ancient temple ruins at sunrise"
        },
        {
            "type": "text",
            "content": "The lesser-known temples of Koh Ker, once a brief capital of the Khmer Empire, lie hidden in the jungle north of Siem Reap. The main structure, Prasat Thom, resembles a seven-tiered pyramid rising dramatically from the forest floor. Unlike the more manicured grounds of Angkor, many of Koh Ker's structures remain partially reclaimed by the jungle, with massive tree roots embracing ancient stones in a haunting display of nature's power. The remote location means visitors can often explore these magnificent structures in solitude, imagining themselves as early explorers discovering these ruins for the first time."
        },
        {
            "type": "image",
            "url": "https://images.unsplash.com/photo-1555855815-43e2a068efdd",
            "alt": "Temple ruins being reclaimed by jungle"
        },
        {
            "type": "heading",
            "content": "Local Cuisine and Cultural Experiences"
        },
        {
            "type": "text",
            "content": "One of the most rewarding aspects of exploring Southeast Asia's hidden corners is discovering authentic local cuisine. From street food vendors in remote market towns to family-run restaurants serving recipes passed down through generations, these culinary experiences offer insight into the region's rich cultural heritage. In northern Thailand's Nan Province, traditional Tai Lue cuisine features distinctive dishes like khao kan jin (rice mixed with blood and herbs, steamed in banana leaves) and kaeng no mai (spicy bamboo shoot curry) that reflect the area's unique cultural influences from both Lanna Thai and Laotian traditions."
        },
        {
            "type": "image",
            "url": "https://images.unsplash.com/photo-1504674900247-0877df9cc836",
            "alt": "Traditional Southeast Asian street food market"
        },
        {
            "type": "text",
            "content": "Throughout the region, cooking classes have become a popular way for travelers to connect with local culture. Unlike the standardized tourist-oriented classes in major cities, those found in rural areas often take place in family homes, beginning with trips to local markets or even foraging expeditions to gather wild herbs and vegetables. In Battambang, Cambodia, visitors can learn to prepare amok (fish curry steamed in banana leaves) and prahok ktis (fermented fish dip) using traditional clay pots over wood fires, techniques that have remained unchanged for centuries despite the conveniences of modern cooking equipment."
        },
        {
            "type": "image",
            "url": "https://images.unsplash.com/photo-1455619452474-d2be8b1e70cd",
            "alt": "Traditional cooking class with local ingredients"
        },
        {
            "type": "heading",
            "content": "Sacred Mountains of Myanmar"
        },
        {
            "type": "text",
            "content": "In the remote regions of Myanmar, ancient pilgrimage sites and mountain monasteries offer spiritual encounters far from the tourist crowds. The Golden Rock at Mount Kyaiktiyo defies physical laws, perched impossibly on the edge of a cliff and drawing devout Buddhists from across the country. The journey to this sacred site is as meaningful as the destination itself, with pilgrims walking barefoot up the mountain path while contemplating their faith. Covered in gold leaf applied by generations of worshippers, the rock glows magnificently at sunrise and sunset, seeming to change colors throughout the day as the light shifts across its surface."
        },
        {
            "type": "image",
            "url": "https://images.unsplash.com/photo-1546542868-c0e57ca3dd4f",
            "alt": "Golden Rock at Kyaiktiyo"
        },
        {
            "type": "text",
            "content": "Mount Popa, an extinct volcano rising dramatically from central Myanmar's plains, hosts a monastery perched atop a sheer volcanic plug. Home to the 37 Great Nats (spirits) of Burmese tradition, this site represents the fascinating synthesis of ancient animist beliefs with Buddhist practices. Visitors climb 777 steps to reach the summit, passing shrines and being kept company by the resident monkeys that have made this sacred mountain their home. From the top, panoramic views extend across the plains to distant mountains, with the golden stupas of the monastery gleaming in the sun against the backdrop of endless sky."
        },
        {
            "type": "image",
            "url": "https://images.unsplash.com/photo-1598127004624-62c8c1cc68b4",
            "alt": "Sunrise at Mount Popa monastery"
        },
        {
            "type": "text",
            "content": "In the Shan State highlands, the remote monastery of Kakku contains over 2,000 stupas crowded together in a relatively small area, creating a mesmerizing forest of spires that dates back to the 3rd century BCE. Many of these stupas feature intricate carvings of celestial beings and mythological creatures that have survived centuries of weathering. Unlike Myanmar's more famous archaeological sites, Kakku remains largely unknown to international travelers, allowing for contemplative exploration without crowds or commercial development."
        },
        {
            "type": "image",
            "url": "https://images.unsplash.com/photo-1519882189396-71b93cb3e428",
            "alt": "Buddhist monastery in misty mountains"
        },
        {
            "type": "heading",
            "content": "Island Life in the Philippines"
        },
        {
            "type": "text",
            "content": "With over 7,000 islands, the Philippines offers endless opportunities for exploration beyond the well-trodden paths of Boracay and Palawan. The Batanes Islands, located in the northernmost part of the Philippines where the Pacific Ocean meets the South China Sea, feature a dramatic landscape of rolling hills, stone villages, and rugged coastlines that more closely resemble the Scottish Highlands than tropical Southeast Asia. The indigenous Ivatan people have developed unique stone architecture designed to withstand the powerful typhoons that frequently pass through the region."
        },
        {
            "type": "image",
            "url": "https://images.unsplash.com/photo-1518509562904-e7ef99cdcc86",
            "alt": "Rolling green hills meeting the ocean in Batanes"
        },
        {
            "type": "text",
            "content": "The Bacuit Archipelago near El Nido contains dozens of limestone islands with hidden lagoons that can only be accessed at specific tide levels through small openings in the rock. Kayaking through these secret passages reveals pristine enclosed beaches and lagoons of impossibly blue water surrounded entirely by towering limestone cliffs draped in vegetation. Marine life thrives in these protected waters, with vibrant coral gardens and reef fish visible just below the surface."
        },
        {
            "type": "image",
            "url": "https://images.unsplash.com/photo-1559444839-b936c2cbf637",
            "alt": "Hidden lagoon with turquoise water surrounded by limestone cliffs"
        },
        {
            "type": "heading",
            "content": "Sustainable Travel and Local Communities"
        },
        {
            "type": "text",
            "content": "As these hidden destinations gain recognition, sustainable tourism practices become increasingly important. Many communities are developing eco-friendly accommodations and responsible tourism initiatives that allow visitors to experience these special places while preserving their natural and cultural heritage for future generations. In Thailand's Koh Prasat Yai, a community-based tourism initiative has helped the island develop small-scale, locally-owned accommodations built using traditional methods and materials, with profits being reinvested into conservation projects and educational opportunities for local children."
        },
        {
            "type": "image",
            "url": "https://images.unsplash.com/photo-1625472603517-1b0dc72846ab",
            "alt": "Eco-friendly bamboo accommodation in the jungle"
        },
        {
            "type": "text",
            "content": "The Cambodia Community-Based Ecotourism Network (CCBEN) connects travelers with village homestays in remote regions, where visitors can participate in conservation efforts such as wildlife monitoring or reforestation projects. These initiatives provide alternative livelihoods for communities that might otherwise depend on environmentally destructive practices like logging or wildlife poaching. By creating economic value for intact ecosystems and cultural traditions, sustainable tourism can help preserve Southeast Asia's hidden gems for generations to come."
        },
        {
            "type": "image",
            "url": "https://images.unsplash.com/photo-1513836279014-a89f7a76ae86",
            "alt": "Community-led reforestation project"
        },
        {
            "type": "expert-panel",
            "title": "Local Insights",
            "experts": [
                {
                    "name": "Mai Tran",
                    "role": "Local Guide, Ha Giang",
                    "quote": "The real magic of our region lies in the small moments - sharing tea with a local family, learning traditional crafts, or discovering a viewpoint that takes your breath away.",
                    "image": "https://images.unsplash.com/photo-1494790108377-be9c29b29330"
                },
                {
                    "name": "Somchai Prak",
                    "role": "Cultural Heritage Expert",
                    "quote": "These lesser-known destinations preserve authentic traditions that are increasingly rare in more touristy areas. They offer a window into Southeast Asia's rich cultural heritage.",
                    "image": "https://images.unsplash.com/photo-1507003211169-0a1dd7228f2d"
                }
            ]
        }
    ],
    "image": "https://images.unsplash.com/photo-1552465011-b4e21bf6e79a"
}

# Save the default template to file
if os.path.exists(os.path.join(TEMPLATES_FOLDER, "tech_template.json")):
    os.remove(os.path.join(TEMPLATES_FOLDER, "tech_template.json"))
with open(os.path.join(TEMPLATES_FOLDER, "tech_template.json"), "w") as f:
    json.dump(tech_blog_template, f)

# Save the finance template to file
if os.path.exists(os.path.join(TEMPLATES_FOLDER, "finance_template.json")):
    os.remove(os.path.join(TEMPLATES_FOLDER, "finance_template.json"))
with open(os.path.join(TEMPLATES_FOLDER, "finance_template.json"), "w") as f:
    json.dump(finance_blog_template, f)

# Save the travel template to file
if os.path.exists(os.path.join(TEMPLATES_FOLDER, "travel_template.json")):
    os.remove(os.path.join(TEMPLATES_FOLDER, "travel_template.json"))
with open(os.path.join(TEMPLATES_FOLDER, "travel_template.json"), "w") as f:
    json.dump(travel_blog_template, f)

# Load template from file without adding randomization
def load_template(template_id):
    try:
        # Extract the base template id without any timestamps or random strings
        base_template_id = template_id.split('_')[0] if '_' in template_id else template_id
        
        template_path = os.path.join(TEMPLATES_FOLDER, f"{base_template_id}_template.json")
        if os.path.exists(template_path):
            with open(template_path, "r") as f:
                template = json.load(f)
                return template
        else:
            # Return default template if file doesn't exist
            if base_template_id == "finance":
                return finance_blog_template
            elif base_template_id == "travel":
                return travel_blog_template
            else:
                return tech_blog_template
    except Exception as e:
        print(f"Error loading template: {str(e)}")
        return tech_blog_template

# Add new function to clean up template files
def cleanup_template_files():
    """Clean up old template files, keeping only the original template files"""
    try:
        # Get all files in template directory
        template_files = os.listdir(TEMPLATES_FOLDER)
        
        # Keep only the base template files
        base_templates = ["tech_template.json", "finance_template.json", "travel_template.json"]
        
        # Delete all other template files
        for file in template_files:
            if file not in base_templates:
                file_path = os.path.join(TEMPLATES_FOLDER, file)
                try:
                    os.remove(file_path)
                    print(f"Deleted template file: {file}")
                except Exception as e:
                    print(f"Error deleting file {file}: {str(e)}")
        
        # Also check for any temporary files in the uploads directory
        uploads_dir = "uploads"
        if os.path.exists(uploads_dir):
            for filename in os.listdir(uploads_dir):
                # Keep only files from the last hour to avoid disrupting current sessions
                file_path = os.path.join(uploads_dir, filename)
                try:
                    file_modification_time = os.path.getmtime(file_path)
                    if time.time() - file_modification_time > 3600:  # 1 hour
                        os.remove(file_path)
                        print(f"Removed old upload file: {filename}")
                except Exception as e:
                    print(f"Error checking/removing file {filename}: {str(e)}")
        
        return True
    except Exception as e:
        print(f"Error cleaning up template files: {str(e)}")
        return False

# Modify save_template function to use the base template name only
def save_template(template_data):
    try:
        template_id = template_data.get("id", "tech")
        
        # Extract the base template id without any timestamps or random strings
        base_template_id = template_id.split('_')[0] if '_' in template_id else template_id
        
        # Add metadata for tracking
        if "metadata" not in template_data:
            template_data["metadata"] = {}
            
        template_data["metadata"]["last_saved"] = int(time.time())
        
        # Save only to the base template file
        template_path = os.path.join(TEMPLATES_FOLDER, f"{base_template_id}_template.json")
        
        with open(template_path, "w") as f:
            json.dump(template_data, f)
            
        # Run cleanup to remove any extra files
        cleanup_template_files()
            
        return True
    except Exception as e:
        print(f"Error saving template: {str(e)}")
        return False

# Enhanced chatbot function to handle user queries with randomization
def query_chatbot(user_message, template_id="tech"):
    try:
        print(f"[DEBUG] query_chatbot called with user_message={user_message}, template_id={template_id}")
        if not GROQ_API_KEY:
            return "Error: Groq API key is missing.", None
        current_template = load_template(template_id)
        client = groq.Groq(api_key=GROQ_API_KEY)
        random_seed = int(time.time())
        random_factor = random.randint(1, 1000)
        is_file_upload = user_message.startswith("Generate a complete blog post")
        if is_file_upload:
            # Enhanced system prompt for file uploads with more detailed instructions
            system_prompt = """You are BlogGen Assistant, a specialized AI that generates high-quality blog content based on user uploads.
            Current template type: {template_id}
            Generation seed: {random_seed}
            Variation factor: {random_factor}
            
            YOUR TASK IS TO GENERATE COMPLETELY UNIQUE CONTENT EACH TIME, EVEN IF THE INPUTS ARE SIMILAR.
            Use the generation seed and variation factor to create highly varied content.
            
            Your task is to create a fully-formed blog post using the uploaded file content while maintaining the structure of the selected template.
            
            IMPORTANT INSTRUCTIONS:
            1. Generate a compelling title related to the uploaded content - make it unique and different each time
            2. Create appropriate headings and subheadings that organize the content logically - vary structure based on random seed
            3. Expand the content to be thorough and informative, at least 1200 words total
            4. Suggest appropriate alt text for images based on the section content
            5. Maintain the overall structure and flow of the template type
            6. ENSURE YOU CREATE DIFFERENT CONTENT EACH TIME even for similar inputs
            7. Return your complete response as a JSON object matching the template structure
            
            The JSON must follow this structure:
            ```json
            {{
              "title": "Generated Title",
              "content": [
                {{"type": "text", "content": "Introduction paragraph..."}},
                {{"type": "image", "url": "IMAGE_URL", "alt": "Descriptive alt text"}},
                {{"type": "heading", "content": "First Section Heading"}},
                {{"type": "text", "content": "First section content..."}},
                ...
              ]
            }}
            ```
            """
            # For uploads, use larger token limit to handle complete blog generation
            max_tokens = 2000
        else:
            # Enhanced system prompt for blog modifications
            system_prompt = """You are BlogGen Assistant, a helpful AI that assists users with blog creation and content modifications.
        Current template: {template_id}
        Generation seed: {random_seed}
        Variation factor: {random_factor}
            
            IMPORTANT: Ensure your responses and generated content are unique each time, even for similar requests.
            
            When users request changes to their blog:
            1. Understand their specific requirements
            2. Modify only the requested parts while maintaining the overall structure
            3. Return changes in a JSON block with the specific updates
            4. Keep the existing content and structure intact for parts not being modified
            5. Use the random seeds to vary your approach to content creation
            
            For example, if user asks to change the title, return:
            ```json
            {{
              "title": "New Title"
            }}
            ```
            
            If user asks to modify a specific section, return:
            ```json
            {{
              "content": [
                {{"type": "text", "content": "Updated content..."}},
                {{"type": "heading", "content": "Updated heading"}}
              ]
            }}
            ```
            
            Keep responses clear and helpful. If user asks for specific blog help, offer creative suggestions.
            """
            max_tokens = 500

        # Try to get a response from Groq API
        try:
            response = client.chat.completions.create(
                model="llama3-8b-8192",  # Use a widely available model name
                messages=[
                    {"role": "system", "content": system_prompt.format(
                        template_id=template_id,
                        random_seed=random_seed,
                        random_factor=random_factor
                    )},
                    {"role": "user", "content": user_message}
                ],
                temperature=0.9,  # Increase temperature for more variety
                max_tokens=max_tokens
            )
            response_text = response.choices[0].message.content
        except Exception as groq_error:
            import traceback
            print("[ERROR] Groq API call failed:")
            traceback.print_exc()
            print(f"[ERROR] Exception details: {groq_error}")
            # Fallback to Gemini if Groq fails
            if GEMINI_API_KEY:
                try:
                    genai.configure(api_key=GEMINI_API_KEY)
                    model = genai.GenerativeModel(model_name="gemma-3n-e4b-it")
                    prompt = f"""System: {system_prompt.format(
                        template_id=template_id,
                        random_seed=random_seed,
                        random_factor=random_factor
                    )}\nUser: {user_message}"""
                    response = model.generate_content(prompt)
                    response_text = response.text
                except Exception as gemini_error:
                    print("[ERROR] Gemini API call failed:")
                    traceback.print_exc()
                    print(f"[ERROR] Exception details: {gemini_error}")
                    return f"Sorry, I'm having trouble connecting to my AI services. (Gemini error: {str(gemini_error)})", None
            else:
                return f"Sorry, I'm having trouble responding right now. (Groq error: {str(groq_error)})", None
        
        # Rest of the function remains the same
        # Process the response
        updated_template = current_template.copy()  # Create a copy to avoid reference issues
        
        # Add randomization metadata to template
        if "metadata" not in updated_template:
            updated_template["metadata"] = {}
        updated_template["metadata"]["generated_at"] = int(time.time())
        updated_template["metadata"]["random_seed"] = random_seed
        updated_template["metadata"]["random_factor"] = random_factor
        
        # Extract JSON updates if present
        json_blocks = re.findall(r'```json(.*?)```', response_text, re.DOTALL)
        if json_blocks:
            for block in json_blocks:
                try:
                    updates = json.loads(block.strip())
                    
                    # Handle title updates
                    if "title" in updates:
                        updated_template["title"] = updates["title"]
                    
                    # Handle author updates
                    if "author" in updates:
                        updated_template["author"] = updates["author"]
                    
                    # Handle date updates
                    if "date" in updates:
                        updated_template["date"] = updates["date"]
                    
                    # Handle content updates
                    if "content" in updates:
                        # Check if this is a full content replacement or partial update
                        if any(item.get("type") == "text" for item in updates["content"]):
                            # This is likely a full content replacement
                            updated_template["content"] = updates["content"]
                        else:
                            # This is likely a partial update - merge with existing content
                            existing_content = updated_template.get("content", [])
                            for new_item in updates["content"]:
                                # Find matching item in existing content and update it
                                for i, existing_item in enumerate(existing_content):
                                    if existing_item.get("id") == new_item.get("id"):
                                        existing_content[i] = new_item
                                        break
                                else:
                                    # If no match found, append the new item
                                    existing_content.append(new_item)
                    
                    # Handle image updates
                    if "image" in updates:
                        updated_template["image"] = updates["image"]
                    
                    # Save the updated template
                    save_template(updated_template)
                except json.JSONDecodeError as e:
                    print(f"[ERROR] JSON decode error: {str(e)}\nRaw AI response block:\n{block.strip()}")
                    print(f"[ERROR] Full AI response:\n{response_text}")
                    return (f"Sorry, I encountered an error processing the AI's response. Please try again or rephrase your request. (Details: {str(e)})", None)
                except Exception as e:
                    print(f"[ERROR] Template update error: {str(e)}\nRaw AI response block:\n{block.strip()}")
                    print(f"[ERROR] Full AI response:\n{response_text}")
                    return (f"Sorry, I encountered an error processing the AI's response. Please try again or rephrase your request. (Details: {str(e)})", None)
            # Clean response and remove JSON blocks
            clean_response = re.sub(r'```json.*?```', '', response_text, flags=re.DOTALL).strip()
            # --- Title extraction logic ---
            if any(word in user_message.lower() for word in ["title", "headline"]):
                title = extract_concrete_title(clean_response)
                if title:
                    return title, updated_template
            return clean_response, updated_template
        else:
            # No JSON block found, just return the AI's plain text response
            clean_response = re.sub(r'```.*?```', '', response_text, flags=re.DOTALL).strip()
            # --- Title extraction logic ---
            if any(word in user_message.lower() for word in ["title", "headline"]):
                title = extract_concrete_title(clean_response)
                if title:
                    return title, updated_template
            return clean_response, updated_template
    except Exception as e:
        print(f"Error in query_chatbot: {str(e)}")
        return f"Error: {str(e)}", None

# Helper function to extract a concrete title from AI response
import re

def extract_concrete_title(response_text):
    """
    Extract a likely title from the AI response, even if the response is generic or includes a preamble.
    Returns a string title or None if not found.
    """
    if not response_text:
        return None
    # Look for common title patterns
    # 1. Markdown or plain: Title: ...
    match = re.search(r"(?:^|\n)\s*Title\s*[:\-–]\s*(.+)", response_text, re.IGNORECASE)
    if match:
        title = match.group(1).strip().strip('"')
        # Stop at first line break or period if present
        title = re.split(r'[\n\r\.]', title)[0].strip()
        return title
    # 2. Quoted or bolded line at the top
    match = re.search(r'^[\"\']?([A-Z][^\n\r]{10,80})[\"\']?$', response_text.strip(), re.MULTILINE)
    if match:
        return match.group(1).strip()
    # 3. First non-empty line if it looks like a title
    lines = response_text.strip().splitlines()
    for line in lines:
        line = line.strip().strip('"')
        if 10 <= len(line) <= 80 and line[0].isupper() and not line.endswith((':', '-', '—')):
            return line
    # 4. Fallback: first sentence, truncated
    sentence = re.split(r'[\.!?]', response_text)[0]
    if 10 < len(sentence) < 80:
        return sentence.strip()
    return None

# Generate HTML for preview from template data
def generate_html_preview(template_data):
    """Generate HTML preview for a blog template"""
    try:
        html = f"""
        <!DOCTYPE html>
        <html lang="en">
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
            <title>{template_data.get('title', 'Blog Preview')}</title>
            <style>
                body {{
                    font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
                    line-height: 1.6;
                    color: #333;
                    max-width: 800px;
                    margin: 0 auto;
                    padding: 20px;
                }}
                img {{
                    max-width: 100%;
                    height: auto;
                    border-radius: 8px;
                    margin: 20px 0;
                }}
                h1, h2, h3 {{
                    color: #2c3e50;
                }}
                .meta {{
                    color: #666;
                    margin-bottom: 30px;
                }}
                .content-section {{
                    margin-bottom: 30px;
                }}
                blockquote {{
                    border-left: 4px solid #ddd;
                    margin: 0;
                    padding-left: 20px;
                    font-style: italic;
                }}
            </style>
        </head>
        <body>
            <h1>{template_data.get('title', 'Blog Title')}</h1>
            <div class="meta">
                <p>By {template_data.get('author', 'Anonymous')} | {template_data.get('date', 'Date')}</p>
            </div>
        """
        
        # Add content sections
        for section in template_data.get('content', []):
            section_type = section.get('type', '')
            if section_type == 'text':
                html += f"""
                <div class="content-section">
                    <p>{section.get('content', '')}</p>
                </div>
                """
            elif section_type == 'heading':
                html += f"""
                <div class="content-section">
                    <h2>{section.get('content', '')}</h2>
                </div>
                """
            elif section_type == 'image':
                html += f"""
                <div class="content-section">
                    <img src="{section.get('url', '')}" alt="{section.get('alt', '')}" loading="lazy">
                    {f'<p class="caption">{section["caption"]}</p>' if 'caption' in section else ''}
                </div>
                """
            elif section_type == 'quote':
                html += f"""
                <div class="content-section">
                    <blockquote>{section.get('content', '')}</blockquote>
                </div>
                """
            elif section_type == 'split-section':
                layout = section.get('layout', 'image-left')
                image_html = f"""
                <div class="split-image">
                    <img src="{section['image']['url']}" alt="{section['image'].get('alt', '')}" loading="lazy">
                </div>
                """ if 'image' in section else ''
                
                html += f"""
                <div class="split-section {layout}">
                    {image_html}
                    <div class="split-content">
                        {section.get('content', '')}
                    </div>
                </div>
                """
        
        html += '</div></div>'
        
        return html
    
    except Exception as e:
        print(f"Error generating HTML preview: {str(e)}")
        return f"<p>Error generating preview: {str(e)}</p>"

@app.route("/")
def home():
    # Clean up any unnecessary template files on application startup
    cleanup_template_files()
    return render_template("index.html")

@app.route("/templates")
def templates():
    return render_template("templates.html")

@app.route("/News")
def news():
    return render_template("news.html")

@app.route("/get_blogs", methods=["GET"])
def get_blogs():
    try:
        # Get all template files
        template_files = [f for f in os.listdir(TEMPLATES_FOLDER) if f.endswith('_template.json')]
        blogs = []
        
        for file in template_files:
            with open(os.path.join(TEMPLATES_FOLDER, file), 'r') as f:
                template = json.load(f)
                blogs.append(template)
        
        return jsonify({"success": True, "blogs": blogs})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)})

@app.route('/chat', methods=['POST'])
def chat():
    try:
        data = request.get_json()
        if not data or 'message' not in data:
            return jsonify({"error": "No message provided"}), 400
        user_message = data['message']
        template_id = data.get('template_id', 'tech')  # Default to tech template

        # Get chatbot response
        response_text, updated_template = query_chatbot(user_message, template_id)
        # Debug output for diagnosis
        print(f"[DEBUG] /query_chatbot: user_message={user_message}, template_id={template_id}, response_text={response_text}, updated_template={'present' if updated_template is not None else 'None'}")
        # Always return a valid response, even if the chatbot fails
        if not response_text or (isinstance(response_text, str) and response_text.lower().startswith('error')):
            import traceback
            print(f"[ERROR] query_chatbot failed. user_message: {user_message}, template_id: {template_id}, response_text: {response_text}")
            traceback.print_stack()
            response_text = "Sorry, I encountered an error processing your request. Please try again."
        response = {"response": response_text}
        if updated_template is not None:
            response["template"] = updated_template
            response["templateId"] = template_id
        return jsonify(response)
    except Exception as e:
        print(f"Error in /chat route: {str(e)}")
        return jsonify({"response": "Sorry, I encountered an error processing your request. Please try again."}), 500

@app.route('/get_template/<template_id>', methods=['GET'])
def get_template(template_id):
    template = load_template(template_id)
    return jsonify(template)

@app.route('/api/placeholder/<int:width>/<int:height>')
def placeholder_image(width, height):
    """Generate a placeholder image with the specified dimensions"""
    try:
        # Create a new image with a white background
        image = Image.new('RGB', (width, height), 'white')
        draw = ImageDraw.Draw(image)
        
        # Draw a border
        draw.rectangle([0, 0, width-1, height-1], outline='#ddd')
        
        # Draw diagonal lines
        draw.line([(0, 0), (width, height)], fill='#ddd')
        draw.line([(0, height), (width, 0)], fill='#ddd')
        
        # Add dimensions text
        text = f'{width}x{height}'
        # Note: In a production environment, you'd want to use a proper font
        # For this example, we'll use a simple text
        text_bbox = draw.textbbox((0, 0), text)
        text_width = text_bbox[2] - text_bbox[0]
        text_height = text_bbox[3] - text_bbox[1]
        text_x = (width - text_width) // 2
        text_y = (height - text_height) // 2
        draw.text((text_x, text_y), text, fill='#999')
        
        # Save to bytes
        img_io = io.BytesIO()
        image.save(img_io, 'PNG')
        img_io.seek(0)
        
        return send_file(img_io, mimetype='image/png')
    except Exception as e:
        return str(e), 500

# Handle file uploads and generate blog content with improved randomization
@app.route('/upload_generate', methods=['POST'])
def upload_generate():
    try:
        print("[DEBUG] upload_generate called", flush=True)
        # Check if a file was uploaded
        if 'file' not in request.files:
            print("[DEBUG] No file part in request.files", flush=True)
            return jsonify({"error": "No file part"}), 400
        
        file = request.files['file']
        print(f"[DEBUG] file object: {file}", flush=True)
        
        # If the user does not select a file, the browser submits an empty file
        if file.filename == '':
            print("[DEBUG] No selected file (empty filename)", flush=True)
            return jsonify({"error": "No selected file"}), 400
        print(f"[DEBUG] Uploaded filename: {file.filename}", flush=True)
        print(f"[DEBUG] File content type: {file.content_type}", flush=True)
        
        # Create strong randomization to ensure content variation
        timestamp = int(time.time())
        random_id = uuid.uuid4().hex
        # Create a unique seed for this specific generation
        unique_seed = f"{timestamp}_{random_id}_{random.randint(1000000, 9999999)}"
        random.seed(unique_seed)
        
        # Force content variation with strong randomization factors
        variation_factors = {
            "timestamp": timestamp,
            "seed": random_id,
            "random_number": random.randint(1000, 9999),
            "content_variation": random.choice(["detailed", "concise", "analytical", "creative", "technical", "narrative", "instructional", "conversational"]),
            "tone_variation": random.choice(["formal", "casual", "enthusiastic", "professional", "conversational", "authoritative", "friendly", "inspirational"]),
            "structure_seed": random.randint(1, 1000),
            "perspective": random.choice(["first_person", "second_person", "third_person"]),
            "focus": random.choice(["problem_solution", "how_to", "list_based", "story_based", "comparison", "timeline"]),
            "unique_id": str(uuid.uuid4())
        }
        
        # Clear any cached results for this file name
        cache_key = f"cached_blog_{file.filename}"
        if cache_key in globals():
            print(f"[DEBUG] Clearing cache for {cache_key}", flush=True)
            del globals()[cache_key]
        
        # Initialize variables for content and images
        file_content = ""
        extracted_images = []
        image_index = 0  # Initialize image_index at the top level
        doc_text = ""
        
        if file:
            print("[DEBUG] File object exists, proceeding to save.", flush=True)
            # Save the file temporarily with a unique name to avoid caching
            upload_folder = "uploads"
            os.makedirs(upload_folder, exist_ok=True)
            unique_filename = f"{timestamp}_{random_id}_{secure_filename(file.filename)}"
            file_path = os.path.join(upload_folder, unique_filename)
            file.save(file_path)
            print(f"[DEBUG] File saved to {file_path}", flush=True)
            # Check if file exists and print its size
            import os
            if os.path.exists(file_path):
                file_size = os.path.getsize(file_path)
                print(f"[DEBUG] File exists after save. Size: {file_size} bytes", flush=True)
            else:
                print(f"[DEBUG] File does NOT exist after save!", flush=True)
            
            # Extract content and images based on file type
            ext = os.path.splitext(file.filename)[1].lower()
            print(f"[DEBUG] File extension: {ext}", flush=True)
            
            # Helper function to add an image to extracted_images
            def add_extracted_image(image_url, alt_text, source):
                nonlocal image_index
                extracted_images.append({
                    "url": image_url,
                    "alt": alt_text,
                    "source": source,
                    "unique_id": f"{timestamp}_{image_index}_{random.randint(1000, 9999)}"  # Add uniqueness
                })
                image_index += 1
            
            # RAG: Extract document text for supported types
            if ext in ['.txt', '.md', '.docx', '.pdf']:
                print(f"[DEBUG] Attempting to extract text from file (ext: {ext})", flush=True)
                doc_text = extract_text_from_file(file_path, ext)
                print(f"[DEBUG] Extracted doc_text (first 500 chars): {doc_text[:500]}", flush=True)
                if not doc_text:
                    print("[DEBUG] doc_text is empty after extraction!", flush=True)
            
            # Existing extraction logic for images and file_content
            if ext in ['.txt', '.md']:
                print("[DEBUG] Reading as plain text or markdown", flush=True)
                with open(file_path, 'r', encoding='utf-8') as f:
                    file_content = f.read()
                print(f"[DEBUG] file_content (first 500 chars): {file_content[:500]}", flush=True)
            elif ext in ['.docx']:
                print("[DEBUG] Reading as DOCX", flush=True)
                try:
                    import docx
                    doc = docx.Document(file_path)
                    
                    # Extract text content
                    paragraphs = []
                    for para in doc.paragraphs:
                        if para.text.strip():
                            paragraphs.append(para.text)
                    file_content = "\n\n".join(paragraphs)
                    print(f"[DEBUG] file_content from DOCX (first 500 chars): {file_content[:500]}", flush=True)
                    
                    # Extract images from docx with improved handling
                    for rel in doc.part.rels.values():
                        if "image" in rel.target_ref:
                            try:
                                image_data = rel.target_part.blob
                                image_ext = rel.target_ref.split(".")[-1] if "." in rel.target_ref else "png"
                                if image_ext not in ['png', 'jpg', 'jpeg', 'gif']:
                                    image_ext = 'png'  # Default to PNG if extension is not recognized
                                
                                # Save the image to the upload folder with unique name
                                image_filename = f"docx_image_{uuid.uuid4().hex[:8]}.{image_ext}"
                                image_path = os.path.join(upload_folder, image_filename)
                                
                                with open(image_path, "wb") as img_file:
                                    img_file.write(image_data)
                                
                                # Create a URL for the saved image
                                image_url = f"/uploads/{image_filename}"
                                add_extracted_image(image_url, f"Image {image_index + 1} from document", "docx")
                                print(f"[DEBUG] Extracted image from DOCX: {image_filename}", flush=True)
                            except Exception as img_err:
                                print(f"Error extracting image from DOCX: {str(img_err)}", flush=True)
                                
                except ImportError:
                    print("[DEBUG] DOCX support not available. Install python-docx package.", flush=True)
                    file_content = "DOCX support not available. Install python-docx package."
                except Exception as docx_err:
                    print(f"[DEBUG] Exception reading DOCX: {docx_err}", flush=True)
            elif ext in ['.pdf']:
                print("[DEBUG] Reading as PDF", flush=True)
                # Extract text and images from PDF with improved handling
                try:
                    pdf_text = []
                    from pypdf import PdfReader
                    pdf_reader = PdfReader(file_path)
                    print(f"[DEBUG] PDF has {len(pdf_reader.pages)} pages", flush=True)
                    # Extract text from each page
                    for page_num in range(len(pdf_reader.pages)):
                        page = pdf_reader.pages[page_num]
                        page_text = page.extract_text()
                        pdf_text.append(page_text)
                        print(f"[DEBUG] Extracted text from PDF page {page_num+1} (first 200 chars): {str(page_text)[:200]}", flush=True)
                    # Join all text
                    file_content = "\n\n".join([text for text in pdf_text if text])
                    print(f"[DEBUG] file_content from PDF (first 500 chars): {file_content[:500]}", flush=True)
                    # ... existing code for image extraction ...
                except Exception as pdf_err:
                    file_content = f"Error extracting PDF content: {str(pdf_err)}"
                    print(f"[DEBUG] Exception reading PDF: {pdf_err}", flush=True)
                # If we couldn't extract any content, use a fallback
                if not file_content or file_content.strip() == "":
                    print("[DEBUG] PDF file uploaded but text could not be extracted.", flush=True)
                    file_content = f"PDF file {file.filename} was uploaded but text could not be extracted."
            elif ext in ['.jpg', '.jpeg', '.png', '.gif']:
                print("[DEBUG] File is an image, not extracting text.", flush=True)
                unique_filename = f"{uuid.uuid4().hex[:8]}_{file.filename}"
                unique_path = os.path.join(upload_folder, unique_filename)
                with open(file_path, 'rb') as src_file:
                    with open(unique_path, 'wb') as dst_file:
                        dst_file.write(src_file.read())
                add_extracted_image(
                    f"/uploads/{unique_filename}",
                    f"Uploaded image - {file.filename}",
                    "uploaded"
                )
                file_content = f"Content for blog with image {file.filename}"
            else:
                print(f"[DEBUG] File type {ext} is not supported for content extraction.", flush=True)
                file_content = f"File type {ext} is not supported for content extraction."
        else:
            print("[DEBUG] No file object after initial check!", flush=True)
        
        # Get template ID from the form
        template_id = request.form.get('template_id', 'tech')
        
        # Load the template
        template = load_template(template_id)
        
        # Create a modified copy to avoid altering the original
        modified_template = copy.deepcopy(template)
        
        # Initialize metadata if not present
        if "metadata" not in modified_template:
            modified_template["metadata"] = {}
        
        # Add randomization metadata to force content variation
        modified_template["metadata"]["variation_factors"] = variation_factors
        modified_template["metadata"]["generation_time"] = timestamp
        modified_template["metadata"]["unique_id"] = random_id
        
        # Add a version number to avoid caching
        modified_template["version"] = f"v{timestamp}-{random.randint(1000, 9999)}"
        
        # Use Gemini model to generate content
        if GEMINI_API_KEY:
            try:
                genai.configure(api_key=GEMINI_API_KEY)
                model = genai.GenerativeModel(model_name="gemma-3n-e4b-it")
                
                # Create a prompt that forces variation based on the randomization factors
                blog_prompt = f"""Write a completely unique and original blog post about: 
                
                {file_content[:1000] if file_content else file.filename}
                
                IMPORTANT VARIATION INSTRUCTIONS:
                - This is generation attempt #{timestamp % 1000} with unique ID {random_id}
                - Use a {variation_factors['tone_variation']} tone and {variation_factors['content_variation']} style
                - Structure this as a {variation_factors['focus']} type of content
                - Write from a {variation_factors['perspective']} perspective
                - Structure seed: {variation_factors['structure_seed']}
                - Make this blog COMPLETELY DIFFERENT from any previous generations on this topic
                
                Format as JSON with the following structure:
                {{
                    "title": "A catchy title for the blog",
                    "content": [
                        {{"type": "text", "content": "Introduction paragraph"}},
                        {{"type": "heading", "content": "First Section Heading"}},
                        {{"type": "text", "content": "First section content"}},
                        {{"type": "image", "alt": "Description of an image that could go here"}}
                    ]
                }}
                
                Return ONLY valid JSON without any extra text, markdown formatting or explanation.
                
                The blog should have at least 5-6 sections with headings and 3-4 image placeholders.
                """
                
                # Make the API call
                response = model.generate_content(blog_prompt)
                response_text = response.text
                
                # Clean up the response to ensure it's valid JSON
                try:
                    # Remove any markdown formatting or text outside the JSON
                    if "```json" in response_text:
                        json_text = response_text.split("```json")[1].split("```")[0].strip()
                    elif "```" in response_text:
                        json_text = response_text.split("```")[1].strip()
                    else:
                        json_text = response_text.strip()
                    
                    # Parse the JSON content
                    content_updates = json.loads(json_text)
                    
                    # Update the template with the generated content
                    if "title" in content_updates:
                        modified_template["title"] = content_updates["title"]
                    
                    # Process content items and add in images
                    if "content" in content_updates and isinstance(content_updates["content"], list):
                        # Handle image placement
                        new_content = []
                        image_count = 0
                        
                        # Check if the file is an image type
                        is_image_upload = file.content_type and file.content_type.startswith('image/')
                        
                        # Only use user's uploaded image if it's an image
                        if is_image_upload:
                            # Create a path for the uploaded image
                            image_filename = f"user_upload_{timestamp}_{random.randint(1000, 9999)}.jpg"
                            image_path = os.path.join(upload_folder, image_filename)
                            
                            # Try to save a copy of the image
                            try:
                                image = Image.open(file_path)
                                image.save(image_path)
                                
                                # Add the user's image as the first image
                                user_image_url = f"/uploads/{image_filename}"
                                add_extracted_image(user_image_url, "User uploaded image", "user_upload")
                            except Exception as img_err:
                                print(f"Error saving user image: {str(img_err)}")
                        
                        # Process all content items
                        for item in content_updates["content"]:
                            # Clean the content to remove any timestamp references
                            if item.get("type") == "text" and "content" in item:
                                item["content"] = item["content"].replace("Generated content section", "")
                                item["content"] = re.sub(r"with timestamp \d+", "", item["content"])
                                item["content"] = re.sub(r"Generated Section \d+ - \d+", "", item["content"])
                            
                            if item.get("type") == "heading" and "content" in item:
                                item["content"] = item["content"].replace("Generated content section", "")
                                item["content"] = re.sub(r"with timestamp \d+", "", item["content"])
                                item["content"] = re.sub(r"Generated Section \d+ - \d+", "", item["content"])
                            
                            # Add the item to the new content
                            new_content.append(item)
                            # If this is an image item and we have extracted images, use them
                            if item.get("type") == "image" and image_count < len(extracted_images):
                                item["url"] = extracted_images[image_count]["url"]
                                item["unique_id"] = extracted_images[image_count]["unique_id"]
                                image_count += 1
                            
                            modified_template["content"] = new_content
                            try:
                                # Save the template
                                save_template(modified_template)
                            except Exception as save_err:
                                print(f"Error saving template: {str(save_err)}")
                                return jsonify({"error": f"Error saving template: {str(save_err)}"}), 500
                except Exception as json_err:
                    print(f"Error parsing generated content: {str(json_err)}")
                    return jsonify({"error": f"Error parsing generated content: {str(json_err)}"}), 500
                except Exception as gemini_err:
                    print(f"Error generating content with Gemini: {str(gemini_err)}")
                    return jsonify({"error": f"Error generating content: {str(gemini_err)}"}), 500
            except Exception as gemini_err:
                print(f"Error generating content with Gemini: {str(gemini_err)}")
                return jsonify({"error": f"Error generating content: {str(gemini_err)}"}), 500
        
        # Save the generated blog with a timestamp in the filename to avoid caching
        unique_id = f"{template_id}_{timestamp}_{random.randint(1000, 9999)}"
        modified_template["id"] = unique_id
        
        # Save updated template
        save_template(modified_template)
        
        # Return the generated template
        return jsonify({
            "response": f"Successfully generated unique blog based on your uploaded file.",
            "updatedContent": generate_html_preview(modified_template),
            "template": modified_template,
            "templateId": unique_id,
            "timestamp": timestamp,  # Include timestamp for client-side cache busting
            "randomFactor": random.randint(1000000, 9999999)  # Additional randomization
        })
        
    except Exception as e:
        print(f"Error in upload_generate: {str(e)}", flush=True)
        return jsonify({"error": f"Error processing your file: {str(e)}"}), 500

# Helper function to split content into chunks for the blog
def split_content_into_chunks(content, min_chunk_size=200, max_chunk_size=500):
    """Split content into reasonably sized chunks for better readability"""
    # If content is shorter than min_chunk_size, return as is
    if len(content) <= min_chunk_size:
        return [content]
    
    chunks = []
    current_chunk = ""
    sentences = content.split('. ')
    
    for sentence in sentences:
        # Add period back if it was removed by split
        if not sentence.endswith('.'):
            sentence += '.'
        
        # If adding this sentence would make chunk too long, save current chunk and start new one
        if len(current_chunk) + len(sentence) > max_chunk_size and len(current_chunk) >= min_chunk_size:
            chunks.append(current_chunk.strip())
            current_chunk = sentence + ' '
        else:
            current_chunk += sentence + ' '
    
    # Add any remaining content
    if current_chunk:
        chunks.append(current_chunk.strip())
    
    return chunks

# Extract a heading from text
def extract_heading(text, max_length=60):
    """Extract a heading from text, ensuring it's not too long"""
    # Split into sentences
    sentences = text.split('. ')
    
    for sentence in sentences:
        # Clean up the sentence
        heading = sentence.strip()
        
        # Remove any markdown or special characters
        heading = re.sub(r'[#*_]', '', heading)
        
        # If it's a good length, use it
        if 20 <= len(heading) <= max_length:
            return heading
        
        # If it's too long, try to truncate at a word boundary
        elif len(heading) > max_length:
            words = heading.split()
            truncated = ''
            for word in words:
                if len(truncated) + len(word) + 1 <= max_length - 3:  # Leave room for ellipsis
                    truncated += word + ' '
                else:
                    break
            return truncated.strip() + '...'
    
    # If we couldn't find a good sentence, use the first one truncated
   
    if sentences:
       
        heading = sentences[0].strip()
        if len(heading) > max_length:
            return heading[:max_length-3] + '...'
        return heading
    
    return "New Section"  # Fallback if no text provided

# Extract a quote from text
def extract_quote(text, min_length=30, max_length=150):
    """Extract a quote from text that's neither too short nor too long"""
    # Split into sentences
    sentences = text.split('. ')
    
    for sentence in sentences:
        # Clean up the sentence
        quote = sentence.strip()
        
        # Remove any markdown or special characters
        quote = re.sub(r'[#*_]', '', quote)
        
        # Check if it's a good length for a quote
        if min_length <= len(quote) <= max_length:
            return quote
    
    # If we couldn't find a perfect quote, try to combine short sentences or truncate long ones
    current_quote = ""
    for sentence in sentences:
        quote = sentence.strip()
        quote = re.sub(r'[#*_]', '', quote)
        
        if len(current_quote) + len(quote) <= max_length:
            current_quote += quote + '. '
        
        if len(current_quote) >= min_length:
            return current_quote.strip()
    
    # If still no good quote found, take the first sentence and truncate if needed
    if sentences:
        quote = sentences[0].strip()
        quote = re.sub(r'[#*_]', '', quote)
        if len(quote) > max_length:
            return quote[:max_length-3] + '...'
        return quote
    
    return None  # Return None if no suitable quote found

# Extract statistical values from text (numbers, percentages, currency, etc.)
def extract_stat_value(text):
    """Extract a statistical value from text"""
    # Look for numbers with % or common units
    percentage_match = re.search(r'\b(\d+(?:\.\d+)?%)\b', text)
    if percentage_match:
        return percentage_match.group(1)
    
    # Look for currency values
    currency_match = re.search(r'\$\s*(\d+(?:,\d{3})*(?:\.\d{2})?(?:\s*[Kk])?(?:\s*[Mm])?(?:\s*[Bb])?)', text)
    if currency_match:
        return currency_match.group(0)
    
    # Look for numbers with K/M/B suffix
    scale_match = re.search(r'\b(\d+(?:\.\d+)?\s*[KkMmBb])\b', text)
    if scale_match:
        return scale_match.group(1)
    
    # Look for simple numbers
    number_match = re.search(r'\b(\d+(?:,\d{3})*(?:\.\d+)?)\b', text)
    if number_match:
        return number_match.group(1)
    
    return None

# Extract a short piece of text for descriptions
def extract_short_text(text, max_length=100):
    """Extract a short snippet of text, ending at a natural break"""
    if len(text) <= max_length:
        return text
    
    # Try to find a sentence break
    shortened = text[:max_length]
    last_period = shortened.rfind('.')
    if last_period > max_length * 0.5:  # Only use sentence break if it's not too short
        return text[:last_period + 1]
    
    # Try to find a word break
    last_space = shortened.rfind(' ')
    if last_space > 0:
        return text[:last_space] + '...'
    
    # If no good breaks found, just truncate
    return shortened + '...'

# Function to fetch images from Unsplash API
def fetch_online_image(query, fallback_seed=None):
    """Fetch an image URL from Unsplash based on the query"""
    try:
        if not UNSPLASH_API_KEY:
            return generate_placeholder_image(fallback_seed)
            
        headers = {
            'Authorization': f'Client-ID {UNSPLASH_API_KEY}'
        }
        
        # Clean and encode the query
        clean_query = query.replace('#', '').strip()
        encoded_query = requests.utils.quote(clean_query)
        
        response = requests.get(
            f'https://api.unsplash.com/search/photos?query={encoded_query}&per_page=1',
            headers=headers
        )
        
        if response.status_code == 200:
            data = response.json()
            if data['results']:
                return data['results'][0]['urls']['regular']
        
        # If no results or error, fall back to placeholder
        return generate_placeholder_image(fallback_seed)
        
    except Exception as e:
        print(f"Error fetching image: {str(e)}")
        return generate_placeholder_image(fallback_seed)

# Function to generate a placeholder image with more variety and improved quality
def generate_placeholder_image(seed=None, alt_text="Generated image"):
    """
    Generate a high-quality placeholder image when online images are unavailable.
    Creates a visually appealing gradient with text overlay for the given seed.
    
    Args:
        seed: Numeric seed for reproducible generation
        alt_text: Alt text to describe the image
    
    Returns:
        dict with url and alt text
    """
    # Generate a seed if not provided
    if seed is None:
        seed = random.randint(10000, 99999)
    elif isinstance(seed, str) and seed.isdigit():
        seed = int(seed)
    
    # Seed the random generator for reproducible results
    random.seed(seed)
    
    # Generate higher quality images with consistent dimensions for blog posts
    dimensions = [
        (800, 400),   # 2:1 aspect ratio (good for headers)
        (900, 600),   # 3:2 aspect ratio (standard photo)
        (850, 500),   # 17:10 aspect ratio (widescreen)
        (700, 525),   # 4:3 aspect ratio (standard)
        (1000, 500)   # 2:1 aspect ratio (panoramic)
    ]
    
    # Pick one of the predefined dimensions
    width, height = random.choice(dimensions)
    
    # Parse any descriptive text from alt_text to create more relevant placeholders
    if isinstance(alt_text, str) and len(alt_text) > 10:
        # Extract meaningful words from alt text (remove common filler words)
        words = alt_text.split()
        important_words = [w for w in words if len(w) > 3 and w.lower() not in 
                          ('this', 'that', 'with', 'from', 'image', 'picture', 'photo', 'generated')]
        
        # Use up to 3 important words as the image descriptor
        if important_words:
            descriptor = ' '.join(important_words[:3])
            if len(descriptor) > 50:
                descriptor = descriptor[:47] + '...'
        else:
            descriptor = alt_text[:50] + ('...' if len(alt_text) > 50 else '')
    else:
        descriptor = "Generated image"
    
    # Create a more visually appealing URL with context
    image_type = random.choice(['abstract', 'gradient', 'pattern', 'geometric'])
    
    # Clean descriptor for URL (remove special chars, spaces to dashes)
    clean_descriptor = re.sub(r'[^\w\s-]', '', descriptor).strip().lower()
    clean_descriptor = re.sub(r'[\s-]+', '-', clean_descriptor)
    
    # Truncate if too long
    if len(clean_descriptor) > 30:
        clean_descriptor = clean_descriptor[:30]
    
    # Generate a nice URL format
    url_params = f"?type={image_type}&seed={seed}&text={clean_descriptor}"
    
    return {
        "url": f"/api/placeholder/{width}/{height}{url_params}",
        "alt": descriptor
    }

# Set up a route to serve files from the uploads directory
@app.route('/uploads/<filename>')
def uploaded_file(filename):
    return send_file(os.path.join(UPLOAD_FOLDER, filename))

def extract_keywords(text, num_keywords=5):
    """
    Extract relevant keywords from text to use as search terms for images.
    Improved to handle PDF content better by focusing on significant terms.
    
    Args:
        text: The text to extract keywords from
        num_keywords: Number of keywords to extract
        
    Returns:
        List of keyword strings
    """
    if not text or len(text) < 50:
        return ["document", "business", "professional", "concept", "abstract"]
        
    # Expanded stop words list to filter out common terms
    stop_words = {
        'a', 'an', 'the', 'and', 'or', 'but', 'if', 'because', 'as', 'what', 
        'when', 'where', 'how', 'is', 'are', 'was', 'were', 'be', 'been', 'to', 
        'of', 'for', 'by', 'with', 'about', 'against', 'between', 'into', 'through',
        'during', 'before', 'after', 'above', 'below', 'from', 'up', 'down', 'in',
        'out', 'on', 'off', 'over', 'under', 'again', 'further', 'then', 'once',
        'here', 'there', 'all', 'any', 'both', 'each', 'few', 'more', 'most', 'other',
        'some', 'such', 'no', 'nor', 'not', 'only', 'own', 'same', 'so', 'than', 'too',
        'very', 's', 't', 'can', 'will', 'just', 'don', 'should', 'now', 'also',
        'figure', 'fig', 'table', 'copyright', 'page', 'chapter', 'et', 'al',
        'this', 'that', 'these', 'those', 'would', 'could', 'may', 'might'
    }
    
    # Clean and normalize text - remove special characters and extra spaces
    # For PDF text specifically, look for patterns like multiple spaces, page numbers, etc.
    cleaned_text = re.sub(r'[^\w\s]', ' ', text.lower())
    cleaned_text = re.sub(r'\s+', ' ', cleaned_text).strip()
    
    # Handle PDF page numbers and headers - common patterns in PDFs
    cleaned_text = re.sub(r'\d+\s*of\s*\d+', '', cleaned_text)  # Remove "Page X of Y"
    cleaned_text = re.sub(r'^\d+$', '', cleaned_text, flags=re.MULTILINE)  # Remove standalone numbers
    cleaned_text = re.sub(r'www\.\w+\.\w+', '', cleaned_text)  # Remove URLs
    
    # Split into words
    words = cleaned_text.split()
    
    # Filter out stop words and words that are too short or numerical
    filtered_words = []
    for word in words:
        # Skip if it's a stop word, too short, or just a number
        if (word not in stop_words and 
            len(word) > 3 and 
            not word.isdigit() and
            not re.match(r'^\d+\w?$', word)):  # Skip patterns like "10th", "2nd", etc.
            filtered_words.append(word)
    
    # Count word frequencies with more weight given to longer words
    word_counts = {}
    for word in filtered_words:
        # Give more weight to longer words - they're often more significant
        weight = min(len(word) - 2, 3)  # Words of length 5+ get weight boost
        if word in word_counts:
            word_counts[word] += 1 + (weight * 0.2)
        else:
            word_counts[word] = 1 + (weight * 0.2)
    
    # Boost for title-case words (likely proper nouns, more significant)
    title_case_pattern = re.compile(r'^[A-Z][a-z]+$')
    for word in text.split():
        if title_case_pattern.match(word) and len(word) > 3:
            word_lower = word.lower()
            if word_lower in word_counts:
                word_counts[word_lower] *= 1.5  # 50% boost for proper nouns
    
    # Get most common words
    sorted_words = sorted(word_counts.items(), key=lambda x: x[1], reverse=True)
    top_keywords = [word for word, count in sorted_words[:num_keywords*2]]
    
    # Prioritize specific contextual keywords that work well for image search
    contextual_keywords = [
        'technology', 'business', 'nature', 'people', 'science', 'travel', 
        'finance', 'health', 'education', 'food', 'industry', 'architecture',
        'medical', 'digital', 'office', 'wildlife', 'landscape', 'urban',
        'research', 'innovation', 'data', 'analysis', 'strategy', 'leadership',
        'development', 'environment', 'communication', 'marketing', 'design'
    ]
    
    # Push contextual keywords that appear in our list to the front
    # These tend to generate better image results
    prioritized_keywords = []
    remaining_keywords = []
    
    for keyword in top_keywords:
        if keyword in contextual_keywords:
            prioritized_keywords.append(keyword)
        else:
            remaining_keywords.append(keyword)
    
    result_keywords = prioritized_keywords + remaining_keywords
    
    # If we don't have enough keywords, add some generic ones
    generic_keywords = [
        'document', 'professional', 'business', 'concept', 'information',
        'technology', 'analysis', 'research', 'data', 'report'
    ]
    
    while len(result_keywords) < num_keywords:
        for keyword in generic_keywords:
            if keyword not in result_keywords:
                result_keywords.append(keyword)
                if len(result_keywords) >= num_keywords:
                    break
    
    # Return unique keywords with the highest counts
    return list(dict.fromkeys(result_keywords))[:num_keywords]

@app.route('/debug_log', methods=['POST'])
def debug_log():
    data = request.get_json()
    msg = data.get('msg', '')
    print(f'[FRONTEND DEBUG] {msg}', file=sys.stderr, flush=True)
    return jsonify({'status': 'ok'})

@app.route('/query_chatbot', methods=['POST'])
def query_chatbot_route():
    try:
        data = request.get_json(force=True, silent=True)
        if not data:
            return jsonify({"error": "No data provided"}), 400
        user_message = data.get('message', '').strip()
        template_id = data.get('template_id', 'tech')
        if not user_message:
            return jsonify({"error": "No message provided"}), 400
        response_text, updated_template = query_chatbot(user_message, template_id)
        # Debug output for diagnosis
        print("[DEBUG] /query_chatbot: user_message={}, template_id={}, response_text={}, updated_template={}".format(
            user_message, template_id, response_text, 'present' if updated_template is not None else 'None'))
        # Always return a valid response, even if the chatbot fails
        if not response_text or (isinstance(response_text, str) and response_text.lower().startswith('error')):
            import traceback
            print("[ERROR] query_chatbot failed. user_message: {}, template_id: {}, response_text: {}".format(
                user_message, template_id, response_text))
            traceback.print_stack()
            response_text = "Sorry, I encountered an error processing your request. Please try again."
        response = {"response": response_text}
        if updated_template is not None:
            response["template"] = updated_template
            response["templateId"] = template_id
        return jsonify(response)
    except Exception as e:
        print("Error in /query_chatbot route: {}".format(str(e)))
        return jsonify({"error": "Internal server error. Please try again later."}), 500
    
@app.route('/generate_blog', methods=['POST'])
def generate_blog():
    print("=== TEST: generate_blog endpoint called ===", flush=True)
    title = request.form.get('title', 'Untitled Blog')
    description = request.form.get('description', '')
    images = request.files.getlist('images')
    template_id = request.form.get('template_id', 'tech')
    template = load_template(template_id)

    # Debug: Check uploaded images
    print(f"[DEBUG] Number of images uploaded: {len(images)}", flush=True)
    image_captions = []
    for idx, img in enumerate(images):
        print(f"[DEBUG] Image {idx}: filename={img.filename}, content_type={img.content_type}", flush=True)
        # Save image to disk and check existence/size
        if img.filename:
            upload_folder = "uploads"
            import os
            os.makedirs(upload_folder, exist_ok=True)
            image_path = os.path.join(upload_folder, img.filename)
            img.save(image_path)
            print(f"[DEBUG] Saved image to {image_path}", flush=True)
            if os.path.exists(image_path):
                file_size = os.path.getsize(image_path)
                print(f"[DEBUG] Image file exists after save. Size: {file_size} bytes", flush=True)
                # Hugging Face captioning
                try:
                    with open(image_path, "rb") as f:
                        img_bytes = f.read()
                    caption = get_image_caption_hf(img_bytes)
                    print(f"[DEBUG] Hugging Face caption for {img.filename}: {caption}")
                    image_captions.append(caption)
                except Exception as e:
                    print(f"[DEBUG] Error reading image for Hugging Face: {e}")
                    image_captions.append("Image uploaded by user")
            else:
                print(f"[DEBUG] Image file does NOT exist after save!", flush=True)
                image_captions.append("Image uploaded by user")
        else:
            image_captions.append("Image uploaded by user")

    # Debug: Check uploaded content files (PDF, DOC, DOCX)
    content_files = request.files.getlist('content_files')
    print(f"[DEBUG] Number of content files uploaded: {len(content_files)}", flush=True)
    for idx, cfile in enumerate(content_files):
        print(f"[DEBUG] Content file {idx}: filename={cfile.filename}, content_type={cfile.content_type}", flush=True)
        if cfile.filename:
            upload_folder = "uploads"
            import os
            os.makedirs(upload_folder, exist_ok=True)
            content_path = os.path.join(upload_folder, cfile.filename)
            cfile.save(content_path)
            print(f"[DEBUG] Saved content file to {content_path}", flush=True)
            if os.path.exists(content_path):
                file_size = os.path.getsize(content_path)
                print(f"[DEBUG] Content file exists after save. Size: {file_size} bytes", flush=True)
                # Try to extract text if PDF, DOCX, TXT, or MD
                ext = os.path.splitext(cfile.filename)[1].lower()
                if ext in ['.pdf', '.docx', '.txt', '.md']:
                    print(f"[DEBUG] Attempting to extract text from {content_path} (ext: {ext})", flush=True)
                    doc_text = extract_text_from_file(content_path, ext)
                    if doc_text:
                        print(f"[DEBUG] Extracted text (first 500 chars): {doc_text[:500]}", flush=True)
                    else:
                        print(f"[DEBUG] No text extracted from {content_path}", flush=True)
            else:
                print(f"[DEBUG] Content file does NOT exist after save!", flush=True)

    # Use AI model to generate blog content based on title, description, and template
    blog_prompt = f"""Write a complete, original blog post with the following details:
Title: {title}
Description: {description}
Template Type: {template.get('type', template_id)}

- The blog should be well-structured with an introduction, multiple sections, and a conclusion.
- Add relevant headings and subheadings.
- Insert at least 2-3 image placeholders with descriptive alt text.
- Use the style, tone, and structure of a {template.get('type', template_id)} blog.
- Uploaded image captions: {"; ".join(image_captions) if image_captions else "Image uploaded by user"}
- Return the blog as a JSON object with this structure:
{{
  "title": "...",
  "content": [
    {{"type": "text", "content": "..."}},
    {{"type": "heading", "content": "..."}},
    {{"type": "image", "alt": "..."}}
  ]
}}
Return only valid JSON."""

    blog_json = None
    ai_error = None

    # Try to use Gemini if available, else fallback to Groq
    if GEMINI_API_KEY:
        try:
            genai.configure(api_key=GEMINI_API_KEY)
            model = genai.GenerativeModel(model_name="gemma-3n-e4b-it")
            response = model.generate_content(blog_prompt)
            response_text = response.text
            # Extract JSON from response
            if "```json" in response_text:
                json_text = response_text.split("```json")[1].split("```")[0].strip()
            elif "```" in response_text:
                json_text = response_text.split("```")[1].strip()
            else:
                json_text = response_text.strip()
            blog_json = json.loads(json_text)
        except Exception as e:
            ai_error = str(e)
            print("[WARNING] Gemini AI error:", e, flush=True)
    elif GROQ_API_KEY:
        try:
            client = groq.Groq(api_key=GROQ_API_KEY)
            response = client.chat.completions.create(
                model="llama3-8b-8192",
                messages=[
                    {"role": "system", "content": "You are a helpful AI that generates blog posts in JSON as described."},
                    {"role": "user", "content": blog_prompt}
                ],
                temperature=0.9,
                max_tokens=1500
            )
            response_text = response.choices[0].message.content
            if "```json" in response_text:
                json_text = response_text.split("```json")[1].split("```")[0].strip()
            elif "```" in response_text:
                json_text = response_text.split("```")[1].strip()
            else:
                json_text = response_text.strip()
            blog_json = json.loads(json_text)
            print("[DEBUG] AI blog JSON (Groq):", blog_json, flush=True)
        except Exception as e:
            ai_error = str(e)
            print("[WARNING] Groq AI error:", e, flush=True)
    else:
        print("[WARNING] No AI API key set. Falling back to user content.", flush=True)
        ai_error = "No AI API key set. Please set GEMINI_API_KEY or GROQ_API_KEY."

    # Fallback if AI fails
    if not blog_json:
        blog_json = {
            "title": title,
            "content": [
                {"type": "text", "content": description or "No description provided."}
            ]
        }

    # Convert images to data URLs (if any uploaded)
    image_tags = ""
    for img in images:
        img_bytes = img.read()
        mime = img.mimetype or "image/png"
        b64 = base64.b64encode(img_bytes).decode('utf-8')
        image_tags += f'<img src="data:{mime};base64,{b64}" style="max-width:100%;margin:20px 0;border-radius:12px;box-shadow:0 2px 12px #0002;">\n'

    # Generate HTML from blog_json
    html_content = ""
    for section in blog_json.get("content", []):
        if section.get("type") == "text":
            html_content += f'<div class="blog-section"><p>{section.get("content", "")}</p></div>\n'
        elif section.get("type") == "heading":
            html_content += f'<div class="blog-section"><h2>{section.get("content", "")}</h2></div>\n'
        elif section.get("type") == "image":
            # Use placeholder image with alt text
            alt = section.get("alt", "Blog image")
            html_content += f'<div class="blog-section"><img src="/api/placeholder/800/400?text={alt.replace(" ", "+")}" alt="{alt}" style="max-width:100%;margin:20px 0;border-radius:12px;box-shadow:0 2px 12px #0002;"></div>\n'

    # Add uploaded images at the end
    html_content += image_tags

    # Self-contained HTML template
    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <title>{blog_json.get("title", title)}</title>
    <style>
        body {{
            font-family: 'Roboto', Arial, sans-serif;
            background: #181a1b;
            color: #e3e6ea;
            margin: 0;
            padding: 0 0 40px 0;
        }}
        .container {{
            max-width: 700px;
            margin: 40px auto;
            background: #23262b;
            border-radius: 16px;
            box-shadow: 0 4px 32px #000a, 0 1.5px 8px #0004;
            padding: 32px 28px;
        }}
        h1 {{
            color: #4f8cff;
            font-size: 2.2rem;
            margin-bottom: 18px;
        }}
        .blog-section {{
            font-size: 1.13rem;
            line-height: 1.7;
            color: #e3e6ea;
            margin-bottom: 24px;
        }}
        img {{
            display: block;
            margin: 24px auto;
            max-width: 100%;
            border-radius: 12px;
            box-shadow: 0 2px 12px #0002;
        }}
        @media (max-width: 800px) {{
            .container {{ padding: 18px 5vw; }}
        }}
    </style>
</head>
<body>
    <div class="container">
        <h1>{blog_json.get("title", title)}</h1>
        {html_content}
    </div>
</body>
</html>
"""

    # Return as downloadable file
    mem = BytesIO()
    mem.write(html.encode('utf-8'))
    mem.seek(0)
    return send_file(
        mem,
        mimetype='text/html',
        as_attachment=True,
        download_name='blog.html'
    )

# Helper function to extract text from uploaded files (PDF, DOCX, TXT, MD)
def extract_text_from_file(file_path, ext):
    if ext == '.pdf':
        try:
            from pypdf import PdfReader
            pdf_reader = PdfReader(file_path)
            pdf_text = []
            for page in pdf_reader.pages:
                pdf_text.append(page.extract_text())
            return "\n\n".join([t for t in pdf_text if t])
        except Exception as e:
            return ""
    elif ext == '.docx':
        try:
            import docx
            doc = docx.Document(file_path)
            return "\n\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        except Exception as e:
            return ""
    elif ext in ['.txt', '.md']:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    return ""

def extract_texts_from_uploaded_files(content_files, upload_dir="uploads"):
    """Extract and concatenate text from all uploaded content files."""
    texts = []
    for cfile in content_files:
        if cfile and cfile.filename:
            ext = os.path.splitext(cfile.filename)[1].lower()
            temp_path = os.path.join(upload_dir, secure_filename(cfile.filename))
            cfile.save(temp_path)
            print(f"[DEBUG] Saved uploaded file to {temp_path}")
            text = extract_text_from_file(temp_path, ext)
            print(f"[DEBUG] Extracted text from {temp_path} (first 300 chars): {text[:300]}")
            if text:
                texts.append(text)
            else:
                print(f"[WARNING] No text extracted from {temp_path}")
    full_text = "\n\n".join(texts)
    print(f"[DEBUG] Full concatenated extracted text (first 1000 chars): {full_text[:1000]}")
    return full_text

@app.route('/generate_social_posts', methods=['POST'])
def generate_social_posts():
    try:
        print("[DEBUG] /generate_social_posts endpoint called.")
        title = request.form.get('title', '').strip()
        description = request.form.get('description', '').strip()
        hashtags = request.form.get('hashtags', '').strip()
        generate_by_ai = request.form.get('generate_by_ai', 'false').lower() == 'true'
        platforms = request.form.get('platforms', '[]')
        import json as _json
        try:
            platforms = _json.loads(platforms)
            print(f"[DEBUG] Parsed platforms: {platforms}")
        except Exception as e:
            print(f"[WARNING] Could not parse platforms: {e}")
            platforms = []
        images = request.files.getlist('images')
        print(f"[DEBUG] Number of images received: {len(images)}")
        content_files = request.files.getlist('content_files')
        print(f"[DEBUG] Number of content files received: {len(content_files)}")
        rag_context = extract_texts_from_uploaded_files(content_files) if content_files else ""
        if rag_context:
            print(f"[DEBUG] FINAL rag_context to AI (first 1000 chars): {rag_context[:1000]}")
            debug_context_source = 'document'
        elif description:
            rag_context = description
            print("[DEBUG] No document text extracted, falling back to description.")
            print(f"[DEBUG] FINAL rag_context to AI (first 1000 chars): {rag_context[:1000]}")
            debug_context_source = 'description'
        else:
            rag_context = 'No content provided.'
            print("[WARNING] No document or description provided. Using fallback content.")
            debug_context_source = 'fallback'
            print(f"[DEBUG] FINAL rag_context to AI (first 1000 chars): {rag_context[:1000]}")

        posts_html = ''
        blog_html = ''
        platform_styles = {
            'instagram': {'icon': 'fab fa-instagram', 'color': '#E1306C', 'label': 'Instagram'},
            'facebook': {'icon': 'fab fa-facebook-f', 'color': '#1877F3', 'label': 'Facebook'},
            'twitter': {'icon': 'fab fa-twitter', 'color': '#1DA1F2', 'label': 'Twitter'},
            'linkedin': {'icon': 'fab fa-linkedin-in', 'color': '#0077B5', 'label': 'LinkedIn'},
            'blog': {'icon': 'fas fa-blog', 'color': '#4f8cff', 'label': 'Blog'},
        }
        if generate_by_ai:
            print(f"[DEBUG] AI generation requested for social posts. Using context from: {debug_context_source}")
            ai_posts = {}
            for platform in platforms:
                if platform == 'blog':
                    continue
                style = platform_styles.get(platform, {})
                prompt = f"""
You are an expert social media copywriter.

Below is the content from a document uploaded by the user. Base the post entirely on this content. Do not use generic content unless it is present in the document.

Document content:
{rag_context}

Platform: {platform.capitalize()}
Title: {title}
Hashtags: {hashtags}

Instructions:
- Summarize or adapt the document content for a {platform} post.
- Do NOT invent or add unrelated content.
- Use the user's document as the main source.
- Add the hashtags at the end.
- Return only the post text (no markdown, no explanation).
"""
                print(f"[DEBUG] Prompt for {platform} (first 500 chars): {prompt[:500]}")
                ai_content = None
                if GEMINI_API_KEY:
                    try:
                        genai.configure(api_key=GEMINI_API_KEY)
                        model = genai.GenerativeModel(model_name="gemma-3n-e4b-it")
                        response = model.generate_content(prompt)
                        ai_content = response.text.strip()
                        print(f"[DEBUG] AI social post for {platform} (Gemini, first 300 chars): {ai_content[:300]}")
                    except Exception as e:
                        print(f"[WARNING] Gemini AI error for {platform}:", e)
                        ai_content = f"[AI Error: {str(e)}] {description}"
                elif GROQ_API_KEY:
                    try:
                        client = groq.Groq(api_key=GROQ_API_KEY)
                        response = client.chat.completions.create(
                            model="llama3-8b-8192",
                            messages=[
                                {"role": "system", "content": "You are a helpful AI that generates social media posts."},
                                {"role": "user", "content": prompt}
                            ],
                            temperature=0.8,
                            max_tokens=300
                        )
                        ai_content = response.choices[0].message.content.strip()
                        print(f"[DEBUG] AI social post for {platform} (Groq, first 300 chars): {ai_content[:300]}")
                    except Exception as e:
                        print(f"[WARNING] Groq AI error for {platform}:", e)
                        ai_content = f"[AI Error: {str(e)}] {description}"
                else:
                    print(f"[WARNING] No AI API key set. Falling back to user content for {platform}.")
                    ai_content = description
                ai_posts[platform] = ai_content
            for platform in platforms:
                if platform == 'blog':
                    continue
                style = platform_styles.get(platform, {})
                icon = style.get('icon', '')
                color = style.get('color', '#4f8cff')
                label = style.get('label', platform.capitalize())
                content = ai_posts.get(platform, description)
                content = clean_ai_content(content)
                hashtags_html = f'<div style="color:var(--accent-hover);margin-bottom:0.5em;">{hashtags}</div>' if hashtags.strip() else ''
                images_html = ''
                if not content and not hashtags_html and not images_html:
                    posts_html += f'''
                    <div class="post-card" style="background:var(--bg-card);padding:1.2em;border-radius:12px;margin-bottom:1.2em;">
                        <div style="font-weight:600;color:{color};margin-bottom:0.5em;text-transform:capitalize;">
                            <i class="{icon}"></i> {label}
                        </div>
                        <h3 style="margin:0.5em 0 0.3em 0;">{title}</h3>
                    </div>
                    '''
                else:
                    posts_html += f'''
                    <div class="post-card" style="background:var(--bg-card);padding:1.2em;border-radius:12px;margin-bottom:1.2em;">
                        <div style="font-weight:600;color:{color};margin-bottom:0.5em;text-transform:capitalize;">
                            <i class="{icon}"></i> {label}
                        </div>
                        {images_html}
                        <h3 style="margin:0.5em 0 0.3em 0;">{title}</h3>
                        {f'<p style="margin:0.3em 0 0.3em 0;">{content}</p>' if content else ''}
                        {hashtags_html}
                    </div>
                    '''
        else:
            print("[DEBUG] Not using AI for social posts (generate_by_ai is False).")
        return jsonify({'posts_html': posts_html, 'blog_html': blog_html})
    except Exception as e:
        print("[ERROR] Exception in generate_social_posts:", e)
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

def clean_ai_content(content):
    import re
    content = content.strip()
    content = re.sub(r'^```[a-zA-Z]*\n?', '', content)
    content = re.sub(r'```$', '', content)
    return content.strip()

# Helper: Hugging Face image captioning
import requests as _requests

def get_image_caption_hf(image_bytes):
    if not HUGGINGFACE_API_KEY:
        print("[DEBUG] Hugging Face API key not found in environment.")
        return "Image uploaded by user"
    api_url = "https://api-inference.huggingface.co/models/Salesforce/blip-image-captioning-base"
    headers = {"Authorization": f"Bearer {HUGGINGFACE_API_KEY}"}
    try:
        resp = _requests.post(api_url, headers=headers, files={"file": image_bytes})
        if resp.status_code == 200:
            result = resp.json()
            if isinstance(result, list) and result and 'generated_text' in result[0]:
                print(f"[DEBUG] Hugging Face caption: {result[0]['generated_text']}")
                return result[0]['generated_text']
            print(f"[DEBUG] Unexpected Hugging Face response: {result}")
            return "Image uploaded by user"
        else:
            print(f"[DEBUG] Hugging Face API error: {resp.status_code} {resp.text}")
            return "Image uploaded by user"
    except Exception as e:
        print(f"[DEBUG] Exception during Hugging Face API call: {e}")
        return "Image uploaded by user"

if __name__ == "__main__":
    app.run(debug=True)
