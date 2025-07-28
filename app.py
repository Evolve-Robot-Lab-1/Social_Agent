from flask import Flask, render_template, request, jsonify, url_for, send_file, redirect, send_from_directory
from io import BytesIO
import base64
import requests
import google.generativeai as genai
import groq
from dotenv import load_dotenv
import json
import os
import re
from PIL import Image, ImageDraw
import io
import random
import uuid
import time
import openai
import math
from werkzeug.utils import secure_filename
import copy
import traceback
import sys

# Add new imports for PDF support
try:
    from pypdf import PdfReader
except ImportError:
    print("pypdf not installed. PDF support will be limited.")
    print("To install: pip install pypdf")

# Try to import pdf2image for enhanced PDF support
try:
    from pdf2image import convert_from_path
except ImportError:
    print("pdf2image not installed. PDF image extraction will be limited.")
    print("To install: pip install pdf2image")
    print("\nIMPORTANT: pdf2image requires poppler to be installed:")
    print("1. Download poppler from: https://github.com/oschwartz10612/poppler-windows/releases/")
    print("2. Extract to C:\\poppler")
    print("3. Add C:\\poppler\\Library\\bin to your PATH environment variable")
    print("   You can do this by running: $env:PATH += \";C:\\poppler\\Library\\bin\"")
    print("   For permanent PATH addition, add it through System Properties > Environment Variables")

# Try to verify poppler is working
try:
    import os  # Ensure os is imported
    # Instead of pdf2image.poppler_path (which does not exist), just check if the poppler binary exists
    poppler_bin_path = os.path.join(os.path.dirname(__file__), 'poppler', 'poppler-23.11.0', 'Library', 'bin')
    if os.path.exists(poppler_bin_path):
        print(f"Poppler found at: {poppler_bin_path}")
    else:
        print("Poppler path not set or does not exist. PDF to image conversion may not work correctly.")
        print("You may need to specify the path manually in convert_from_path() calls.")
except Exception as e:
    print(f"Could not verify poppler installation: {str(e)}")
    print("PDF to image conversion may not work correctly.")

load_dotenv()

# Initialize Flask app
app = Flask(__name__)

# Configure upload and templates folders
UPLOAD_FOLDER = "uploads"
TEMPLATES_FOLDER = "static/templates"
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# Ensure directories exist
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(TEMPLATES_FOLDER, exist_ok=True)

# Load API Keys from Environment Variables
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
NEWS_API_KEY = os.getenv("NEWS_API_KEY")
UNSPLASH_API_KEY = os.getenv("UNSPLASH_API_KEY", "")  
HUGGINGFACE_API_KEY = os.getenv("HUGGINGFACE_API_KEY")

# Initialize Groq client
client = groq.Groq(api_key=GROQ_API_KEY)

# Template storage
TEMPLATES_FOLDER = "static/templates"
os.makedirs(TEMPLATES_FOLDER, exist_ok=True)

# Default tech blog template
tech_blog_template = {
    "id": "tech",
    "name": "Tech Blog",
    "title": "The Future of Artificial Intelligence and Machine Learning",
    "author": "Tech Insights Team",
    "date": "March 12, 2025",
    "content": [
        {
            "type": "text",
            "content": "Artificial intelligence continues to evolve at a remarkable pace, transforming industries and reshaping our daily lives. The convergence of improved algorithms, massive datasets, and unprecedented computing power has propelled AI from research laboratories into our everyday experiences, from voice assistants to autonomous vehicles, from fraud detection systems to medical diagnostics."
        },
        {
            "type": "image",
            "url": "https://images.unsplash.com/photo-1677442136019-21780ecad995",
            "alt": "AI Technology Visualization"
        },
        {
            "type": "heading",
            "content": "Recent Breakthroughs in AI"
        },
        {
            "type": "text",
            "content": "Recent breakthroughs in neural networks and machine learning have accelerated progress in fields like natural language processing and computer vision. Large language models now generate text virtually indistinguishable from human writing, while diffusion models create photorealistic images from text prompts with astonishing fidelity. These capabilities are not merely impressive technical achievements—they're fundamentally changing how creative work is produced and consumed."
        },
        {
            "type": "image",
            "url": "https://images.unsplash.com/photo-1620712943543-bcc4688e7485",
            "alt": "Neural network visualization"
        },
        {
            "type": "heading",
            "content": "Enterprise AI Adoption"
        },
        {
            "type": "text",
            "content": "As AI technologies mature, we're witnessing unprecedented enterprise adoption. Organizations are moving beyond experimentation to deeply integrate AI throughout their operations. Process automation, intelligent document processing, and predictive analytics are becoming standard components of the business toolkit. According to recent industry surveys, over 65% of enterprises have increased their AI investments substantially in the past year, with the greatest growth in retail, financial services, and healthcare sectors."
        },
        {
            "type": "image",
            "url": "https://images.unsplash.com/photo-1661956602868-6ae368943878",
            "alt": "Enterprise AI implementation"
        },
        {
            "type": "heading",
            "content": "Ethical Considerations and Responsible AI"
        },
        {
            "type": "text",
            "content": "As AI systems become more powerful, the ethical considerations surrounding their deployment grow increasingly complex. Issues of bias, fairness, transparency, and accountability have moved from academic discussions to board rooms. Leading organizations are developing comprehensive responsible AI frameworks that address these concerns while enabling innovation. Regulatory frameworks are beginning to emerge globally, with the EU's AI Act representing the most comprehensive attempt to classify AI applications according to risk and establish appropriate governance mechanisms."
        },
        {
            "type": "image",
            "url": "https://images.unsplash.com/photo-1668001027022-c9e377de649c",
            "alt": "AI ethics concept image"
        },
        {
            "type": "heading",
            "content": "The Future of Human-AI Collaboration"
        },
        {
            "type": "text",
            "content": "Perhaps the most promising direction for AI development is not autonomy but augmentation—systems designed to enhance human capabilities rather than replace them. We're seeing the emergence of 'centaur' models across industries, where humans and AI systems collaborate to achieve outcomes superior to what either could accomplish alone. In medicine, AI-enhanced diagnostic systems empower doctors to detect conditions earlier and with greater accuracy. In engineering, generative design tools suggest novel solutions that might never occur to human designers. These collaborative approaches address many of the concerns about displacement while maximizing the unique strengths of both human and artificial intelligence."
        },
        {
            "type": "image",
            "url": "https://images.unsplash.com/photo-1531746790731-6c087fecd65a",
            "alt": "Human and AI collaboration concept"
        }
    ],
    "image": "https://images.unsplash.com/photo-1485827404703-89b55fcc595e"
}

# Finance blog template with expanded content
finance_blog_template = {
    "id": "finance",
    "name": "Finance Blog",
    "title": "Economic Update: Market Expectations and Bond Trends for 2025",
    "author": "Financial Analysis Team",
    "date": "February 27, 2025",
    "content": [
        {
            "type": "text",
            "content": "As we navigate through 2025, the global economic landscape presents a complex interplay of market dynamics, technological disruption, and evolving investment strategies. Traditional financial models are being challenged by new paradigms, while the intersection of artificial intelligence and finance continues to reshape how we think about money, investments, and economic growth. Financial institutions worldwide are adapting to these changes, with many embracing digital transformation as a core strategic initiative rather than simply a technological upgrade."
        },
        {
            "type": "image",
            "url": "https://images.unsplash.com/photo-1611974789855-9c2a0a7236a3",
            "alt": "Modern financial district with skyscrapers"
        },
        {
            "type": "heading",
            "content": "Global Market Trends and Economic Indicators"
        },
        {
            "type": "text",
            "content": "Recent data from bond markets suggests elevated inflation expectations for the years ahead. The gap between standard U.S. Treasury securities and inflation-protected bonds, known as the breakeven rate, has shown significant movement. Current bond prices imply a 2.7% annual Consumer Price Index inflation for the next five years. This divergence reflects growing investor concern about inflationary pressures despite central bank assertions that recent price increases are transitory."
        },
        {
            "type": "image",
            "url": "https://images.unsplash.com/photo-1642543492481-44e81e3914a6",
            "alt": "Digital financial data visualization"
        },
        {
            "type": "text",
            "content": "The yield curve, another critical market indicator, has undergone meaningful shifts. The flattening of the curve in certain segments suggests changing market expectations about future economic growth and monetary policy. Specifically, the 10-year Treasury yield stabilizing around 4.2% indicates a market consensus forming around medium-term growth and inflation prospects. Portfolio managers are increasingly allocating assets based on these yield curve signals, with many rotating into sectors that historically perform well during periods of higher inflation."
        },
        {
            "type": "split-section",
            "layout": "image-right",
            "image": {
                "url": "https://images.unsplash.com/photo-1563986768609-322da13575f3",
                "alt": "Market analyst examining financial charts"
            },
            "content": {
                "type": "text",
                "content": "Market analysts point to several key factors driving these trends: the continued digital transformation of financial services, evolving regulatory frameworks, and changing consumer preferences. The rise of decentralized finance (DeFi) and digital assets has introduced new variables into traditional market analysis. Central banks worldwide are accelerating research into Central Bank Digital Currencies (CBDCs), with pilot programs already underway in China, Sweden, and the Bahamas, potentially reshaping the future of monetary policy implementation."
            }
        },
        {
            "type": "heading",
            "content": "Technological Innovation in Financial Services"
        },
        {
            "type": "text",
            "content": "The financial technology sector continues to evolve rapidly, with artificial intelligence and blockchain technology leading the transformation. Traditional banks are investing heavily in digital infrastructure, while fintech startups are challenging established business models with innovative solutions. This technological revolution is not just changing how we conduct transactions, but also how we think about value and financial relationships. The integration of machine learning algorithms into credit decisioning has improved access to financial services for previously underserved populations, while simultaneously reducing default rates for lenders through more nuanced risk assessment models."
        },
        {
            "type": "image",
            "url": "https://images.unsplash.com/photo-1640340434855-6084b1f4901c",
            "alt": "Advanced trading terminal with multiple screens"
        },
        {
            "type": "text",
            "content": "Blockchain technology continues to mature beyond its cryptocurrency origins, finding applications in trade finance, securities settlement, and identity verification. Enterprise blockchain solutions are being deployed by major financial institutions to reduce friction in cross-border payments and streamline back-office operations. Smart contracts are automating complex financial agreements, reducing overhead costs and minimizing the potential for disputes. These innovations are particularly transformative for emerging markets, where traditional financial infrastructure limitations can be leapfrogged through distributed ledger technologies."
        },
        {
            "type": "grid-section",
            "columns": 2,
            "items": [
                {
                    "type": "stat-card",
                    "title": "Digital Payments Growth",
                    "value": "127%",
                    "trend": "up",
                    "description": "Year-over-year increase in digital transaction volume"
                },
                {
                    "type": "stat-card",
                    "title": "AI Adoption in Finance",
                    "value": "85%",
                    "trend": "up",
                    "description": "Financial institutions implementing AI solutions"
                },
                {
                    "type": "stat-card",
                    "title": "Sustainable Investment",
                    "value": "$31.7T",
                    "trend": "up",
                    "description": "Global ESG assets under management"
                },
                {
                    "type": "stat-card",
                    "title": "Fintech Market Share",
                    "value": "23%",
                    "trend": "up",
                    "description": "Of global financial services revenue"
                }
            ]
        },
        {
            "type": "heading",
            "content": "Investment Strategies for the Digital Age"
        },
        {
            "type": "text",
            "content": "Modern investment strategies are evolving to incorporate both traditional wisdom and new technological capabilities. Quantitative analysis, powered by machine learning algorithms, is becoming increasingly sophisticated, enabling investors to identify patterns and opportunities that were previously invisible. The democratization of investment tools through mobile apps and robo-advisors has made sophisticated investment strategies accessible to retail investors. Factor-based investing has gained significant traction, with multi-factor models demonstrating resilience across various market conditions by balancing exposure to value, momentum, quality, and low volatility characteristics."
        },
        {
            "type": "image",
            "url": "https://images.unsplash.com/photo-1639322537228-f710d846310a",
            "alt": "Investment strategy planning session"
        },
        {
            "type": "text",
            "content": "Alternative data sources are becoming increasingly central to investment decision-making. Satellite imagery is being used to monitor agricultural yields and retail parking lot occupancy. Natural language processing algorithms analyze earnings call transcripts to detect sentiment shifts among corporate executives. Mobile payment data provides real-time insights into consumer spending patterns. These non-traditional data streams, combined with computational power to analyze them, are creating new forms of information asymmetry that sophisticated investors can leverage for alpha generation."
        },
        {
            "type": "heading",
            "content": "The Future of Sustainable Finance"
        },
        {
            "type": "text",
            "content": "Environmental, Social, and Governance (ESG) considerations are becoming central to investment decision-making. Asset managers are developing sophisticated frameworks to evaluate companies' sustainability practices, while regulators are implementing new disclosure requirements. This shift towards sustainable finance is driving innovation in financial products and creating new opportunities for both investors and entrepreneurs. The EU Taxonomy for sustainable activities has established clear criteria for what constitutes environmentally sustainable economic activities, bringing standardization to a previously fragmented market and facilitating the flow of capital towards genuinely sustainable projects."
        },
        {
            "type": "image",
            "url": "https://images.unsplash.com/photo-1623659873902-3b11d35d5c33",
            "alt": "Renewable energy investment project"
        },
        {
            "type": "text",
            "content": "Green bonds continue to see record issuance, with proceeds dedicated to environmental projects ranging from renewable energy infrastructure to biodiversity conservation. Sustainability-linked loans and bonds, which tie financing costs to the achievement of specific sustainability targets, are incentivizing corporations to accelerate their transition to more sustainable business models. Impact investing is moving from the periphery to the mainstream, with major asset managers launching dedicated funds focused on generating measurable social and environmental benefits alongside financial returns."
        },
        {
            "type": "heading",
            "content": "Regulatory Developments and Compliance Challenges"
        },
        {
            "type": "text",
            "content": "The regulatory landscape for financial services continues to evolve at a rapid pace. Global harmonization efforts, such as the Basel IV framework for banking regulation, aim to establish consistent standards across jurisdictions. However, regional variations in implementation create compliance complexity for multinational financial institutions. The EU's Digital Operational Resilience Act (DORA) and the UK's Operational Resilience Framework are elevating expectations regarding financial institutions' ability to withstand, absorb and recover from operational disruptions, with an emphasis on critical business services rather than just systems."
        },
        {
            "type": "image",
            "url": "https://images.unsplash.com/photo-1589829085413-56de8ae18c73",
            "alt": "Financial regulatory compliance meeting"
        },
        {
            "type": "text",
            "content": "Cryptocurrency regulation represents a significant frontier, with jurisdictions taking divergent approaches. Some are creating comprehensive regulatory frameworks to encourage innovation while protecting consumers, while others are adopting more restrictive stances. Anti-money laundering and counter-terrorism financing requirements continue to expand, with beneficial ownership registries becoming more common globally. Financial institutions are investing heavily in RegTech solutions to manage these compliance challenges efficiently, with artificial intelligence playing an increasingly important role in transaction monitoring and fraud detection."
        },
        {
            "type": "expert-panel",
            "title": "Expert Insights",
            "experts": [
                {
                    "name": "Dr. Sarah Chen",
                    "role": "Chief Investment Strategist",
                    "quote": "The convergence of AI and financial markets is creating unprecedented opportunities for sophisticated risk management and alpha generation.",
                    "image": "https://images.unsplash.com/photo-1573497019940-1c28c88b4f3e"
                },
                {
                    "name": "Michael Rodriguez",
                    "role": "Global Markets Director",
                    "quote": "Sustainable finance is no longer just a trend - it's becoming a fundamental factor in investment decision-making.",
                    "image": "https://images.unsplash.com/photo-1472099645785-5658abf4ff4e"
                }
            ]
        }
    ],
    "image": "https://images.unsplash.com/photo-1611974789855-9c2a0a7236a3"
}

# Travel blog template with significantly expanded content and more images
travel_blog_template = {
    "id": "travel",
    "name": "Travel Blog",
    "title": "Hidden Gems of Southeast Asia: A Journey Beyond the Tourist Trail",
    "author": "Travel Explorer Team",
    "date": "March 15, 2024",
    "content": [
        {
            "type": "text",
            "content": "Southeast Asia has long captivated travelers with its rich tapestry of cultures, stunning landscapes, and warm hospitality. While destinations like Bangkok, Bali, and Singapore draw millions of visitors annually, the region harbors countless hidden treasures waiting to be discovered by those willing to venture off the beaten path. These lesser-known destinations offer authentic experiences, unburdened by mass tourism, where travelers can forge genuine connections with local communities and immerse themselves in the region's incredible diversity."
        },
        {
            "type": "image",
            "url": "https://images.unsplash.com/photo-1552465011-b4e21bf6e79a",
            "alt": "Pristine beach with traditional long-tail boats"
        },
        {
            "type": "heading",
            "content": "Undiscovered Islands of Thailand"
        },
        {
            "type": "text",
            "content": "Beyond the well-known islands of Phuket and Koh Samui lies an archipelago of lesser-visited gems. Koh Yao Noi, situated in Phang Nga Bay, offers a glimpse of traditional Thai island life, with rubber plantations, fishing villages, and empty beaches. The island's western coast provides spectacular sunset views over the limestone karsts that dot the bay, while local communities welcome visitors with genuine warmth and hospitality. The pace of life here remains blissfully unhurried, with bicycles and motorbikes serving as the primary modes of transportation on the island's quiet roads."
        },
        {
            "type": "image",
            "url": "https://images.unsplash.com/photo-1552733407-5d5c46c3bb3b",
            "alt": "Aerial view of pristine island beaches"
        },
        {
            "type": "text",
            "content": "Further south, the Trang Islands remain one of Thailand's best-kept secrets. Koh Mook's Emerald Cave (Tham Morakot) requires swimming through a dark passageway that suddenly opens to reveal a hidden beach surrounded by towering limestone cliffs and lush vegetation – a true pirate's hideaway. Nearby Koh Kradan offers some of Thailand's most pristine beaches and spectacular snorkeling directly from shore, with vibrant coral reefs just a short swim from the powdery sand. Koh Ngai's eastern shore features shallow, crystal-clear waters perfect for families, while its western side offers sunset vistas that rival any in the world."
        },
        {
            "type": "image",
            "url": "https://images.unsplash.com/photo-1544644181-1484b3fdfc32",
            "alt": "Emerald Cave hidden beach surrounded by limestone cliffs"
        },
        {
            "type": "heading",
            "content": "Vietnam's Hidden Highlands"
        },
        {
            "type": "text",
            "content": "In the northern reaches of Vietnam, beyond the tourist hotspots of Hanoi and Halong Bay, the Ha Giang province offers an untouched landscape of dramatic mountain passes, terraced rice fields, and authentic cultural experiences. Home to various ethnic minority groups, this region provides opportunities for homestays, allowing travelers to immerse themselves in local traditions and daily life. The Dong Van Karst Plateau Geopark, recognized by UNESCO for its geological significance, features a otherworldly landscape of limestone formations that have been shaped over millions of years."
        },
        {
            "type": "split-section",
            "layout": "image-left",
            "image": {
                "url": "https://images.unsplash.com/photo-1528127269322-539801943592",
                "alt": "Traditional Vietnamese village in the mountains"
            },
            "content": {
                "type": "text",
                "content": "The winding roads of the Ha Giang Loop offer some of Asia's most spectacular motorcycle routes, taking riders through remote villages and breathtaking mountain vistas. Local markets burst with color and activity, where different ethnic groups gather to trade goods and maintain centuries-old traditions. The Sunday market in Bac Ha is particularly renowned, with Flower Hmong women in vibrant traditional dress selling textiles, medicinal herbs, and local produce. This high-altitude region also produces exceptional tea, with ancient tea forests containing trees over 400 years old that produce complex, nuanced flavors found nowhere else."
            }
        },
        {
            "type": "image",
            "url": "https://images.unsplash.com/photo-1518802170774-e31e2bb6b6eb",
            "alt": "Terraced rice fields in northern Vietnam highlands"
        },
        {
            "type": "text",
            "content": "Venturing into central Vietnam, the ancient Cham towers of Qui Nhon stand as testament to the once-powerful Champa civilization that ruled these lands for centuries. These brick structures, adorned with intricate carvings and sculptures, have weathered centuries of monsoons and wars yet remain remarkably intact. The nearby coastal city of Qui Nhon itself offers pristine beaches without the crowds of more famous Vietnamese coastal destinations, along with a thriving seafood culture where fishermen bring their daily catch directly to beachside restaurants."
        },
        {
            "type": "image",
            "url": "https://images.unsplash.com/photo-1557750255-c76072a7aad1",
            "alt": "Ancient Cham towers against blue sky"
        },
        {
            "type": "heading",
            "content": "Cambodia's Forgotten Temples"
        },
        {
            "type": "text",
            "content": "While Angkor Wat draws millions of visitors, Cambodia's Preah Vihear province houses equally impressive but far less visited temple complexes. Perched atop a cliff in the Dangrek Mountains, the Preah Vihear temple offers stunning views across the Cambodian plains and a peaceful atmosphere that's increasingly rare at more popular sites. Built between the 9th and 12th centuries, this Hindu temple dedicated to Shiva features some of the most impressive carved lintels and pediments of any Khmer structure, telling ancient stories of gods and kings through intricate stone carvings."
        },
        {
            "type": "image",
            "url": "https://images.unsplash.com/photo-1508009603885-50cf7c579365",
            "alt": "Ancient temple ruins at sunrise"
        },
        {
            "type": "text",
            "content": "The lesser-known temples of Koh Ker, once a brief capital of the Khmer Empire, lie hidden in the jungle north of Siem Reap. The main structure, Prasat Thom, resembles a seven-tiered pyramid rising dramatically from the forest floor. Unlike the more manicured grounds of Angkor, many of Koh Ker's structures remain partially reclaimed by the jungle, with massive tree roots embracing ancient stones in a haunting display of nature's power. The remote location means visitors can often explore these magnificent structures in solitude, imagining themselves as early explorers discovering these ruins for the first time."
        },
        {
            "type": "image",
            "url": "https://images.unsplash.com/photo-1555855815-43e2a068efdd",
            "alt": "Temple ruins being reclaimed by jungle"
        },
        {
            "type": "heading",
            "content": "Local Cuisine and Cultural Experiences"
        },
        {
            "type": "text",
            "content": "One of the most rewarding aspects of exploring Southeast Asia's hidden corners is discovering authentic local cuisine. From street food vendors in remote market towns to family-run restaurants serving recipes passed down through generations, these culinary experiences offer insight into the region's rich cultural heritage. In northern Thailand's Nan Province, traditional Tai Lue cuisine features distinctive dishes like khao kan jin (rice mixed with blood and herbs, steamed in banana leaves) and kaeng no mai (spicy bamboo shoot curry) that reflect the area's unique cultural influences from both Lanna Thai and Laotian traditions."
        },
        {
            "type": "image",
            "url": "https://images.unsplash.com/photo-1504674900247-0877df9cc836",
            "alt": "Traditional Southeast Asian street food market"
        },
        {
            "type": "text",
            "content": "Throughout the region, cooking classes have become a popular way for travelers to connect with local culture. Unlike the standardized tourist-oriented classes in major cities, those found in rural areas often take place in family homes, beginning with trips to local markets or even foraging expeditions to gather wild herbs and vegetables. In Battambang, Cambodia, visitors can learn to prepare amok (fish curry steamed in banana leaves) and prahok ktis (fermented fish dip) using traditional clay pots over wood fires, techniques that have remained unchanged for centuries despite the conveniences of modern cooking equipment."
        },
        {
            "type": "image",
            "url": "https://images.unsplash.com/photo-1455619452474-d2be8b1e70cd",
            "alt": "Traditional cooking class with local ingredients"
        },
        {
            "type": "heading",
            "content": "Sacred Mountains of Myanmar"
        },
        {
            "type": "text",
            "content": "In the remote regions of Myanmar, ancient pilgrimage sites and mountain monasteries offer spiritual encounters far from the tourist crowds. The Golden Rock at Mount Kyaiktiyo defies physical laws, perched impossibly on the edge of a cliff and drawing devout Buddhists from across the country. The journey to this sacred site is as meaningful as the destination itself, with pilgrims walking barefoot up the mountain path while contemplating their faith. Covered in gold leaf applied by generations of worshippers, the rock glows magnificently at sunrise and sunset, seeming to change colors throughout the day as the light shifts across its surface."
        },
        {
            "type": "image",
            "url": "https://images.unsplash.com/photo-1546542868-c0e57ca3dd4f",
            "alt": "Golden Rock at Kyaiktiyo"
        },
        {
            "type": "text",
            "content": "Mount Popa, an extinct volcano rising dramatically from central Myanmar's plains, hosts a monastery perched atop a sheer volcanic plug. Home to the 37 Great Nats (spirits) of Burmese tradition, this site represents the fascinating synthesis of ancient animist beliefs with Buddhist practices. Visitors climb 777 steps to reach the summit, passing shrines and being kept company by the resident monkeys that have made this sacred mountain their home. From the top, panoramic views extend across the plains to distant mountains, with the golden stupas of the monastery gleaming in the sun against the backdrop of endless sky."
        },
        {
            "type": "image",
            "url": "https://images.unsplash.com/photo-1598127004624-62c8c1cc68b4",
            "alt": "Sunrise at Mount Popa monastery"
        },
        {
            "type": "text",
            "content": "In the Shan State highlands, the remote monastery of Kakku contains over 2,000 stupas crowded together in a relatively small area, creating a mesmerizing forest of spires that dates back to the 3rd century BCE. Many of these stupas feature intricate carvings of celestial beings and mythological creatures that have survived centuries of weathering. Unlike Myanmar's more famous archaeological sites, Kakku remains largely unknown to international travelers, allowing for contemplative exploration without crowds or commercial development."
        },
        {
            "type": "image",
            "url": "https://images.unsplash.com/photo-1519882189396-71b93cb3e428",
            "alt": "Buddhist monastery in misty mountains"
        },
        {
            "type": "heading",
            "content": "Island Life in the Philippines"
        },
        {
            "type": "text",
            "content": "With over 7,000 islands, the Philippines offers endless opportunities for exploration beyond the well-trodden paths of Boracay and Palawan. The Batanes Islands, located in the northernmost part of the Philippines where the Pacific Ocean meets the South China Sea, feature a dramatic landscape of rolling hills, stone villages, and rugged coastlines that more closely resemble the Scottish Highlands than tropical Southeast Asia. The indigenous Ivatan people have developed unique stone architecture designed to withstand the powerful typhoons that frequently pass through the region."
        },
        {
            "type": "image",
            "url": "https://images.unsplash.com/photo-1518509562904-e7ef99cdcc86",
            "alt": "Rolling green hills meeting the ocean in Batanes"
        },
        {
            "type": "text",
            "content": "The Bacuit Archipelago near El Nido contains dozens of limestone islands with hidden lagoons that can only be accessed at specific tide levels through small openings in the rock. Kayaking through these secret passages reveals pristine enclosed beaches and lagoons of impossibly blue water surrounded entirely by towering limestone cliffs draped in vegetation. Marine life thrives in these protected waters, with vibrant coral gardens and reef fish visible just below the surface."
        },
        {
            "type": "image",
            "url": "https://images.unsplash.com/photo-1559444839-b936c2cbf637",
            "alt": "Hidden lagoon with turquoise water surrounded by limestone cliffs"
        },
        {
            "type": "heading",
            "content": "Sustainable Travel and Local Communities"
        },
        {
            "type": "text",
            "content": "As these hidden destinations gain recognition, sustainable tourism practices become increasingly important. Many communities are developing eco-friendly accommodations and responsible tourism initiatives that allow visitors to experience these special places while preserving their natural and cultural heritage for future generations. In Thailand's Koh Prasat Yai, a community-based tourism initiative has helped the island develop small-scale, locally-owned accommodations built using traditional methods and materials, with profits being reinvested into conservation projects and educational opportunities for local children."
        },
        {
            "type": "image",
            "url": "https://images.unsplash.com/photo-1625472603517-1b0dc72846ab",
            "alt": "Eco-friendly bamboo accommodation in the jungle"
        },
        {
            "type": "text",
            "content": "The Cambodia Community-Based Ecotourism Network (CCBEN) connects travelers with village homestays in remote regions, where visitors can participate in conservation efforts such as wildlife monitoring or reforestation projects. These initiatives provide alternative livelihoods for communities that might otherwise depend on environmentally destructive practices like logging or wildlife poaching. By creating economic value for intact ecosystems and cultural traditions, sustainable tourism can help preserve Southeast Asia's hidden gems for generations to come."
        },
        {
            "type": "image",
            "url": "https://images.unsplash.com/photo-1513836279014-a89f7a76ae86",
            "alt": "Community-led reforestation project"
        },
        {
            "type": "expert-panel",
            "title": "Local Insights",
            "experts": [
                {
                    "name": "Mai Tran",
                    "role": "Local Guide, Ha Giang",
                    "quote": "The real magic of our region lies in the small moments - sharing tea with a local family, learning traditional crafts, or discovering a viewpoint that takes your breath away.",
                    "image": "https://images.unsplash.com/photo-1494790108377-be9c29b29330"
                },
                {
                    "name": "Somchai Prak",
                    "role": "Cultural Heritage Expert",
                    "quote": "These lesser-known destinations preserve authentic traditions that are increasingly rare in more touristy areas. They offer a window into Southeast Asia's rich cultural heritage.",
                    "image": "https://images.unsplash.com/photo-1507003211169-0a1dd7228f2d"
                }
            ]
        }
    ],
    "image": "https://images.unsplash.com/photo-1552465011-b4e21bf6e79a"
}

# Save the default template to file
if os.path.exists(os.path.join(TEMPLATES_FOLDER, "tech_template.json")):
    os.remove(os.path.join(TEMPLATES_FOLDER, "tech_template.json"))
with open(os.path.join(TEMPLATES_FOLDER, "tech_template.json"), "w") as f:
    json.dump(tech_blog_template, f)

# Save the finance template to file
if os.path.exists(os.path.join(TEMPLATES_FOLDER, "finance_template.json")):
    os.remove(os.path.join(TEMPLATES_FOLDER, "finance_template.json"))
with open(os.path.join(TEMPLATES_FOLDER, "finance_template.json"), "w") as f:
    json.dump(finance_blog_template, f)

# Save the travel template to file
if os.path.exists(os.path.join(TEMPLATES_FOLDER, "travel_template.json")):
    os.remove(os.path.join(TEMPLATES_FOLDER, "travel_template.json"))
with open(os.path.join(TEMPLATES_FOLDER, "travel_template.json"), "w") as f:
    json.dump(travel_blog_template, f)

# Load template from file without adding randomization
def load_template(template_id):
    try:
        # Extract the base template id without any timestamps or random strings
        base_template_id = template_id.split('_')[0] if '_' in template_id else template_id
        
        template_path = os.path.join(TEMPLATES_FOLDER, f"{base_template_id}_template.json")
        if os.path.exists(template_path):
            with open(template_path, "r") as f:
                template = json.load(f)
                return template
        else:
            # Return default template if file doesn't exist
            if base_template_id == "finance":
                return finance_blog_template
            elif base_template_id == "travel":
                return travel_blog_template
            else:
                return tech_blog_template
    except Exception as e:
        print(f"Error loading template: {str(e)}")
        return tech_blog_template

# Add new function to clean up template files
def cleanup_template_files():
    """Clean up old template files, keeping only the original template files"""
    try:
        # Get all files in template directory
        template_files = os.listdir(TEMPLATES_FOLDER)
        
        # Keep only the base template files
        base_templates = ["tech_template.json", "finance_template.json", "travel_template.json"]
        
        # Delete all other template files
        for file in template_files:
            if file not in base_templates:
                file_path = os.path.join(TEMPLATES_FOLDER, file)
                try:
                    os.remove(file_path)
                    print(f"Deleted template file: {file}")
                except Exception as e:
                    print(f"Error deleting file {file}: {str(e)}")
        
        # Also check for any temporary files in the uploads directory
        uploads_dir = "uploads"
        if os.path.exists(uploads_dir):
            for filename in os.listdir(uploads_dir):
                # Keep only files from the last hour to avoid disrupting current sessions
                file_path = os.path.join(uploads_dir, filename)
                try:
                    file_modification_time = os.path.getmtime(file_path)
                    if time.time() - file_modification_time > 3600:  # 1 hour
                        os.remove(file_path)
                        print(f"Removed old upload file: {filename}")
                except Exception as e:
                    print(f"Error checking/removing file {filename}: {str(e)}")
        
        return True
    except Exception as e:
        print(f"Error cleaning up template files: {str(e)}")
        return False

# Modify save_template function to use the base template name only
def save_template(template_data):
    try:
        template_id = template_data.get("id", "tech")
        
        # Extract the base template id without any timestamps or random strings
        base_template_id = template_id.split('_')[0] if '_' in template_id else template_id
        
        # Add metadata for tracking
        if "metadata" not in template_data:
            template_data["metadata"] = {}
            
        template_data["metadata"]["last_saved"] = int(time.time())
        
        # Save only to the base template file
        template_path = os.path.join(TEMPLATES_FOLDER, f"{base_template_id}_template.json")
        
        with open(template_path, "w") as f:
            json.dump(template_data, f)
            
        # Run cleanup to remove any extra files
        cleanup_template_files()
            
        return True
    except Exception as e:
        print(f"Error saving template: {str(e)}")
        return False

# Enhanced chatbot function to handle user queries with randomization
def query_chatbot(user_message, template_id="tech"):
    try:
        print(f"[DEBUG] query_chatbot called with user_message={user_message}, template_id={template_id}")
        if not GROQ_API_KEY:
            return "Error: Groq API key is missing.", None
        current_template = load_template(template_id)
        client = groq.Groq(api_key=GROQ_API_KEY)
        random_seed = int(time.time())
        random_factor = random.randint(1, 1000)
        is_file_upload = user_message.startswith("Generate a complete blog post")
        if is_file_upload:
            # Enhanced system prompt for file uploads with more detailed instructions
            system_prompt = """You are BlogGen Assistant, a specialized AI that generates high-quality blog content based on user uploads.
            Current template type: {template_id}
            Generation seed: {random_seed}
            Variation factor: {random_factor}
            
            YOUR TASK IS TO GENERATE COMPLETELY UNIQUE CONTENT EACH TIME, EVEN IF THE INPUTS ARE SIMILAR.
            Use the generation seed and variation factor to create highly varied content.
            
            Your task is to create a fully-formed blog post using the uploaded file content while maintaining the structure of the selected template.
            
            IMPORTANT INSTRUCTIONS:
            1. Generate a compelling title related to the uploaded content - make it unique and different each time
            2. Create appropriate headings and subheadings that organize the content logically - vary structure based on random seed
            3. Expand the content to be thorough and informative, at least 1200 words total
            4. Suggest appropriate alt text for images based on the section content
            5. Maintain the overall structure and flow of the template type
            6. ENSURE YOU CREATE DIFFERENT CONTENT EACH TIME even for similar inputs
            7. Return your complete response as a JSON object matching the template structure
            
            The JSON must follow this structure:
            ```json
            {{
              "title": "Generated Title",
              "content": [
                {{"type": "text", "content": "Introduction paragraph..."}},
                {{"type": "image", "url": "IMAGE_URL", "alt": "Descriptive alt text"}},
                {{"type": "heading", "content": "First Section Heading"}},
                {{"type": "text", "content": "First section content..."}},
                ...
              ]
            }}
            ```
            """
            # For uploads, use larger token limit to handle complete blog generation
            max_tokens = 2000
        else:
            # Enhanced system prompt for blog modifications
            system_prompt = """You are BlogGen Assistant, a helpful AI that assists users with blog creation and content modifications.
        Current template: {template_id}
        Generation seed: {random_seed}
        Variation factor: {random_factor}
            
            IMPORTANT: Ensure your responses and generated content are unique each time, even for similar requests.
            
            When users request changes to their blog:
            1. Understand their specific requirements
            2. Modify only the requested parts while maintaining the overall structure
            3. Return changes in a JSON block with the specific updates
            4. Keep the existing content and structure intact for parts not being modified
            5. Use the random seeds to vary your approach to content creation
            
            For example, if user asks to change the title, return:
            ```json
            {{
              "title": "New Title"
            }}
            ```
            
            If user asks to modify a specific section, return:
            ```json
            {{
              "content": [
                {{"type": "text", "content": "Updated content..."}},
                {{"type": "heading", "content": "Updated heading"}}
              ]
            }}
            ```
            
            Keep responses clear and helpful. If user asks for specific blog help, offer creative suggestions.
            """
            max_tokens = 500

        # Try to get a response from Groq API
        try:
            response = client.chat.completions.create(
                model="meta-llama/llama-4-scout-17b-16e-instruct",  # Use a widely available model name
                messages=[
                    {"role": "system", "content": system_prompt.format(
                        template_id=template_id,
                        random_seed=random_seed,
                        random_factor=random_factor
                    )},
                    {"role": "user", "content": user_message}
                ],
                temperature=0.9,  # Increase temperature for more variety
                max_tokens=max_tokens
            )
            response_text = response.choices[0].message.content
        except Exception as groq_error:
            import traceback
            print("[ERROR] Groq API call failed:")
            traceback.print_exc()
            print(f"[ERROR] Exception details: {groq_error}")
            # Fallback to Gemini if Groq fails
            if GEMINI_API_KEY:
                try:
                    genai.configure(api_key=GEMINI_API_KEY)
                    model = genai.GenerativeModel(model_name="gemma-3n-e4b-it")
                    prompt = f"""System: {system_prompt.format(
                        template_id=template_id,
                        random_seed=random_seed,
                        random_factor=random_factor
                    )}\nUser: {user_message}"""
                    response = model.generate_content(prompt)
                    response_text = response.text
                except Exception as gemini_error:
                    print("[ERROR] Gemini API call failed:")
                    traceback.print_exc()
                    print(f"[ERROR] Exception details: {gemini_error}")
                    return f"Sorry, I'm having trouble connecting to my AI services. (Gemini error: {str(gemini_error)})", None
            else:
                return f"Sorry, I'm having trouble responding right now. (Groq error: {str(groq_error)})", None
        
        # Rest of the function remains the same
        # Process the response
        updated_template = current_template.copy()  # Create a copy to avoid reference issues
        
        # Add randomization metadata to template
        if "metadata" not in updated_template:
            updated_template["metadata"] = {}
        updated_template["metadata"]["generated_at"] = int(time.time())
        updated_template["metadata"]["random_seed"] = random_seed
        updated_template["metadata"]["random_factor"] = random_factor
        
        # Extract JSON updates if present
        json_blocks = re.findall(r'```json(.*?)```', response_text, re.DOTALL)
        if json_blocks:
            for block in json_blocks:
                try:
                    updates = json.loads(block.strip())
                    
                    # Handle title updates
                    if "title" in updates:
                        updated_template["title"] = updates["title"]
                    
                    # Handle author updates
                    if "author" in updates:
                        updated_template["author"] = updates["author"]
                    
                    # Handle date updates
                    if "date" in updates:
                        updated_template["date"] = updates["date"]
                    
                    # Handle content updates
                    if "content" in updates:
                        # Check if this is a full content replacement or partial update
                        if any(item.get("type") == "text" for item in updates["content"]):
                            # This is likely a full content replacement
                            updated_template["content"] = updates["content"]
                        else:
                            # This is likely a partial update - merge with existing content
                            existing_content = updated_template.get("content", [])
                            for new_item in updates["content"]:
                                # Find matching item in existing content and update it
                                for i, existing_item in enumerate(existing_content):
                                    if existing_item.get("id") == new_item.get("id"):
                                        existing_content[i] = new_item
                                        break
                                else:
                                    # If no match found, append the new item
                                    existing_content.append(new_item)
                    
                    # Handle image updates
                    if "image" in updates:
                        updated_template["image"] = updates["image"]
                    
                    # Save the updated template
                    save_template(updated_template)
                except json.JSONDecodeError as e:
                    print(f"[ERROR] JSON decode error: {str(e)}\nRaw AI response block:\n{block.strip()}")
                    print(f"[ERROR] Full AI response:\n{response_text}")
                    return (f"Sorry, I encountered an error processing the AI's response. Please try again or rephrase your request. (Details: {str(e)})", None)
                except Exception as e:
                    print(f"[ERROR] Template update error: {str(e)}\nRaw AI response block:\n{block.strip()}")
                    print(f"[ERROR] Full AI response:\n{response_text}")
                    return (f"Sorry, I encountered an error processing the AI's response. Please try again or rephrase your request. (Details: {str(e)})", None)
            # Clean response and remove JSON blocks
            clean_response = re.sub(r'```json.*?```', '', response_text, flags=re.DOTALL).strip()
            # --- Title extraction logic ---
            if any(word in user_message.lower() for word in ["title", "headline"]):
                title = extract_concrete_title(clean_response)
                if title:
                    return title, updated_template
            return clean_response, updated_template
        else:
            # No JSON block found, just return the AI's plain text response
            clean_response = re.sub(r'```.*?```', '', response_text, flags=re.DOTALL).strip()
            # --- Title extraction logic ---
            if any(word in user_message.lower() for word in ["title", "headline"]):
                title = extract_concrete_title(clean_response)
                if title:
                    return title, updated_template
            return clean_response, updated_template
    except Exception as e:
        print(f"Error in query_chatbot: {str(e)}")
        return f"Error: {str(e)}", None

# Helper function to extract a concrete title from AI response
import re

def extract_concrete_title(response_text):
    """
    Extract a likely title from the AI response, even if the response is generic or includes a preamble.
    Returns a string title or None if not found.
    """
    if not response_text:
        return None
    # Look for common title patterns
    # 1. Markdown or plain: Title: ...
    match = re.search(r"(?:^|\n)\s*Title\s*[:\-–]\s*(.+)", response_text, re.IGNORECASE)
    if match:
        title = match.group(1).strip().strip('"')
        # Stop at first line break or period if present
        title = re.split(r'[\n\r\.]', title)[0].strip()
        return title
    # 2. Quoted or bolded line at the top
    match = re.search(r'^[\"\']?([A-Z][^\n\r]{10,80})[\"\']?$', response_text.strip(), re.MULTILINE)
    if match:
        return match.group(1).strip()
    # 3. First non-empty line if it looks like a title
    lines = response_text.strip().splitlines()
    for line in lines:
        line = line.strip().strip('"')
        if 10 <= len(line) <= 80 and line[0].isupper() and not line.endswith((':', '-', '—')):
            return line
    # 4. Fallback: first sentence, truncated
    sentence = re.split(r'[\.!?]', response_text)[0]
    if 10 < len(sentence) < 80:
        return sentence.strip()
    return None

# Generate HTML for preview from template data
def generate_html_preview(template_data):
    """Generate HTML preview for a blog template"""
    try:
        html = f"""
        <!DOCTYPE html>
        <html lang="en">
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
            <title>{template_data.get('title', 'Blog Preview')}</title>
            <style>
                body {{
                    font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
                    line-height: 1.6;
                    color: #333;
                    max-width: 800px;
                    margin: 0 auto;
                    padding: 20px;
                }}
                img {{
                    max-width: 100%;
                    height: auto;
                    border-radius: 8px;
                    margin: 20px 0;
                }}
                h1, h2, h3 {{
                    color: #2c3e50;
                }}
                h1 {{
                    text-align: center;
                }}
                .meta {{
                    color: #666;
                    margin-bottom: 30px;
                }}
                .content-section {{
                    margin-bottom: 30px;
                }}
                blockquote {{
                    border-left: 4px solid #ddd;
                    margin: 0;
                    padding-left: 20px;
                    font-style: italic;
                }}
            </style>
        </head>
        <body>
            <h1>{template_data.get('title', 'Blog Title')}</h1>
            <div class="meta">
                <p>By {template_data.get('author', 'Anonymous')} | {template_data.get('date', 'Date')}</p>
            </div>
        """
        
        # Add content sections
        for section in template_data.get('content', []):
            section_type = section.get('type', '')
            if section_type == 'text':
                html += f"""
                <div class="content-section">
                    <p>{section.get('content', '')}</p>
                </div>
                """
            elif section_type == 'heading':
                html += f"""
                <div class="content-section">
                    <h2>{section.get('content', '')}</h2>
                </div>
                """
            elif section_type == 'image':
                html += f"""
                <div class="content-section">
                    <img src="{section.get('url', '')}" alt="{section.get('alt', '')}" loading="lazy">
                    {f'<p class="caption">{section["caption"]}</p>' if 'caption' in section else ''}
                </div>
                """
            elif section_type == 'quote':
                html += f"""
                <div class="content-section">
                    <blockquote>{section.get('content', '')}</blockquote>
                </div>
                """
            elif section_type == 'split-section':
                layout = section.get('layout', 'image-left')
                image_html = f"""
                <div class="split-image">
                    <img src="{section['image']['url']}" alt="{section['image'].get('alt', '')}" loading="lazy">
                </div>
                """ if 'image' in section else ''
                
                html += f"""
                <div class="split-section {layout}">
                    {image_html}
                    <div class="split-content">
                        {section.get('content', '')}
                    </div>
                </div>
                """
        
        html += '</div></div>'
        
        return html
    
    except Exception as e:
        print(f"Error generating HTML preview: {str(e)}")
        return f"<p>Error generating preview: {str(e)}</p>"

@app.route("/")
def home():
    # Clean up any unnecessary template files on application startup
    cleanup_template_files()
    return render_template("index.html")

@app.route("/templates")
def templates():
    return render_template("templates.html")

@app.route("/News")
def news():
    return render_template("news.html")

@app.route("/get_blogs", methods=["GET"])
def get_blogs():
    try:
        # Get all template files
        template_files = [f for f in os.listdir(TEMPLATES_FOLDER) if f.endswith('_template.json')]
        blogs = []
        
        for file in template_files:
            with open(os.path.join(TEMPLATES_FOLDER, file), 'r') as f:
                template = json.load(f)
                blogs.append(template)
        
        return jsonify({"success": True, "blogs": blogs})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)})

@app.route('/chat', methods=['POST'])
def chat():
    try:
        data = request.get_json()
        if not data or 'message' not in data:
            return jsonify({"error": "No message provided"}), 400
        user_message = data['message']
        template_id = data.get('template_id', 'tech')  # Default to tech template

        # Get chatbot response
        response_text, updated_template = query_chatbot(user_message, template_id)
        # Debug output for diagnosis
        print(f"[DEBUG] /query_chatbot: user_message={user_message}, template_id={template_id}, response_text={response_text}, updated_template={'present' if updated_template is not None else 'None'}")
        # Always return a valid response, even if the chatbot fails
        if not response_text or (isinstance(response_text, str) and response_text.lower().startswith('error')):
            import traceback
            print(f"[ERROR] query_chatbot failed. user_message: {user_message}, template_id: {template_id}, response_text: {response_text}")
            traceback.print_stack()
            response_text = "Sorry, I encountered an error processing your request. Please try again."
        response = {"response": response_text}
        if updated_template is not None:
            response["template"] = updated_template
            response["templateId"] = template_id
        return jsonify(response)
    except Exception as e:
        print(f"Error in /chat route: {str(e)}")
        return jsonify({"response": "Sorry, I encountered an error processing your request. Please try again."}), 500

@app.route('/get_template/<template_id>', methods=['GET'])
def get_template(template_id):
    template = load_template(template_id)
    return jsonify(template)

@app.route('/template_html/<template_id>', methods=['GET'])
def get_template_html(template_id):
    """Serve the actual HTML template file from the template folder"""
    try:
        # Look for the HTML file in the template folder
        template_folder = os.path.join('static', 'templates', template_id)
        
        # Define possible HTML file names for each template
        possible_html_files = [
            f'{template_id}.html',  # Standard naming (tech.html, travel.html, etc.)
            f'{template_id}_blog.html',  # For blog templates
            f'{template_id}_template.html',  # For template files
            'index.html',  # For tool_installition
            'healtth.html' if template_id == 'health' else None,  # Handle typo in health template
            'comparison_template.html' if template_id == 'comparision' else None,  # Handle comparision template
            'internship_journey_blog.html' if template_id == 'one_line_story' else None,  # Handle one_line_story template
        ]
        
        # Remove None values
        possible_html_files = [f for f in possible_html_files if f is not None]
        
        html_content = None
        found_file = None
        
        # Try each possible filename
        for html_file_name in possible_html_files:
            html_file = os.path.join(template_folder, html_file_name)
            if os.path.exists(html_file):
                with open(html_file, 'r', encoding='utf-8') as f:
                    html_content = f.read()
                found_file = html_file_name
                break
        
        if html_content:
            # Fix image paths to be relative to the template folder
            # Replace relative image paths with absolute paths
            html_content = html_content.replace('src="', f'src="/static/templates/{template_id}/')
            html_content = html_content.replace('src=\'', f'src=\'/static/templates/{template_id}/')
            
            return jsonify({
                'template_html': html_content,
                'template_id': template_id,
                'found_file': found_file,
                'status': 'success'
            })
        else:
            return jsonify({
                'error': f'Template HTML file not found in {template_folder}. Tried: {possible_html_files}',
                'template_id': template_id,
                'status': 'error'
            }), 404
    except Exception as e:
        return jsonify({
            'error': f'Error loading template: {str(e)}',
            'template_id': template_id,
            'status': 'error'
        }), 500

@app.route('/api/placeholder/<int:width>/<int:height>')
def placeholder_image(width, height):
    """Generate a placeholder image with the specified dimensions"""
    try:
        # Create a new image with a white background
        image = Image.new('RGB', (width, height), 'white')
        draw = ImageDraw.Draw(image)
        
        # Draw a border
        draw.rectangle([0, 0, width-1, height-1], outline='#ddd')
        
        # Draw diagonal lines
        draw.line([(0, 0), (width, height)], fill='#ddd')
        draw.line([(0, height), (width, 0)], fill='#ddd')
        
        # Add dimensions text
        text = f'{width}x{height}'
        # Note: In a production environment, you'd want to use a proper font
        # For this example, we'll use a simple text
        text_bbox = draw.textbbox((0, 0), text)
        text_width = text_bbox[2] - text_bbox[0]
        text_height = text_bbox[3] - text_bbox[1]
        text_x = (width - text_width) // 2
        text_y = (height - text_height) // 2
        draw.text((text_x, text_y), text, fill='#999')
        
        # Save to bytes
        img_io = io.BytesIO()
        image.save(img_io, 'PNG')
        img_io.seek(0)
        
        return send_file(img_io, mimetype='image/png')
    except Exception as e:
        return str(e), 500

# Handle file uploads and generate blog content with improved randomization
@app.route('/upload_generate', methods=['POST'])
def upload_generate():
    try:
        print("[DEBUG] upload_generate called", flush=True)
        # Check if a file was uploaded
        file = request.files.get('file')
        if file and file.filename != '':
            print(f"[DEBUG] file object: {file}", flush=True)
            print(f"[DEBUG] Uploaded filename: {file.filename}", flush=True)
            print(f"[DEBUG] File content type: {file.content_type}", flush=True)
            file_uploaded = True
        else:
            print("[DEBUG] No file uploaded, proceeding with default prompt.", flush=True)
            file_uploaded = False        
        # Create strong randomization to ensure content variation
        timestamp = int(time.time())
        random_id = uuid.uuid4().hex
        # Create a unique seed for this specific generation
        unique_seed = f"{timestamp}_{random_id}_{random.randint(1000000, 9999999)}"
        random.seed(unique_seed)
        
        # Force content variation with strong randomization factors
        variation_factors = {
            "timestamp": timestamp,
            "seed": random_id,
            "random_number": random.randint(1000, 9999),
            "content_variation": random.choice(["detailed", "concise", "analytical", "creative", "technical", "narrative", "instructional", "conversational"]),
            "tone_variation": random.choice(["formal", "casual", "enthusiastic", "professional", "conversational", "authoritative", "friendly", "inspirational"]),
            "structure_seed": random.randint(1, 1000),
            "perspective": random.choice(["first_person", "second_person", "third_person"]),
            "focus": random.choice(["problem_solution", "how_to", "list_based", "story_based", "comparison", "timeline"]),
            "unique_id": str(uuid.uuid4())
        }
        
        # Clear any cached results for this file name
        cache_key = f"cached_blog_{file.filename}"
        if cache_key in globals():
            print(f"[DEBUG] Clearing cache for {cache_key}", flush=True)
            del globals()[cache_key]
        
        # Initialize variables for content and images
        file_content = ""
        extracted_images = []
        image_index = 0  # Initialize image_index at the top level
        doc_text = ""
        
        if file:
            print("[DEBUG] File object exists, proceeding to save.", flush=True)
            # Save the file temporarily with a unique name to avoid caching
            upload_folder = "uploads"
            os.makedirs(upload_folder, exist_ok=True)
            unique_filename = f"{timestamp}_{random_id}_{secure_filename(file.filename)}"
            file_path = os.path.join(upload_folder, unique_filename)
            file.save(file_path)
            print(f"[DEBUG] File saved to {file_path}", flush=True)
            # Check if file exists and print its size
            import os
            if os.path.exists(file_path):
                file_size = os.path.getsize(file_path)
                print(f"[DEBUG] File exists after save. Size: {file_size} bytes", flush=True)
            else:
                print(f"[DEBUG] File does NOT exist after save!", flush=True)
            
            # Extract content and images based on file type
            ext = os.path.splitext(file.filename)[1].lower()
            print(f"[DEBUG] File extension: {ext}", flush=True)
            
            # Helper function to add an image to extracted_images
            def add_extracted_image(image_url, alt_text, source):
                nonlocal image_index
                extracted_images.append({
                    "url": image_url,
                    "alt": alt_text,
                    "source": source,
                    "unique_id": f"{timestamp}_{image_index}_{random.randint(1000, 9999)}"  # Add uniqueness
                })
                image_index += 1
            
            # RAG: Extract document text for supported types
            if ext in ['.txt', '.md', '.docx', '.pdf']:
                print(f"[DEBUG] Attempting to extract text from file (ext: {ext})", flush=True)
                doc_text = extract_text_from_file(file_path, ext)
                print(f"[DEBUG] Extracted doc_text (first 500 chars): {doc_text[:500]}", flush=True)
                if not doc_text:
                    print("[DEBUG] doc_text is empty after extraction!", flush=True)
            
            # Existing extraction logic for images and file_content
            if ext in ['.txt', '.md']:
                print("[DEBUG] Reading as plain text or markdown", flush=True)
                with open(file_path, 'r', encoding='utf-8') as f:
                    file_content = f.read()
                print(f"[DEBUG] file_content (first 500 chars): {file_content[:500]}", flush=True)
            elif ext in ['.docx']:
                print("[DEBUG] Reading as DOCX", flush=True)
                try:
                    import docx
                    doc = docx.Document(file_path)
                    
                    # Extract text content
                    paragraphs = []
                    for para in doc.paragraphs:
                        if para.text.strip():
                            paragraphs.append(para.text)
                    file_content = "\n\n".join(paragraphs)
                    print(f"[DEBUG] file_content from DOCX (first 500 chars): {file_content[:500]}", flush=True)
                    
                    # Extract images from docx with improved handling
                    for rel in doc.part.rels.values():
                        if "image" in rel.target_ref:
                            try:
                                image_data = rel.target_part.blob
                                image_ext = rel.target_ref.split(".")[-1] if "." in rel.target_ref else "png"
                                if image_ext not in ['png', 'jpg', 'jpeg', 'gif']:
                                    image_ext = 'png'  # Default to PNG if extension is not recognized
                                
                                # Save the image to the upload folder with unique name
                                image_filename = f"docx_image_{uuid.uuid4().hex[:8]}.{image_ext}"
                                image_path = os.path.join(upload_folder, image_filename)
                                
                                with open(image_path, "wb") as img_file:
                                    img_file.write(image_data)
                                
                                # Create a URL for the saved image
                                image_url = f"/uploads/{image_filename}"
                                add_extracted_image(image_url, f"Image {image_index + 1} from document", "docx")
                                print(f"[DEBUG] Extracted image from DOCX: {image_filename}", flush=True)
                            except Exception as img_err:
                                print(f"Error extracting image from DOCX: {str(img_err)}", flush=True)
                                
                except ImportError:
                    print("[DEBUG] DOCX support not available. Install python-docx package.", flush=True)
                    file_content = "DOCX support not available. Install python-docx package."
                except Exception as docx_err:
                    print(f"[DEBUG] Exception reading DOCX: {docx_err}", flush=True)
            elif ext in ['.pdf']:
                print("[DEBUG] Reading as PDF", flush=True)
                # Extract text and images from PDF with improved handling
                try:
                    pdf_text = []
                    from pypdf import PdfReader
                    pdf_reader = PdfReader(file_path)
                    print(f"[DEBUG] PDF has {len(pdf_reader.pages)} pages", flush=True)
                    # Extract text from each page
                    for page_num in range(len(pdf_reader.pages)):
                        page = pdf_reader.pages[page_num]
                        page_text = page.extract_text()
                        pdf_text.append(page_text)
                        print(f"[DEBUG] Extracted text from PDF page {page_num+1} (first 200 chars): {str(page_text)[:200]}", flush=True)
                    # Join all text
                    file_content = "\n\n".join([text for text in pdf_text if text])
                    print(f"[DEBUG] file_content from PDF (first 500 chars): {file_content[:500]}", flush=True)
                    # ... existing code for image extraction ...
                except Exception as pdf_err:
                    file_content = f"Error extracting PDF content: {str(pdf_err)}"
                    print(f"[DEBUG] Exception reading PDF: {pdf_err}", flush=True)
                # If we couldn't extract any content, use a fallback
                if not file_content or file_content.strip() == "":
                    print("[DEBUG] PDF file uploaded but text could not be extracted.", flush=True)
                    file_content = f"PDF file {file.filename} was uploaded but text could not be extracted."
            elif ext in ['.jpg', '.jpeg', '.png', '.gif']:
                print("[DEBUG] File is an image, not extracting text.", flush=True)
                unique_filename = f"{uuid.uuid4().hex[:8]}_{file.filename}"
                unique_path = os.path.join(upload_folder, unique_filename)
                with open(file_path, 'rb') as src_file:
                    with open(unique_path, 'wb') as dst_file:
                        dst_file.write(src_file.read())
                add_extracted_image(
                    f"/uploads/{unique_filename}",
                    f"Uploaded image - {file.filename}",
                    "uploaded"
                )
                file_content = f"Content for blog with image {file.filename}"
            else:
                print(f"[DEBUG] File type {ext} is not supported for content extraction.", flush=True)
                file_content = f"File type {ext} is not supported for content extraction."
        else:
            print("[DEBUG] No file object after initial check!", flush=True)
        
        # Get template ID from the form
        template_id = request.form.get('template_id', 'tech')
        
        # Load the template
        template = load_template(template_id)
        
        # Create a modified copy to avoid altering the original
        modified_template = copy.deepcopy(template)
        
        # Initialize metadata if not present
        if "metadata" not in modified_template:
            modified_template["metadata"] = {}
        
        # Add randomization metadata to force content variation
        modified_template["metadata"]["variation_factors"] = variation_factors
        modified_template["metadata"]["generation_time"] = timestamp
        modified_template["metadata"]["unique_id"] = random_id
        
        # Add a version number to avoid caching
        modified_template["version"] = f"v{timestamp}-{random.randint(1000, 9999)}"
        
        # Use Gemini model to generate content
        if GEMINI_API_KEY:
            try:
                genai.configure(api_key=GEMINI_API_KEY)
                model = genai.GenerativeModel(model_name="gemma-3n-e4b-it")
                
                {file_content[:1000] if file_content else file.filename}
                blog_prompt = f"""
                IMPORTANT VARIATION INSTRUCTIONS:
                - This is generation attempt #{timestamp % 1000} with unique ID {random_id}
                - Use a {variation_factors['tone_variation']} tone and {variation_factors['content_variation']} style
                - Structure this as a {variation_factors['focus']} type of content
                - Write from a {variation_factors['perspective']} perspective
                - Structure seed: {variation_factors['structure_seed']}
                - Make this blog COMPLETELY DIFFERENT from any previous generations on this topic
                
                Format as JSON with the following structure:
                {{
                    "title": "A catchy title for the blog",
                    "content": [
                        {{"type": "text", "content": "Introduction paragraph"}},
                        {{"type": "heading", "content": "First Section Heading"}},
                        {{"type": "text", "content": "First section content"}},
                        {{"type": "image", "alt": "Description of an image that could go here"}}
                    ]
                }}
                
                Return ONLY valid JSON without any extra text, markdown formatting or explanation.
                
                The blog should have at least 5-6 sections with headings and 3-4 image placeholders.
                """
                
                # Make the API call
                response = model.generate_content(blog_prompt)
                response_text = response.text
                
                # Clean up the response to ensure it's valid JSON
                try:
                    # Remove any markdown formatting or text outside the JSON
                    if "```json" in response_text:
                        json_text = response_text.split("```json")[1].split("```")[0].strip()
                    elif "```" in response_text:
                        json_text = response_text.split("```")[1].strip()
                    else:
                        json_text = response_text.strip()
                    
                    # Parse the JSON content
                    content_updates = json.loads(json_text)
                    
                    # Update the template with the generated content
                    if "title" in content_updates:
                        modified_template["title"] = content_updates["title"]
                    
                    # Process content items and add in images
                    if "content" in content_updates and isinstance(content_updates["content"], list):
                        # Handle image placement
                        new_content = []
                        image_count = 0
                        
                        # Check if the file is an image type
                        is_image_upload = file.content_type and file.content_type.startswith('image/')
                        
                        # Only use user's uploaded image if it's an image
                        if is_image_upload:
                            # Create a path for the uploaded image
                            image_filename = f"user_upload_{timestamp}_{random.randint(1000, 9999)}.jpg"
                            image_path = os.path.join(upload_folder, image_filename)
                            
                            # Try to save a copy of the image
                            try:
                                image = Image.open(file_path)
                                image.save(image_path)
                                
                                # Add the user's image as the first image
                                user_image_url = f"/uploads/{image_filename}"
                                add_extracted_image(user_image_url, "User uploaded image", "user_upload")
                            except Exception as img_err:
                                print(f"Error saving user image: {str(img_err)}")
                        
                        # Process all content items
                        for item in content_updates["content"]:
                            # Clean the content to remove any timestamp references
                            if item.get("type") == "text" and "content" in item:
                                item["content"] = item["content"].replace("Generated content section", "")
                                item["content"] = re.sub(r"with timestamp \d+", "", item["content"])
                                item["content"] = re.sub(r"Generated Section \d+ - \d+", "", item["content"])
                            
                            if item.get("type") == "heading" and "content" in item:
                                item["content"] = item["content"].replace("Generated content section", "")
                                item["content"] = re.sub(r"with timestamp \d+", "", item["content"])
                                item["content"] = re.sub(r"Generated Section \d+ - \d+", "", item["content"])
                            
                            # Add the item to the new content
                            new_content.append(item)
                            # If this is an image item and we have extracted images, use them
                            if item.get("type") == "image" and image_count < len(extracted_images):
                                item["url"] = extracted_images[image_count]["url"]
                                item["unique_id"] = extracted_images[image_count]["unique_id"]
                                image_count += 1
                            
                            modified_template["content"] = new_content
                            try:
                                # Save the template
                                save_template(modified_template)
                            except Exception as save_err:
                                print(f"Error saving template: {str(save_err)}")
                                return jsonify({"error": f"Error saving template: {str(save_err)}"}), 500
                except Exception as json_err:
                    print(f"Error parsing generated content: {str(json_err)}")
                    return jsonify({"error": f"Error parsing generated content: {str(json_err)}"}), 500
                except Exception as gemini_err:
                    print(f"Error generating content with Gemini: {str(gemini_err)}")
                    return jsonify({"error": f"Error generating content: {str(gemini_err)}"}), 500
            except Exception as gemini_err:
                print(f"Error generating content with Gemini: {str(gemini_err)}")
                return jsonify({"error": f"Error generating content: {str(gemini_err)}"}), 500
        
        # Save the generated blog with a timestamp in the filename to avoid caching
        unique_id = f"{template_id}_{timestamp}_{random.randint(1000, 9999)}"
        modified_template["id"] = unique_id
        
        # Save updated template
        save_template(modified_template)
        
        # Return the generated template
        return jsonify({
            "response": f"Successfully generated unique blog based on your uploaded file.",
            "updatedContent": generate_html_preview(modified_template),
            "template": modified_template,
            "templateId": unique_id,
            "timestamp": timestamp,  # Include timestamp for client-side cache busting
            "randomFactor": random.randint(1000000, 9999999)  # Additional randomization
        })
        
    except Exception as e:
        print(f"Error in upload_generate: {str(e)}", flush=True)
        return jsonify({"error": f"Error processing your file: {str(e)}"}), 500

# Helper function to split content into chunks for the blog
def split_content_into_chunks(content, min_chunk_size=200, max_chunk_size=500):
    """Split content into reasonably sized chunks for better readability"""
    # If content is shorter than min_chunk_size, return as is
    if len(content) <= min_chunk_size:
        return [content]
    
    chunks = []
    current_chunk = ""
    sentences = content.split('. ')
    
    for sentence in sentences:
        # Add period back if it was removed by split
        if not sentence.endswith('.'):
            sentence += '.'
        
        # If adding this sentence would make chunk too long, save current chunk and start new one
        if len(current_chunk) + len(sentence) > max_chunk_size and len(current_chunk) >= min_chunk_size:
            chunks.append(current_chunk.strip())
            current_chunk = sentence + ' '
        else:
            current_chunk += sentence + ' '
    
    # Add any remaining content
    if current_chunk:
        chunks.append(current_chunk.strip())
    
    return chunks

# Extract a heading from text
def extract_heading(text, max_length=60):
    """Extract a heading from text, ensuring it's not too long"""
    # Split into sentences
    sentences = text.split('. ')
    
    for sentence in sentences:
        # Clean up the sentence
        heading = sentence.strip()
        
        # Remove any markdown or special characters
        heading = re.sub(r'[#*_]', '', heading)
        
        # If it's a good length, use it
        if 20 <= len(heading) <= max_length:
            return heading
        
        # If it's too long, try to truncate at a word boundary
        elif len(heading) > max_length:
            words = heading.split()
            truncated = ''
            for word in words:
                if len(truncated) + len(word) + 1 <= max_length - 3:  # Leave room for ellipsis
                    truncated += word + ' '
                else:
                    break
            return truncated.strip() + '...'
    
    # If we couldn't find a good sentence, use the first one truncated
   
    if sentences:
       
        heading = sentences[0].strip()
        if len(heading) > max_length:
            return heading[:max_length-3] + '...'
        return heading
    
    return "New Section"  # Fallback if no text provided

# Extract a quote from text
def extract_quote(text, min_length=30, max_length=150):
    """Extract a quote from text that's neither too short nor too long"""
    # Split into sentences
    sentences = text.split('. ')
    
    for sentence in sentences:
        # Clean up the sentence
        quote = sentence.strip()
        
        # Remove any markdown or special characters
        quote = re.sub(r'[#*_]', '', quote)
        
        # Check if it's a good length for a quote
        if min_length <= len(quote) <= max_length:
            return quote
    
    # If we couldn't find a perfect quote, try to combine short sentences or truncate long ones
    current_quote = ""
    for sentence in sentences:
        quote = sentence.strip()
        quote = re.sub(r'[#*_]', '', quote)
        
        if len(current_quote) + len(quote) <= max_length:
            current_quote += quote + '. '
        
        if len(current_quote) >= min_length:
            return current_quote.strip()
    
    # If still no good quote found, take the first sentence and truncate if needed
    if sentences:
        quote = sentences[0].strip()
        quote = re.sub(r'[#*_]', '', quote)
        if len(quote) > max_length:
            return quote[:max_length-3] + '...'
        return quote
    
    return None  # Return None if no suitable quote found

# Extract statistical values from text (numbers, percentages, currency, etc.)
def extract_stat_value(text):
    """Extract a statistical value from text"""
    # Look for numbers with % or common units
    percentage_match = re.search(r'\b(\d+(?:\.\d+)?%)\b', text)
    if percentage_match:
        return percentage_match.group(1)
    
    # Look for currency values
    currency_match = re.search(r'\$\s*(\d+(?:,\d{3})*(?:\.\d{2})?(?:\s*[Kk])?(?:\s*[Mm])?(?:\s*[Bb])?)', text)
    if currency_match:
        return currency_match.group(0)
    
    # Look for numbers with K/M/B suffix
    scale_match = re.search(r'\b(\d+(?:\.\d+)?\s*[KkMmBb])\b', text)
    if scale_match:
        return scale_match.group(1)
    
    # Look for simple numbers
    number_match = re.search(r'\b(\d+(?:,\d{3})*(?:\.\d+)?)\b', text)
    if number_match:
        return number_match.group(1)
    
    return None

# Extract a short piece of text for descriptions
def extract_short_text(text, max_length=100):
    """Extract a short snippet of text, ending at a natural break"""
    if len(text) <= max_length:
        return text
    
    # Try to find a sentence break
    shortened = text[:max_length]
    last_period = shortened.rfind('.')
    if last_period > max_length * 0.5:  # Only use sentence break if it's not too short
        return text[:last_period + 1]
    
    # Try to find a word break
    last_space = shortened.rfind(' ')
    if last_space > 0:
        return text[:last_space] + '...'
    
    # If no good breaks found, just truncate
    return shortened + '...'

# Function to fetch images from Unsplash API with enhanced content-based matching
def fetch_online_image(query, fallback_seed=None, content_context=None):
    """
    Fetch an image URL from Unsplash based on the query, optimized for title+template approach
    
    Args:
        query: Primary search query (now typically title + template)
        fallback_seed: Seed for placeholder generation
        content_context: Additional context (now typically None for focused results)
    """
    try:
        if not UNSPLASH_API_KEY:
            return generate_placeholder_image(fallback_seed, query)
            
        headers = {
            'Authorization': f'Client-ID {UNSPLASH_API_KEY}'
        }
        
        # Template-specific keyword mappings for better image results
        template_keywords = {
            'tech': ['technology', 'computer', 'digital', 'software', 'innovation', 'artificial intelligence', 'coding', 'data'],
            'finance': ['business', 'money', 'financial', 'investment', 'economy', 'market', 'banking', 'growth'],
            'health': ['medical', 'healthcare', 'wellness', 'fitness', 'science', 'research', 'hospital', 'medicine'],
            'travel': ['adventure', 'destination', 'tourism', 'landscape', 'journey', 'explore', 'vacation', 'nature'],
            'creative': ['art', 'design', 'artistic', 'creative', 'visual', 'aesthetic', 'inspiration', 'gallery'],
            'news': ['journalism', 'media', 'information', 'communication', 'breaking', 'current', 'events', 'press']
        }
        
        # Enhanced query generation for title+template approach
        search_queries = []
        
        # Clean the primary query
        clean_query = query.replace('#', '').strip()
        
        # Extract template type from query if present
        template_type = None
        for template in template_keywords.keys():
            if template in clean_query.lower():
                template_type = template
                break
        
        # Build focused search queries
        if template_type and template_type in template_keywords:
            # Use template-specific keywords combined with title words
            title_words = [word for word in clean_query.split() if word.lower() not in template_keywords.keys()]
            template_terms = template_keywords[template_type]
            
            # Primary: template keywords only (most focused)
            search_queries.append(template_terms[0])
            search_queries.append(f"{template_terms[0]} {template_terms[1]}")
            
            # Secondary: combine title with template keywords
            if title_words:
                main_title_word = title_words[0] if title_words else ""
                if main_title_word:
                    search_queries.append(f"{main_title_word} {template_terms[0]}")
                    search_queries.append(f"{template_terms[0]} {main_title_word}")
        else:
            # Fallback: use the original query
            if clean_query:
                search_queries.append(clean_query)
                # Also try just the first few words if query is long
                words = clean_query.split()
                if len(words) > 2:
                    search_queries.append(' '.join(words[:2]))
        
        # Try each search query until we find results
        for search_query in search_queries:
            if not search_query.strip():
                continue
                
            encoded_query = requests.utils.quote(search_query.strip())
            
            try:
                response = requests.get(
                    f'https://api.unsplash.com/search/photos?query={encoded_query}&per_page=5&orientation=landscape',
                    headers=headers,
                    timeout=5
                )
                
                if response.status_code == 200:
                    data = response.json()
                    if data['results']:
                        # Get a random image from top 5 results for variety
                        image_index = random.randint(0, min(4, len(data['results']) - 1))
                        image_data = data['results'][image_index]
                        return {
                            "url": image_data['urls']['regular'],
                            "alt": image_data.get('alt_description', search_query),
                            "credit": f"Photo by {image_data['user']['name']} on Unsplash"
                        }
                        
            except requests.RequestException as e:
                print(f"Request error for query '{search_query}': {str(e)}")
                continue
        
        # If no results from any query, fall back to placeholder
        return generate_placeholder_image(fallback_seed, query)
        
    except Exception as e:
        print(f"Error fetching image: {str(e)}")
        return generate_placeholder_image(fallback_seed, query)

def extract_image_keywords(text, num_keywords=5):
    """
    Extract keywords specifically optimized for image search from blog content
    
    Args:
        text: The blog content to extract keywords from
        num_keywords: Number of keywords to extract
        
    Returns:
        List of keyword strings optimized for image search
    """
    if not text or len(text) < 20:
        return ["abstract", "modern", "professional"]
    
    # Image-friendly keywords that work well with Unsplash
    image_friendly_terms = {
        'technology', 'tech', 'artificial', 'intelligence', 'ai', 'machine', 'learning',
        'data', 'digital', 'computer', 'software', 'programming', 'coding', 'innovation',
        'business', 'finance', 'financial', 'money', 'investment', 'market', 'economy',
        'health', 'medical', 'healthcare', 'fitness', 'wellness', 'science', 'research',
        'education', 'learning', 'study', 'knowledge', 'academic', 'university',
        'travel', 'adventure', 'explore', 'journey', 'destination', 'tourism',
        'nature', 'environment', 'green', 'sustainable', 'ecology', 'landscape',
        'creative', 'design', 'art', 'artistic', 'aesthetic', 'visual', 'graphic',
        'communication', 'network', 'connection', 'social', 'community', 'team',
        'growth', 'success', 'achievement', 'progress', 'development', 'strategy'
    }
    
    # Common stop words
    stop_words = {
        'a', 'an', 'the', 'and', 'or', 'but', 'if', 'because', 'as', 'what', 
        'when', 'where', 'how', 'is', 'are', 'was', 'were', 'be', 'been', 'to', 
        'of', 'for', 'by', 'with', 'about', 'can', 'will', 'would', 'could',
        'this', 'that', 'these', 'those', 'have', 'has', 'had', 'do', 'does', 'did'
    }
    
    # Clean text and extract words
    cleaned_text = re.sub(r'[^\w\s]', ' ', text.lower())
    words = cleaned_text.split()
    
    # Filter and score words
    word_scores = {}
    for word in words:
        if (len(word) > 3 and 
            word not in stop_words and 
            not word.isdigit()):
            
            # Base score
            score = 1
            
            # Boost for image-friendly terms
            if word in image_friendly_terms:
                score += 3
            
            # Boost for longer words (often more specific)
            if len(word) > 6:
                score += 1
            
            word_scores[word] = word_scores.get(word, 0) + score
    
    # Get top keywords
    sorted_words = sorted(word_scores.items(), key=lambda x: x[1], reverse=True)
    keywords = [word for word, score in sorted_words[:num_keywords]]
    
    # Ensure we have some fallback keywords
    if len(keywords) < 2:
        keywords.extend(['modern', 'professional', 'abstract'][:num_keywords - len(keywords)])
    
    return keywords

# Function to generate a placeholder image with more variety and improved quality
def generate_placeholder_image(seed=None, alt_text="Generated image"):
    """
    Generate a high-quality placeholder image when online images are unavailable.
    Creates a visually appealing gradient with text overlay for the given seed.
    
    Args:
        seed: Numeric seed for reproducible generation
        alt_text: Alt text to describe the image
    
    Returns:
        dict with url and alt text
    """
    # Generate a seed if not provided
    if seed is None:
        seed = random.randint(10000, 99999)
    elif isinstance(seed, str) and seed.isdigit():
        seed = int(seed)
    
    # Seed the random generator for reproducible results
    random.seed(seed)
    
    # Generate higher quality images with consistent dimensions for blog posts
    dimensions = [
        (800, 400),   # 2:1 aspect ratio (good for headers)
        (900, 600),   # 3:2 aspect ratio (standard photo)
        (850, 500),   # 17:10 aspect ratio (widescreen)
        (700, 525),   # 4:3 aspect ratio (standard)
        (1000, 500)   # 2:1 aspect ratio (panoramic)
    ]
    
    # Pick one of the predefined dimensions
    width, height = random.choice(dimensions)
    
    # Parse any descriptive text from alt_text to create more relevant placeholders
    if isinstance(alt_text, str) and len(alt_text) > 10:
        # Extract meaningful words from alt text (remove common filler words)
        words = alt_text.split()
        important_words = [w for w in words if len(w) > 3 and w.lower() not in 
                          ('this', 'that', 'with', 'from', 'image', 'picture', 'photo', 'generated')]
        
        # Use up to 3 important words as the image descriptor
        if important_words:
            descriptor = ' '.join(important_words[:3])
            if len(descriptor) > 50:
                descriptor = descriptor[:47] + '...'
        else:
            descriptor = alt_text[:50] + ('...' if len(alt_text) > 50 else '')
    else:
        descriptor = "Generated image"
    
    # Create a more visually appealing URL with context
    image_type = random.choice(['abstract', 'gradient', 'pattern', 'geometric'])
    
    # Clean descriptor for URL (remove special chars, spaces to dashes)
    clean_descriptor = re.sub(r'[^\w\s-]', '', descriptor).strip().lower()
    clean_descriptor = re.sub(r'[\s-]+', '-', clean_descriptor)
    
    # Truncate if too long
    if len(clean_descriptor) > 30:
        clean_descriptor = clean_descriptor[:30]
    
    # Generate a nice URL format
    url_params = f"?type={image_type}&seed={seed}&text={clean_descriptor}"
    
    return {
        "url": f"/api/placeholder/{width}/{height}{url_params}",
        "alt": descriptor
    }

# Set up a route to serve files from the uploads directory
@app.route('/uploads/<filename>')
def uploaded_file(filename):
    file_path = os.path.join(UPLOAD_FOLDER, filename)
    print(f"[DEBUG] Upload request for: {filename}")
    print(f"[DEBUG] Full file path: {file_path}")
    print(f"[DEBUG] File exists: {os.path.exists(file_path)}")
    if os.path.exists(file_path):
        print(f"[DEBUG] File size: {os.path.getsize(file_path)} bytes")
        return send_file(file_path)
    else:
        print(f"[DEBUG] File not found in uploads!")
        # Try social_uploads as fallback
        social_file_path = os.path.join('static', 'social_uploads', filename)
        print(f"[DEBUG] Trying social_uploads path: {social_file_path}")
        if os.path.exists(social_file_path):
            print(f"[DEBUG] Found in social_uploads!")
            return send_file(social_file_path)
        else:
            print(f"[DEBUG] File not found in social_uploads either!")
            return "File not found", 404

def extract_keywords(text, num_keywords=5):
    """
    Extract relevant keywords from text to use as search terms for images.
    Improved to handle PDF content better by focusing on significant terms.
    
    Args:
        text: The text to extract keywords from
        num_keywords: Number of keywords to extract
        
    Returns:
        List of keyword strings
    """
    if not text or len(text) < 50:
        return ["document", "business", "professional", "concept", "abstract"]
        
    # Expanded stop words list to filter out common terms
    stop_words = {
        'a', 'an', 'the', 'and', 'or', 'but', 'if', 'because', 'as', 'what', 
        'when', 'where', 'how', 'is', 'are', 'was', 'were', 'be', 'been', 'to', 
        'of', 'for', 'by', 'with', 'about', 'against', 'between', 'into', 'through',
        'during', 'before', 'after', 'above', 'below', 'from', 'up', 'down', 'in',
        'out', 'on', 'off', 'over', 'under', 'again', 'further', 'then', 'once',
        'here', 'there', 'all', 'any', 'both', 'each', 'few', 'more', 'most', 'other',
        'some', 'such', 'no', 'nor', 'not', 'only', 'own', 'same', 'so', 'than', 'too',
        'very', 's', 't', 'can', 'will', 'just', 'don', 'should', 'now', 'also',
        'figure', 'fig', 'table', 'copyright', 'page', 'chapter', 'et', 'al',
        'this', 'that', 'these', 'those', 'would', 'could', 'may', 'might'
    }
    
    # Clean and normalize text - remove special characters and extra spaces
    # For PDF text specifically, look for patterns like multiple spaces, page numbers, etc.
    cleaned_text = re.sub(r'[^\w\s]', ' ', text.lower())
    cleaned_text = re.sub(r'\s+', ' ', cleaned_text).strip()
    
    # Handle PDF page numbers and headers - common patterns in PDFs
    cleaned_text = re.sub(r'\d+\s*of\s*\d+', '', cleaned_text)  # Remove "Page X of Y"
    cleaned_text = re.sub(r'^\d+$', '', cleaned_text, flags=re.MULTILINE)  # Remove standalone numbers
    cleaned_text = re.sub(r'www\.\w+\.\w+', '', cleaned_text)  # Remove URLs
    
    # Split into words
    words = cleaned_text.split()
    
    # Filter out stop words and words that are too short or numerical
    filtered_words = []
    for word in words:
        # Skip if it's a stop word, too short, or just a number
        if (word not in stop_words and 
            len(word) > 3 and 
            not word.isdigit() and
            not re.match(r'^\d+\w?$', word)):  # Skip patterns like "10th", "2nd", etc.
            filtered_words.append(word)
    
    # Count word frequencies with more weight given to longer words
    word_counts = {}
    for word in filtered_words:
        # Give more weight to longer words - they're often more significant
        weight = min(len(word) - 2, 3)  # Words of length 5+ get weight boost
        if word in word_counts:
            word_counts[word] += 1 + (weight * 0.2)
        else:
            word_counts[word] = 1 + (weight * 0.2)
    
    # Boost for title-case words (likely proper nouns, more significant)
    title_case_pattern = re.compile(r'^[A-Z][a-z]+$')
    for word in text.split():
        if title_case_pattern.match(word) and len(word) > 3:
            word_lower = word.lower()
            if word_lower in word_counts:
                word_counts[word_lower] *= 1.5  # 50% boost for proper nouns
    
    # Get most common words
    sorted_words = sorted(word_counts.items(), key=lambda x: x[1], reverse=True)
    top_keywords = [word for word, count in sorted_words[:num_keywords*2]]
    
    # Prioritize specific contextual keywords that work well for image search
    contextual_keywords = [
        'technology', 'business', 'nature', 'people', 'science', 'travel', 
        'finance', 'health', 'education', 'food', 'industry', 'architecture',
        'medical', 'digital', 'office', 'wildlife', 'landscape', 'urban',
        'research', 'innovation', 'data', 'analysis', 'strategy', 'leadership',
        'development', 'environment', 'communication', 'marketing', 'design'
    ]
    
    # Push contextual keywords that appear in our list to the front
    # These tend to generate better image results
    prioritized_keywords = []
    remaining_keywords = []
    
    for keyword in top_keywords:
        if keyword in contextual_keywords:
            prioritized_keywords.append(keyword)
        else:
            remaining_keywords.append(keyword)
    
    result_keywords = prioritized_keywords + remaining_keywords
    
    # If we don't have enough keywords, add some generic ones
    generic_keywords = [
        'document', 'professional', 'business', 'concept', 'information',
        'technology', 'analysis', 'research', 'data', 'report'
    ]
    
    while len(result_keywords) < num_keywords:
        for keyword in generic_keywords:
            if keyword not in result_keywords:
                result_keywords.append(keyword)
                if len(result_keywords) >= num_keywords:
                    break
    
    # Return unique keywords with the highest counts
    return list(dict.fromkeys(result_keywords))[:num_keywords]

@app.route('/debug_log', methods=['POST'])
def debug_log():
    data = request.get_json()
    msg = data.get('msg', '')
    print(f'[FRONTEND DEBUG] {msg}', file=sys.stderr, flush=True)
    return jsonify({'status': 'ok'})

@app.route('/query_chatbot', methods=['POST'])
def query_chatbot_route():
    try:
        data = request.get_json(force=True, silent=True)
        if not data:
            return jsonify({"error": "No data provided"}), 400
        user_message = data.get('message', '').strip()
        template_id = data.get('template_id', 'tech')
        if not user_message:
            return jsonify({"error": "No message provided"}), 400
        response_text, updated_template = query_chatbot(user_message, template_id)
        # Debug output for diagnosis
        print("[DEBUG] /query_chatbot: user_message={}, template_id={}, response_text={}, updated_template={}".format(
            user_message, template_id, response_text, 'present' if updated_template is not None else 'None'))
        # Always return a valid response, even if the chatbot fails
        if not response_text or (isinstance(response_text, str) and response_text.lower().startswith('error')):
            import traceback
            print("[ERROR] query_chatbot failed. user_message: {}, template_id: {}, response_text: {}".format(
                user_message, template_id, response_text))
            traceback.print_stack()
            response_text = "Sorry, I encountered an error processing your request. Please try again."
        response = {"response": response_text}
        if updated_template is not None:
            response["template"] = updated_template
            response["templateId"] = template_id
        return jsonify(response)
    except Exception as e:
        print("Error in /query_chatbot route: {}".format(str(e)))
        return jsonify({"error": "Internal server error. Please try again later."}), 500
    
@app.route('/generate_blog', methods=['POST'])
def generate_blog():
    print("=== TEST: generate_blog endpoint called ===", flush=True)
    title = request.form.get('title', 'Untitled Blog')
    description = request.form.get('description', '')
    images = request.files.getlist('images')
    template_id = request.form.get('template_id', 'tech')
    template = load_template(template_id)

    # Debug: Check uploaded images
    print(f"[DEBUG] Number of images uploaded: {len(images)}", flush=True)
    image_captions = []
    for idx, img in enumerate(images):
        print(f"[DEBUG] Image {idx}: filename={img.filename}, content_type={img.content_type}", flush=True)
        # Save image to disk and check existence/size
        if img.filename:
            upload_folder = "uploads"
            import os
            os.makedirs(upload_folder, exist_ok=True)
            image_path = os.path.join(upload_folder, img.filename)
            img.save(image_path)
            print(f"[DEBUG] Saved image to {image_path}", flush=True)
            if os.path.exists(image_path):
                file_size = os.path.getsize(image_path)
                print(f"[DEBUG] Image file exists after save. Size: {file_size} bytes", flush=True)
                # Hugging Face captioning
                try:
                    with open(image_path, "rb") as f:
                        img_bytes = f.read()
                    caption = get_image_caption_hf(img_bytes)
                    print(f"[DEBUG] Hugging Face caption for {img.filename}: {caption}")
                    image_captions.append(caption)
                except Exception as e:
                    print(f"[DEBUG] Error reading image for Hugging Face: {e}")
                    image_captions.append("Image uploaded by user")
            else:
                print(f"[DEBUG] Image file does NOT exist after save!", flush=True)
                image_captions.append("Image uploaded by user")
        else:
            image_captions.append("Image uploaded by user")

    # Debug: Check uploaded content files (PDF, DOC, DOCX) and extract text
    content_files = request.files.getlist('content_files')
    extracted_texts = []
    print(f"[DEBUG] Number of content files uploaded: {len(content_files)}", flush=True)
    for idx, cfile in enumerate(content_files):
        print(f"[DEBUG] Content file {idx}: filename={cfile.filename}, content_type={cfile.content_type}", flush=True)
        if cfile.filename:
            upload_folder = "uploads"
            import os
            os.makedirs(upload_folder, exist_ok=True)
            content_path = os.path.join(upload_folder, cfile.filename)
            cfile.save(content_path)
            print(f"[DEBUG] Saved content file to {content_path}", flush=True)
            if os.path.exists(content_path):
                file_size = os.path.getsize(content_path)
                print(f"[DEBUG] Content file exists after save. Size: {file_size} bytes", flush=True)
                # Try to extract text if PDF, DOCX, TXT, or MD
                ext = os.path.splitext(cfile.filename)[1].lower()
                if ext in ['.pdf', '.docx', '.txt', '.md']:
                    print(f"[DEBUG] Attempting to extract text from {content_path} (ext: {ext})", flush=True)
                    doc_text = extract_text_from_file(content_path, ext)
                    if doc_text:
                        extracted_texts.append(doc_text)
                        print(f"[DEBUG] Extracted text (first 500 chars): {doc_text[:500]}", flush=True)
                    else:
                        print(f"[DEBUG] No text extracted from {content_path}", flush=True)
            else:
                print(f"[DEBUG] Content file does NOT exist after save!", flush=True)

    # Enhanced blog generation with uploaded content integration
    uploaded_content_summary = ""
    if image_captions:
        uploaded_content_summary += f"Uploaded Images: {', '.join(image_captions)}. "
    if extracted_texts:
        uploaded_content_summary += f"Uploaded Documents: {len(extracted_texts)} document(s) with relevant content. "
    
    blog_prompt = f"""Write a comprehensive, in-depth blog post that integrates uploaded content with the following details:
Title: {title}
Description: {description}
Template Type: {template.get('type', template_id)}

UPLOADED CONTENT TO INTEGRATE:
{uploaded_content_summary if uploaded_content_summary else "No uploaded content provided"}

CONTENT REQUIREMENTS:
- Create a substantial blog post that naturally incorporates the uploaded content
- If images were uploaded, reference them in relevant sections with descriptive alt text
- If documents were uploaded, extract key insights and integrate them into the blog
- Include a compelling introduction that hooks the reader
- Develop 5-7 main sections with detailed explanations, examples, and insights
- Add relevant statistics, case studies, or real-world examples where appropriate
- Include practical tips, actionable advice, or step-by-step guidance
- Create 4-6 descriptive headings and subheadings to structure the content
- Insert image placeholders that reference the uploaded images when relevant
- Add a comprehensive conclusion that summarizes key points and provides next steps
- Use the professional style and tone appropriate for a {template.get('type', template_id)} blog
- Each text section should be 150-300 words for substantial content depth
- Include quotes, statistics, or expert insights where relevant

SPECIFIC INSTRUCTIONS FOR UPLOADED CONTENT:
- If images were uploaded: Reference them in appropriate sections with descriptive captions
- If documents were uploaded: Extract key points and integrate them naturally into the blog
- Make the uploaded content feel like a natural part of the blog, not forced
- Use the uploaded content to enhance the blog's credibility and relevance

Return the blog as a JSON object with this exact structure (aim for 15-20 content items total):
{{
  "title": "Enhanced title if needed",
  "content": [
    {{"type": "text", "content": "Compelling introduction paragraph that hooks readers..."}},
    {{"type": "heading", "content": "Understanding the Fundamentals"}},
    {{"type": "text", "content": "Comprehensive explanation of core concepts..."}},
    {{"type": "image", "alt": "Descriptive alt text for relevant visual"}},
    {{"type": "text", "content": "Deep dive into specific aspects..."}},
    {{"type": "heading", "content": "Key Benefits and Advantages"}},
    {{"type": "text", "content": "Detailed analysis of benefits..."}},
    {{"type": "list", "items": ["Benefit 1 with explanation", "Benefit 2 with details", "Benefit 3 with examples"]}},
    {{"type": "image", "alt": "Visual representation of benefits"}},
    {{"type": "heading", "content": "Common Challenges and Solutions"}},
    {{"type": "text", "content": "Analysis of typical problems and how to solve them..."}},
    {{"type": "quote", "content": "Relevant expert quote or important insight", "author": "Expert Name"}},
    {{"type": "heading", "content": "Best Practices and Implementation"}},
    {{"type": "text", "content": "Step-by-step guidance and practical tips..."}},
    {{"type": "image", "alt": "Implementation example or process diagram"}},
    {{"type": "text", "content": "Advanced strategies and considerations..."}},
    {{"type": "heading", "content": "Future Trends and Predictions"}},
    {{"type": "text", "content": "Analysis of upcoming developments..."}},
    {{"type": "image", "alt": "Future trends visualization"}},
    {{"type": "text", "content": "Comprehensive conclusion with actionable next steps..."}}
  ]
}}

IMPORTANT: Return only valid JSON. Make each text content substantial and informative."""

    blog_json = None
    ai_error = None

    # Try to use Groq first (since user has Groq API key), then fallback to Gemini
    if GROQ_API_KEY:
        try:
            client = groq.Groq(api_key=GROQ_API_KEY)
            response = client.chat.completions.create(
                model="meta-llama/llama-4-scout-17b-16e-instruct",
                messages=[
                    {"role": "system", "content": "You are a helpful AI that generates blog posts in JSON as described."},
                    {"role": "user", "content": blog_prompt}
                ],
                temperature=0.9,
                max_tokens=4000
            )
            response_text = response.choices[0].message.content
            if "```json" in response_text:
                json_text = response_text.split("```json")[1].split("```")[0].strip()
            elif "```" in response_text:
                json_text = response_text.split("```")[1].strip()
            else:
                json_text = response_text.strip()
            blog_json = json.loads(json_text)
            print("[DEBUG] AI blog JSON (Groq):", blog_json, flush=True)
        except Exception as e:
            ai_error = str(e)
            print("[WARNING] Groq AI error:", e, flush=True)
    elif GEMINI_API_KEY:
        try:
            genai.configure(api_key=GEMINI_API_KEY)
            model = genai.GenerativeModel('gemini-1.5-flash')
            response = model.generate_content(blog_prompt)
            response_text = response.text
            # Extract JSON from response
            if "```json" in response_text:
                json_text = response_text.split("```json")[1].split("```")[0].strip()
            elif "```" in response_text:
                json_text = response_text.split("```")[1].strip()
            else:
                json_text = response_text.strip()
            blog_json = json.loads(json_text)
        except Exception as e:
            ai_error = str(e)
            print("[WARNING] Gemini AI error:", e, flush=True)
    else:
        print("[WARNING] No AI API key set. Falling back to user content.", flush=True)
        ai_error = "No AI API key set. Please set GEMINI_API_KEY or GROQ_API_KEY."

    # Fallback if AI fails
    if not blog_json:
        blog_json = {
            "title": title,
            "content": [
                {"type": "text", "content": description or "No description provided."}
            ]
        }

    # Convert images to data URLs (if any uploaded)
    image_tags = ""
    for img in images:
        img_bytes = img.read()
        mime = img.mimetype or "image/png"
        b64 = base64.b64encode(img_bytes).decode('utf-8')
        image_tags += f'''<div class="image-container">
            <img src="data:{mime};base64,{b64}" alt="Uploaded image">
            <div class="image-credit">User uploaded image</div>
        </div>\n'''

    # Generate HTML from blog_json with enhanced Unsplash images
    html_content = ""
    blog_text_content = " ".join([section.get("content", "") for section in blog_json.get("content", []) if section.get("type") == "text"])
    
    for section in blog_json.get("content", []):
        if section.get("type") == "text":
            html_content += f'<div class="blog-section"><p>{section.get("content", "")}</p></div>\n'
        elif section.get("type") == "heading":
            html_content += f'<div class="blog-section"><h2>{section.get("content", "")}</h2></div>\n'
        elif section.get("type") == "list":
            items = section.get("items", [])
            list_html = "<ul>"
            for item in items:
                list_html += f"<li>{item}</li>"
            list_html += "</ul>"
            html_content += f'<div class="blog-section">{list_html}</div>\n'
        elif section.get("type") == "quote":
            quote_content = section.get("content", "")
            quote_author = section.get("author", "")
            author_html = f'<cite>— {quote_author}</cite>' if quote_author else ""
            html_content += f'''<div class="blog-section">
                <blockquote>
                    "{quote_content}"
                    {author_html}
                </blockquote>
            </div>\n'''
        elif section.get("type") == "image":
            # Create image upload placeholder instead of Unsplash images
            alt = section.get("alt", "Blog image")
            image_id = f"blog_image_{len(html_content)}"
            
            html_content += f'''<div class="image-upload-container" data-image-id="{image_id}">
                <div class="image-upload-placeholder">
                    <div class="upload-icon">📷</div>
                    <div class="upload-text">Click to upload image</div>
                    <div class="upload-hint">Recommended: 800x400px or larger</div>
                    <input type="file" class="image-upload-input" accept="image/*" data-image-id="{image_id}" style="display: none;">
                </div>
                <div class="image-preview" style="display: none;">
                    <img src="" alt="{alt}" class="uploaded-image">
                    <button type="button" class="remove-image-btn" data-image-id="{image_id}">✕</button>
                </div>
            </div>\n'''

    # Add uploaded images at the end
    html_content += image_tags

    # Beautiful, modern HTML template with enhanced design
    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{blog_json.get("title", title)}</title>
    <link href="https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&family=Playfair+Display:wght@400;500;600;700&display=swap" rel="stylesheet">
    <style>
        * {{
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }}
        
        body {{
            font-family: 'Inter', -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
            background: linear-gradient(135deg, #0f1419 0%, #1a1f2e 100%);
            color: #e8eaed;
            margin: 0;
            padding: 20px 0 60px 0;
            line-height: 1.6;
            -webkit-font-smoothing: antialiased;
            -moz-osx-font-smoothing: grayscale;
        }}
        
        .container {{
            max-width: 900px;
            margin: 0 auto;
            background: linear-gradient(145deg, #1e2328 0%, #252a31 100%);
            border-radius: 24px;
            box-shadow: 
                0 20px 60px rgba(0, 0, 0, 0.4),
                0 8px 25px rgba(0, 0, 0, 0.15),
                inset 0 1px 0 rgba(255, 255, 255, 0.1);
            padding: 50px;
            border: 1px solid rgba(255, 255, 255, 0.1);
            overflow: hidden;
            position: relative;
        }}
        
        .container::before {{
            content: '';
            position: absolute;
            top: 0;
            left: 0;
            right: 0;
            height: 4px;
            background: linear-gradient(90deg, #4f8cff 0%, #7c3aed 50%, #06b6d4 100%);
        }}
        
        h1 {{
            font-family: 'Playfair Display', Georgia, serif;
            color: #ffffff;
            font-size: 3.2rem;
            font-weight: 700;
            margin-bottom: 30px;
            line-height: 1.2;
            text-align: center;
            background: linear-gradient(135deg, #4f8cff 0%, #7c3aed 100%);
            -webkit-background-clip: text;
            -webkit-text-fill-color: transparent;
            background-clip: text;
            position: relative;
        }}
        
        h1::after {{
            content: '';
            position: absolute;
            bottom: -15px;
            left: 50%;
            transform: translateX(-50%);
            width: 80px;
            height: 3px;
            background: linear-gradient(90deg, #4f8cff 0%, #7c3aed 100%);
            border-radius: 2px;
        }}
        
        .blog-section {{
            margin-bottom: 35px;
        }}
        
        .blog-section h2 {{
            font-family: 'Playfair Display', Georgia, serif;
            font-size: 2.2rem;
            font-weight: 600;
            color: #ffffff;
            margin-bottom: 20px;
            margin-top: 45px;
            line-height: 1.3;
            position: relative;
            padding-left: 20px;
        }}
        
        .blog-section h2::before {{
            content: '';
            position: absolute;
            left: 0;
            top: 5px;
            width: 4px;
            height: 35px;
            background: linear-gradient(135deg, #4f8cff 0%, #7c3aed 100%);
            border-radius: 2px;
        }}
        
        .blog-section p {{
            font-size: 1.2rem;
            line-height: 1.8;
            color: #d1d5db;
            margin-bottom: 20px;
            text-align: justify;
            font-weight: 400;
            letter-spacing: 0.01em;
        }}
        
        .blog-section p:first-of-type {{
            font-size: 1.25rem;
            color: #e5e7eb;
            font-weight: 400;
        }}
        
        .blog-section ul {{
            color: #d1d5db;
            font-size: 1.2rem;
            line-height: 1.8;
            margin: 20px 0;
            padding-left: 20px;
        }}
        
        .blog-section ul li {{
            margin-bottom: 12px;
            position: relative;
        }}
        
        .blog-section ul li::marker {{
            color: #4f8cff;
        }}
        
        .blog-section blockquote {{
            background: rgba(79, 140, 255, 0.1);
            border-left: 4px solid #4f8cff;
            margin: 30px 0;
            padding: 25px 30px;
            border-radius: 12px;
            font-size: 1.3rem;
            font-style: italic;
            color: #e5e7eb;
            position: relative;
        }}
        
        .blog-section blockquote::before {{
            content: '"';
            font-size: 4rem;
            color: #4f8cff;
            position: absolute;
            top: -10px;
            left: 15px;
            opacity: 0.3;
        }}
        
        .blog-section blockquote cite {{
            display: block;
            text-align: right;
            margin-top: 15px;
            font-size: 1rem;
            color: #9ca3af;
            font-style: normal;
        }}
        
        .image-upload-container {{
            margin: 45px 0;
            text-align: center;
            background: rgba(255, 255, 255, 0.02);
            border-radius: 20px;
            padding: 20px;
            border: 1px solid rgba(255, 255, 255, 0.05);
        }}
        
        .image-upload-placeholder {{
            border: 2px dashed #4f8cff;
            border-radius: 12px;
            padding: 40px 20px;
            background: linear-gradient(145deg, #1a1f2e 0%, #252a31 100%);
            cursor: pointer;
            transition: all 0.3s ease;
            position: relative;
        }}
        
        .image-upload-placeholder:hover {{
            border-color: #7c3aed;
            background: linear-gradient(145deg, #252a31 0%, #2a2f36 100%);
            transform: translateY(-2px);
        }}
        
        .upload-icon {{
            font-size: 3rem;
            margin-bottom: 15px;
            opacity: 0.7;
        }}
        
        .upload-text {{
            font-size: 1.1rem;
            font-weight: 500;
            color: #4f8cff;
            margin-bottom: 8px;
        }}
        
        .upload-hint {{
            font-size: 0.85rem;
            color: #a0a0a0;
        }}
        
        .image-preview {{
            position: relative;
            display: inline-block;
        }}
        
        .uploaded-image {{
            display: block;
            margin: 0 auto;
            max-width: 100%;
            width: auto;
            max-height: 500px;
            border-radius: 16px;
            box-shadow: 
                0 25px 50px rgba(0, 0, 0, 0.3),
                0 12px 25px rgba(0, 0, 0, 0.2);
            transition: transform 0.3s ease, box-shadow 0.3s ease;
            border: 1px solid rgba(255, 255, 255, 0.1);
        }}
        
        .remove-image-btn {{
            position: absolute;
            top: -10px;
            right: -10px;
            background: #ff4757;
            color: white;
            border: none;
            border-radius: 50%;
            width: 30px;
            height: 30px;
            cursor: pointer;
            font-size: 14px;
            display: flex;
            align-items: center;
            justify-content: center;
            box-shadow: 0 4px 12px rgba(255, 71, 87, 0.3);
            transition: all 0.3s ease;
        }}
        
        .remove-image-btn:hover {{
            background: #ff3742;
            transform: scale(1.1);
        }}
        
        .image-container img:hover {{
            transform: translateY(-5px);
            box-shadow: 
                0 35px 70px rgba(0, 0, 0, 0.4),
                0 15px 35px rgba(0, 0, 0, 0.3);
        }}
        
        .image-credit {{
            color: #9ca3af;
            font-size: 0.9rem;
            margin-top: 12px;
            font-style: italic;
            opacity: 0.8;
        }}
        
        .date-author {{
            text-align: center;
            color: #9ca3af;
            font-size: 1rem;
            margin-bottom: 40px;
            padding-bottom: 20px;
            border-bottom: 1px solid rgba(255, 255, 255, 0.1);
        }}
        
        .reading-progress {{
            position: fixed;
            top: 0;
            left: 0;
            width: 0%;
            height: 3px;
            background: linear-gradient(90deg, #4f8cff 0%, #7c3aed 100%);
            z-index: 1000;
            transition: width 0.3s ease;
        }}
        
        @media (max-width: 768px) {{
            .container {{
                margin: 10px;
                padding: 30px 25px;
                border-radius: 16px;
            }}
            
            h1 {{
                font-size: 2.4rem;
                margin-bottom: 25px;
            }}
            
            .blog-section h2 {{
                font-size: 1.8rem;
                margin-top: 30px;
                margin-bottom: 15px;
            }}
            
            .blog-section p {{
                font-size: 1.1rem;
                line-height: 1.7;
            }}
            
            .image-container {{
                margin: 30px 0;
                padding: 15px;
            }}
            
            .image-container img {{
                max-height: 300px;
            }}
        }}
        
        @media (max-width: 480px) {{
            .container {{
                margin: 5px;
                padding: 20px 15px;
            }}
            
            h1 {{
                font-size: 2rem;
            }}
            
            .blog-section h2 {{
                font-size: 1.6rem;
            }}
            
            .blog-section p {{
                font-size: 1rem;
            }}
        }}
    </style>
</head>
<body>
    <div class="reading-progress"></div>
    <div class="container">
        <h1>{blog_json.get("title", title)}</h1>
        <div class="date-author">
            <span>Published on {time.strftime('%B %d, %Y')}</span>
        </div>
        {html_content}
    </div>
    
    <script>
        // Reading progress indicator
        window.addEventListener('scroll', function() {{
            const winScroll = document.body.scrollTop || document.documentElement.scrollTop;
            const height = document.documentElement.scrollHeight - document.documentElement.clientHeight;
            const scrolled = (winScroll / height) * 100;
            document.querySelector('.reading-progress').style.width = scrolled + '%';
        }});
        
        // Smooth scroll for better UX
        document.documentElement.style.scrollBehavior = 'smooth';
    </script>
</body>
</html>
"""

    # Return JSON response for frontend
    return jsonify({
        'html_content': html,
        'title': blog_json.get("title", title),
        'status': 'success'
    })

# Helper function to extract text from uploaded files (PDF, DOCX, TXT, MD)
def extract_text_from_file(file_path, ext):
    if ext == '.pdf':
        try:
            from pypdf import PdfReader
            pdf_reader = PdfReader(file_path)
            pdf_text = []
            for page in pdf_reader.pages:
                pdf_text.append(page.extract_text())
            return "\n\n".join([t for t in pdf_text if t])
        except Exception as e:
            return ""
    elif ext == '.docx':
        try:
            import docx
            doc = docx.Document(file_path)
            return "\n\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        except Exception as e:
            return ""
    elif ext in ['.txt', '.md']:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    return ""

def extract_texts_from_uploaded_files(content_files, upload_dir="uploads"):
    """Extract and concatenate text from all uploaded content files."""
    texts = []
    for cfile in content_files:
        if cfile and cfile.filename:
            ext = os.path.splitext(cfile.filename)[1].lower()
            temp_path = os.path.join(upload_dir, secure_filename(cfile.filename))
            cfile.save(temp_path)
            print(f"[DEBUG] Saved uploaded file to {temp_path}")
            text = extract_text_from_file(temp_path, ext)
            print(f"[DEBUG] Extracted text from {temp_path} (first 300 chars): {text[:300]}")
            if text:
                texts.append(text)
            else:
                print(f"[WARNING] No text extracted from {temp_path}")
    full_text = "\n\n".join(texts)
    print(f"[DEBUG] Full concatenated extracted text (first 1000 chars): {full_text[:1000]}")
    return full_text

@app.route('/generate_social_posts', methods=['POST'])
def generate_social_posts():
    try:
        print("[DEBUG] /generate_social_posts endpoint called.")
        title = request.form.get('title', '').strip()
        description = request.form.get('description', '').strip()
        hashtags = request.form.get('hashtags', '').strip()
        generate_by_ai = request.form.get('generate_by_ai', 'false').lower() == 'true'
        platforms = request.form.get('platforms', '[]')
        import json as _json
        try:
            platforms = _json.loads(platforms)
            print(f"[DEBUG] Parsed platforms: {platforms}")
        except Exception as e:
            print(f"[WARNING] Could not parse platforms: {e}")
            platforms = []
        
        print(f"[DEBUG] Selected platforms: {platforms}")
        print(f"[DEBUG] Facebook in platforms: {'facebook' in platforms}")
        
        # Handle uploaded images first
        images = request.files.getlist('images')
        print(f"[DEBUG] Number of images received: {len(images)}")
        image_urls = []
        
        # Save uploaded images and get their URLs
        for idx, image in enumerate(images):
            if image and image.filename:
                # Create unique filename to avoid conflicts
                timestamp = int(time.time())
                base_name, ext = os.path.splitext(secure_filename(image.filename))
                image_filename = f"social_{timestamp}_{base_name}{ext}"
                image_path = os.path.join(app.config['UPLOAD_FOLDER'], image_filename)
                
                print(f"[DEBUG] Processing image {idx + 1}: {image.filename}")
                print(f"[DEBUG] Secure filename: {image_filename}")
                print(f"[DEBUG] Full image path: {image_path}")
                
                # Save the image
                image.save(image_path)
                image_url = f"/uploads/{image_filename}"
                image_urls.append(image_url)
                
                print(f"[DEBUG] Saved image: {image_url}")
                print(f"[DEBUG] Image file exists after save: {os.path.exists(image_path)}")
                if os.path.exists(image_path):
                    file_size = os.path.getsize(image_path)
                    print(f"[DEBUG] Image file size: {file_size} bytes")
                    
                    # Verify it's a valid image
                    try:
                        from PIL import Image as PILImage
                        with PILImage.open(image_path) as img:
                            print(f"[DEBUG] Image format: {img.format}, size: {img.size}")
                    except Exception as img_err:
                        print(f"[DEBUG] Image validation error: {img_err}")
                else:
                    print(f"[DEBUG] ERROR: Image file not found after save!")
        
        print(f"[DEBUG] Final image URLs: {image_urls}")
        print(f"[DEBUG] Total images processed: {len(image_urls)}")

        # Handle content files for AI generation if needed
        content_files = request.files.getlist('content_files')
        print(f"[DEBUG] Number of content files received: {len(content_files)}")
        rag_context = extract_texts_from_uploaded_files(content_files) if content_files else ""
        
        if generate_by_ai:
            if rag_context:
                print(f"[DEBUG] FINAL rag_context to AI (first 1000 chars): {rag_context[:1000]}")
                debug_context_source = 'document'
            elif description:
                rag_context = description
                print("[DEBUG] No document text extracted, falling back to description.")
                print(f"[DEBUG] FINAL rag_context to AI (first 1000 chars): {rag_context[:1000]}")
                debug_context_source = 'description'
            else:
                rag_context = 'No content provided.'
                print("[WARNING] No document or description provided. Using fallback content.")
                debug_context_source = 'fallback'
                print(f"[DEBUG] FINAL rag_context to AI (first 1000 chars): {rag_context[:1000]}")

        posts_html = ''
        platform_styles = {
            'instagram': {'icon': 'fab fa-instagram', 'color': '#E1306C', 'label': 'Instagram'},
            'facebook': {'icon': 'fab fa-facebook-f', 'color': '#1877F3', 'label': 'Facebook'},
            'twitter': {'icon': 'fab fa-twitter', 'color': '#1DA1F2', 'label': 'Twitter'},
            'linkedin': {'icon': 'fab fa-linkedin-in', 'color': '#0077B5', 'label': 'LinkedIn'},
            'blog': {'icon': 'fas fa-blog', 'color': '#4f8cff', 'label': 'Blog'},
        }

        if generate_by_ai:
            print(f"[DEBUG] AI generation requested for social posts. Using context from: {debug_context_source}")
            ai_posts = {}
            for platform in platforms:
                if platform == 'blog':
                    continue
                style = platform_styles.get(platform, {})
                prompt = f"""
You are an expert social media copywriter.

Below is the content from a document uploaded by the user. Base the post entirely on this content. Do not use generic content unless it is present in the document.

Document content:
{rag_context}

Platform: {platform.capitalize()}
Title: {title}
Hashtags: {hashtags}

Instructions:
- Summarize or adapt the document content for a {platform} post.
- If no document content is provided, use the title and hashtags to create a post.
- Don't use words like "Here is the post you asked for" or "let me know if any update needed". Just the post content alone.
- If no hashtags are provided, use relevant hashtags based on the content.
- Do NOT invent or add unrelated content.
- Use the user's document as the main source.
- Return only the post text (no markdown, no explanation).
"""

                # Add LinkedIn-specific instructions if platform is LinkedIn
                if platform.lower() == 'linkedin':
                    prompt += """
For LinkedIn specifically:
- Start with an engaging hook like "Exciting news!", "I'm thrilled to share...", "Breaking: ", "Just dropped: ", or similar attention-grabbing phrases
- Include "Provided by [Your Company Name]" at the end if it's a company post
- Keep the tone professional but engaging
- Use line breaks (\\n) to separate paragraphs
- End with a call to action or question to encourage engagement
- Format hashtags at the end, separated by spaces
"""

                print(f"[DEBUG] Prompt for {platform} (first 500 chars): {prompt[:500]}")
                ai_content = None
                try:
                    # If images are uploaded, use Groq
                    if image_urls:
                        if not GROQ_API_KEY:
                            print(f"[WARNING] No Groq API key set. Using user content for {platform}.")
                            ai_content = description
                        else:
                            client = groq.Groq(api_key=GROQ_API_KEY)
                            response = client.chat.completions.create(
                                model="meta-llama/llama-4-scout-17b-16e-instruct",
                                messages=[
                                    {"role": "system", "content": "You are a helpful AI that generates social media posts. Keep the content concise and engaging."},
                                    {"role": "user", "content": prompt}
                                ],
                                temperature=0.7,
                                max_tokens=300
                            )
                            ai_content = response.choices[0].message.content.strip()
                            print(f"[DEBUG] AI social post for {platform} (Groq, first 300 chars): {ai_content[:300]}")
                    # For text-only content, use Gemini
                    else:
                        if not GEMINI_API_KEY:
                            print(f"[WARNING] No Gemini API key set. Using user content for {platform}.")
                            ai_content = description
                        else:
                            try:
                                genai.configure(api_key=GEMINI_API_KEY)
                                model = genai.GenerativeModel(model_name="gemini-pro")
                                response = model.generate_content(prompt)
                                ai_content = response.text.strip()
                                print(f"[DEBUG] AI social post for {platform} (Gemini, first 300 chars): {ai_content[:300]}")
                            except Exception as e:
                                print(f"[WARNING] Gemini AI error for {platform}:", e)
                                # If Gemini fails, try Groq as fallback
                                if GROQ_API_KEY:
                                    try:
                                        client = groq.Groq(api_key=GROQ_API_KEY)
                                        response = client.chat.completions.create(
                                            model="meta-llama/llama-4-scout-17b-16e-instruct",
                                            messages=[
                                                {"role": "system", "content": "You are a helpful AI that generates social media posts. Keep the content concise and engaging."},
                                                {"role": "user", "content": prompt}
                                            ],
                                            temperature=0.7,
                                            max_tokens=300
                                        )
                                        ai_content = response.choices[0].message.content.strip()
                                        print(f"[DEBUG] AI social post for {platform} (Groq fallback, first 300 chars): {ai_content[:300]}")
                                    except Exception as groq_error:
                                        print(f"[ERROR] Groq fallback also failed for {platform}:", groq_error)
                                        ai_content = description
                                else:
                                    ai_content = description
                except Exception as e:
                    print(f"[ERROR] AI generation failed for {platform}: {str(e)}")
                    ai_content = description  # Fallback to description if all AI attempts fail
                
                ai_posts[platform] = ai_content

        # Generate HTML for each platform
        for platform in platforms:
            if platform == 'blog':
                continue

            info = platform_styles.get(platform, {'icon': '', 'color': '#4f8cff', 'label': platform.capitalize()})
            post_title = f"<h3>{title}</h3>" if title else ""
            
            # Use AI content if available, otherwise use description
            post_content = ai_posts.get(platform, description) if generate_by_ai else description
            post_desc = f"<p class='post-text'>{post_content}</p>" if post_content else ""
            
            # Handle images for the post
            post_image = ""
            if image_urls:
                print(f"[DEBUG] Creating carousel for {platform} with {len(image_urls)} images")
                carousel_items = ""
                indicators = ""
                for idx, image_url in enumerate(image_urls):
                    print(f"[DEBUG] Adding image {idx + 1} to carousel: {image_url}")
                    carousel_items += f"""
                        <div class="carousel-item {'active' if idx == 0 else ''}" 
                             style="background-image: url('{image_url}'); 
                                    background-size: {'cover' if platform in ['instagram', 'twitter'] else 'contain'}; 
                                    background-position: center;
                                    background-repeat: no-repeat;">
                        </div>
                    """
                    indicators += f"""
                        <span class="carousel-indicator {'active' if idx == 0 else ''}" data-index="{idx}"></span>
                    """
                
                post_image = f"""
                    <div class="post-carousel">
                        <div class="carousel-container">
                            {carousel_items}
                            {f'''
                                <button class="carousel-prev"><i class="fas fa-chevron-left"></i></button>
                                <button class="carousel-next"><i class="fas fa-chevron-right"></i></button>
                                <div class="carousel-indicators">{indicators}</div>
                            ''' if len(image_urls) > 1 else ''}
                        </div>
                    </div>
                """
                print(f"[DEBUG] Generated carousel HTML for {platform} (length: {len(post_image)})")
            else:
                print(f"[DEBUG] No images for {platform}")

            # Format hashtags
            formatted_hashtags = ' '.join(f"<span class='hashtag'>#{tag.strip('#')}</span>" for tag in hashtags.split() if tag.strip())

            # Platform-specific formatting
            if platform == "instagram":
                print(f"[DEBUG] Generating Instagram post card for platform: {platform}")
                print(f"[DEBUG] Title: {title}")
                print(f"[DEBUG] Post content: {post_content[:100]}...")
                print(f"[DEBUG] Hashtags: {hashtags}")
                print(f"[DEBUG] Image URLs for Instagram: {image_urls}")
                print(f"[DEBUG] Number of images: {len(image_urls)}")
                
                post_html = f"""
                    <div class='post-card' data-platform='instagram'>
                        <div class='post-header'>
                            <div class='platform-icon' style='background: linear-gradient(45deg, #f09433, #e6683c, #dc2743, #cc2366, #bc1888);'>
                                <i class='{info['icon']}' style='color: white;'></i>
                            </div>
                            <div class='account-info'>
                                <span class='account-name'>Your Instagram</span>
                                <span class='post-time'>Just now</span>
                            </div>
                        </div>
                        {post_image}
                        <div class='post-content'>
                            <div class='caption-text'>{post_content}</div>
                            <div class='post-hashtags'>{formatted_hashtags}</div>
                        </div>
                        <div class='post-actions' style='margin-top: 15px; display: flex; gap: 10px; padding: 10px; border-top: 1px solid #333;'>
                            <button class='action-btn instagram-post-btn' onclick='postToInstagramWithData(this)' 
                                    style='background: linear-gradient(45deg, #f09433, #e6683c, #dc2743, #cc2366, #bc1888); color: white; border: none; padding: 8px 16px; border-radius: 6px; cursor: pointer; font-weight: 500;'
                                    data-title='{title.replace("'", "&apos;").replace('"', "&quot;")}' 
                                    data-description='{post_content.replace("'", "&apos;").replace('"', "&quot;")}' 
                                    data-hashtags='{hashtags.replace("'", "&apos;").replace('"', "&quot;")}'
                                    data-image-url='{image_urls[0] if image_urls else ""}'>
                                <i class='fab fa-instagram'></i> Post to Instagram
                            </button>
                            <button class='action-btn edit-btn' style='background: #4f8cff; color: white; border: none; padding: 8px 16px; border-radius: 6px; cursor: pointer;'>Edit</button>
                            <button class='action-btn download-btn' style='background: #28a745; color: white; border: none; padding: 8px 16px; border-radius: 6px; cursor: pointer;'>Download</button>
                            <button class='action-btn delete-btn' style='background: #dc3545; color: white; border: none; padding: 8px 16px; border-radius: 6px; cursor: pointer;'>Delete</button>
                        </div>
                    </div>
                """
                print(f"[DEBUG] Instagram post HTML generated, length: {len(post_html)} characters")
                print(f"[DEBUG] Button onclick attribute: postToInstagramWithData(this)")
                print(f"[DEBUG] Button data-image-url attribute: {image_urls[0] if image_urls else 'EMPTY'}")
            elif platform == "facebook":
                print(f"[DEBUG] Generating Facebook post card for platform: {platform}")
                print(f"[DEBUG] Title: {title}")
                print(f"[DEBUG] Post content: {post_content[:100]}...")
                print(f"[DEBUG] Hashtags: {hashtags}")
                print(f"[DEBUG] Image URLs for Facebook: {image_urls}")
                print(f"[DEBUG] Number of images: {len(image_urls)}")
                
                post_html = f"""
                    <div class='post-card' data-platform='facebook'>
                        <div class='post-header'>
                            <div class='platform-icon' style='background-color: #1877f3;'>
                                <i class='{info['icon']}' style='color: white;'></i>
                            </div>
                            <div class='account-info'>
                                <span class='account-name'>Your Facebook</span>
                                <span class='post-time'>Just now</span>
                            </div>
                        </div>
                        <div class='post-content'>
                            <h3 class='post-title'>{title}</h3>
                            <div class='post-text'>{post_content}</div>
                            {post_image}
                            <div class='post-hashtags'>{formatted_hashtags}</div>
                        </div>
                        <div class='post-actions' style='margin-top: 15px; display: flex; gap: 10px; padding: 10px; border-top: 1px solid #333;'>
                            <button class='action-btn facebook-post-btn' onclick='postToFacebookWithData(this)' 
                                    style='background: #1877f3; color: white; border: none; padding: 8px 16px; border-radius: 6px; cursor: pointer; font-weight: 500;'
                                    data-title='{title.replace("'", "&apos;").replace('"', "&quot;")}' 
                                    data-description='{post_content.replace("'", "&apos;").replace('"', "&quot;")}' 
                                    data-hashtags='{hashtags.replace("'", "&apos;").replace('"', "&quot;")}'
                                    data-image-url='{image_urls[0] if image_urls else ""}'>
                                <i class='fab fa-facebook-f'></i> Post to Facebook
                            </button>
                            <button class='action-btn edit-btn' style='background: #4f8cff; color: white; border: none; padding: 8px 16px; border-radius: 6px; cursor: pointer;'>Edit</button>
                            <button class='action-btn download-btn' style='background: #28a745; color: white; border: none; padding: 8px 16px; border-radius: 6px; cursor: pointer;'>Download</button>
                            <button class='action-btn delete-btn' style='background: #dc3545; color: white; border: none; padding: 8px 16px; border-radius: 6px; cursor: pointer;'>Delete</button>
                        </div>
                    </div>
                """
                print(f"[DEBUG] Facebook post HTML generated, length: {len(post_html)} characters")
                print(f"[DEBUG] Button onclick attribute: postToFacebookWithData(this)")
                print(f"[DEBUG] Button data-image-url attribute: {image_urls[0] if image_urls else 'EMPTY'}")
            elif platform == "twitter":
                post_html = f"""
                    <div class='post-card' data-platform='twitter'>
                        <div class='post-header'>
                            <div class='platform-icon' style='background-color: #1da1f2;'>
                                <i class='{info['icon']}' style='color: white;'></i>
                            </div>
                            <div class='account-info'>
                                <span class='account-name'>Your Twitter</span>
                                <span class='twitter-handle'>@yourhandle</span>
                                <span class='post-time'>· Just now</span>
                            </div>
                        </div>
                        <div class='post-content'>
                            <div class='tweet-text'>{post_content}</div>
                            {post_image}
                            <div class='post-hashtags'>{formatted_hashtags}</div>
                        </div>
                    </div>
                """
            elif platform == "linkedin":
                print(f"[DEBUG] Generating LinkedIn post card for platform: {platform}")
                print(f"[DEBUG] Title: {title}")
                print(f"[DEBUG] Post content: {post_content[:100]}...")
                print(f"[DEBUG] Hashtags: {hashtags}")
                print(f"[DEBUG] Image URLs for LinkedIn: {image_urls}")
                print(f"[DEBUG] Number of images: {len(image_urls)}")
                
                post_html = f"""
                    <div class='post-card' data-platform='linkedin'>
                        <div class='post-header'>
                            <div class='platform-icon' style='background-color: #0077B5;'>
                                <i class='{info['icon']}' style='color: white;'></i>
                            </div>
                            <div class='account-info'>
                                <span class='account-name'>Your LinkedIn</span>
                                <span class='post-time'>Just now</span>
                            </div>
                        </div>
                        <div class='post-content'>
                            <h3 class='post-title'>{title}</h3>
                            <div class='post-text'>{post_content}</div>
                            {post_image}
                            <div class='post-hashtags'>{formatted_hashtags}</div>
                        </div>
                        <div class='post-actions' style='margin-top: 15px; display: flex; gap: 10px; padding: 10px; border-top: 1px solid #333;'>
                            <button class='action-btn linkedin-post-btn' onclick='postToLinkedInWithData(this)' 
                                    style='background: #0077B5; color: white; border: none; padding: 8px 16px; border-radius: 6px; cursor: pointer; font-weight: 500;'
                                    data-title='{title.replace("'", "&apos;").replace('"', "&quot;")}' 
                                    data-description='{post_content.replace("'", "&apos;").replace('"', "&quot;")}' 
                                    data-hashtags='{hashtags.replace("'", "&apos;").replace('"', "&quot;")}'
                                    data-image-url='{image_urls[0] if image_urls else ""}'>
                                <i class='fab fa-linkedin-in'></i> Post to LinkedIn
                            </button>
                            <button class='action-btn edit-btn' style='background: #4f8cff; color: white; border: none; padding: 8px 16px; border-radius: 6px; cursor: pointer;'>Edit</button>
                            <button class='action-btn download-btn' style='background: #28a745; color: white; border: none; padding: 8px 16px; border-radius: 6px; cursor: pointer;'>Download</button>
                            <button class='action-btn delete-btn' style='background: #dc3545; color: white; border: none; padding: 8px 16px; border-radius: 6px; cursor: pointer;'>Delete</button>
                        </div>
                    </div>
                """
                print(f"[DEBUG] LinkedIn post HTML generated, length: {len(post_html)} characters")
                print(f"[DEBUG] Button onclick attribute: postToLinkedInWithData(this)")
                print(f"[DEBUG] Button data-image-url attribute: {image_urls[0] if image_urls else 'EMPTY'}")
            else:
                post_html = f"""
                    <div class='post-card' data-platform='{platform}'>
                        <div class='post-header'>
                            <div class='platform-icon' style='background-color: {info['color']};'>
                                <i class='{info['icon']}' style='color: white;'></i>
                            </div>
                            <div class='account-info'>
                                <span class='account-name'>Your {platform.capitalize()}</span>
                                <span class='post-time'>Just now</span>
                            </div>
                        </div>
                        <div class='post-content'>
                            <h3 class='post-title'>{title}</h3>
                            <div class='post-text'>{post_content}</div>
                            {post_image}
                            <div class='post-hashtags'>{formatted_hashtags}</div>
                        </div>
                    </div>
                """
            posts_html += post_html

        return jsonify({'posts_html': posts_html})
    except Exception as e:
        print("[ERROR] Exception in generate_social_posts:", e)
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

def clean_ai_content(content):
    import re
    content = content.strip()
    content = re.sub(r'^```[a-zA-Z]*\n?', '', content)
    content = re.sub(r'```$', '', content)
    return content.strip()

# Helper: Hugging Face image captioning
import requests as _requests

def get_image_caption_hf(image_bytes):
    if not HUGGINGFACE_API_KEY:
        print("[DEBUG] Hugging Face API key not found in environment.")
        return "Image uploaded by user"
    api_url = "https://api-inference.huggingface.co/models/Salesforce/blip-image-captioning-base"
    headers = {"Authorization": f"Bearer {HUGGINGFACE_API_KEY}"}
    try:
        resp = _requests.post(api_url, headers=headers, files={"file": image_bytes})
        if resp.status_code == 200:
            result = resp.json()
            if isinstance(result, list) and result and 'generated_text' in result[0]:
                print(f"[DEBUG] Hugging Face caption: {result[0]['generated_text']}")
                return result[0]['generated_text']
            print(f"[DEBUG] Unexpected Hugging Face response: {result}")
            return "Image uploaded by user"
        else:
            print(f"[DEBUG] Hugging Face API error: {resp.status_code} {resp.text}")
            return "Image uploaded by user"
    except Exception as e:
        print(f"[DEBUG] Exception during Hugging Face API call: {e}")
        return "Image uploaded by user"

@app.route('/debug_uploads', methods=['GET'])
def debug_uploads():
    """Debug endpoint to check uploads directory status"""
    try:
        uploads_dir = UPLOAD_FOLDER
        exists = os.path.exists(uploads_dir)
        files = []
        if exists:
            files = os.listdir(uploads_dir)
        
        return jsonify({
            'uploads_dir': uploads_dir,
            'exists': exists,
            'files': files,
            'file_count': len(files)
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500



def post_to_facebook_page(page_id, message):
    """
    Posts a message to a Facebook Page using the Graph API.
    Requires FACEBOOK_USER_ACCESS_TOKEN in .env.
    """
    print(f"[DEBUG] post_to_facebook_page called with page_id: {page_id}")
    print(f"[DEBUG] Message length: {len(message)} characters")
    print(f"[DEBUG] Message preview: {message[:100]}...")
    
    user_token = os.getenv('FACEBOOK_USER_ACCESS_TOKEN')
    
    if not user_token:
        print(f"[DEBUG] FACEBOOK_USER_ACCESS_TOKEN not found in environment")
        raise Exception("FACEBOOK_USER_ACCESS_TOKEN not found in environment. Please add your Facebook token to the .env file.")
    
    print(f"[DEBUG] User token found, length: {len(user_token)} characters")
    print(f"[DEBUG] Token starts with: {user_token[:20]}...")
    print(f"[DEBUG] Token ends with: ...{user_token[-20:]}")

    # Step 1: Get Page Access Token
    accounts_url = f"https://graph.facebook.com/v23.0/me/accounts"
    params = {'access_token': user_token}
    print(f"[DEBUG] Making request to: {accounts_url}")
    
    resp = requests.get(accounts_url, params=params)
    print(f"[DEBUG] Accounts API response status: {resp.status_code}")
    
    data = resp.json()
    print(f"[DEBUG] Accounts API response: {data}")
    
    if 'data' not in data:
        print(f"[DEBUG] Error: 'data' key not found in response")
        raise Exception(f"Error fetching pages: {data}")
    
    print(f"[DEBUG] Found {len(data['data'])} pages")
    
    page_token = None
    for page in data['data']:
        print(f"[DEBUG] Checking page: {page.get('name', 'Unknown')} (ID: {page.get('id', 'Unknown')})")
        if page['id'] == page_id:
            page_token = page['access_token']
            print(f"[DEBUG] Found matching page, token length: {len(page_token)} characters")
            break
    
    if not page_token:
        print(f"[DEBUG] Page access token not found for page_id: {page_id}")
        raise Exception("Page access token not found for the given page_id.")

    # Step 2: Post to Page
    post_url = f"https://graph.facebook.com/v23.0/{page_id}/feed"
    post_params = {
        'message': message,
        'access_token': page_token,
        'published': True  # Ensure post is published and visible
    }
    print(f"[DEBUG] Making post request to: {post_url}")
    print(f"[DEBUG] Post params keys: {list(post_params.keys())}")
    
    post_resp = requests.post(post_url, data=post_params)
    print(f"[DEBUG] Post API response status: {post_resp.status_code}")
    
    result = post_resp.json()
    print(f"[DEBUG] Post API response: {result}")
    
    return result

def post_to_facebook_page_with_image(page_id, message, image_url):
    """
    Posts a message with an image to a Facebook Page using the Graph API.
    """
    print(f"[DEBUG] post_to_facebook_page_with_image called")
    print(f"[DEBUG] page_id: {page_id}, message length: {len(message)}, image_url: {image_url}")
    
    user_token = os.getenv('FACEBOOK_USER_ACCESS_TOKEN')
    
    if not user_token:
        raise Exception("FACEBOOK_USER_ACCESS_TOKEN not found in environment. Please add your Facebook token to the .env file.")

    # Step 1: Get Page Access Token
    accounts_url = f"https://graph.facebook.com/v23.0/me/accounts"
    params = {'access_token': user_token}
    resp = requests.get(accounts_url, params=params)
    data = resp.json()
    
    if 'data' not in data:
        raise Exception(f"Error fetching pages: {data}")
    
    page_token = None
    for page in data['data']:
        if page['id'] == page_id:
            page_token = page['access_token']
            break
    
    if not page_token:
        raise Exception("Page access token not found for the given page_id.")

    # Step 2: Handle image URL conversion
    print(f"[DEBUG] Original image_url: {image_url}")
    print(f"[DEBUG] UPLOAD_FOLDER: {UPLOAD_FOLDER}")
    
    if image_url.startswith('/uploads/'):
        # Facebook can't access localhost URLs, so we need to upload the file directly
        filename = os.path.basename(image_url)
        local_file_path = os.path.join(UPLOAD_FOLDER, filename)
        print(f"[DEBUG] Extracted filename: {filename}")
        print(f"[DEBUG] Local file path: {local_file_path}")
        print(f"[DEBUG] Local file path absolute: {os.path.abspath(local_file_path)}")
        
        # List all files in upload directory for debugging
        if os.path.exists(UPLOAD_FOLDER):
            files_in_upload = os.listdir(UPLOAD_FOLDER)
            print(f"[DEBUG] Files in upload directory: {files_in_upload}")
        else:
            print(f"[DEBUG] Upload directory does not exist: {UPLOAD_FOLDER}")
        
        if os.path.exists(local_file_path):
            print(f"[DEBUG] File exists, will upload directly to Facebook")
            # We'll upload the file directly instead of using URL
            return post_to_facebook_with_file_upload(page_id, message, local_file_path, page_token)
        else:
            print(f"[DEBUG] File not found: {local_file_path}")
            # Try alternative paths
            alt_paths = [
                os.path.join(os.getcwd(), 'uploads', filename),
                os.path.join(os.path.dirname(__file__), 'uploads', filename),
                os.path.join('static', 'social_uploads', filename)
            ]
            for alt_path in alt_paths:
                print(f"[DEBUG] Trying alternative path: {alt_path}")
                if os.path.exists(alt_path):
                    print(f"[DEBUG] Found file at alternative path: {alt_path}")
                    return post_to_facebook_with_file_upload(page_id, message, alt_path, page_token)
            
            print(f"[DEBUG] File not found in any location, falling back to text-only post")
            # Fall back to text-only post
            return post_to_facebook_page(page_id, message)
    else:
        # External URL - convert to absolute if needed
        if not image_url.startswith('http'):
            image_url = f"http://localhost:5001{image_url}"
        print(f"[DEBUG] Using external image URL: {image_url}")

    # Step 3: Try to post with image, fallback to text-only if image fails
    post_url = f"https://graph.facebook.com/v23.0/{page_id}/photos"
    post_params = {
        'message': message,
        'url': image_url,
        'access_token': page_token,
        'published': True  # Make sure it's published
    }
    
    print(f"[DEBUG] Attempting to post photo with URL: {image_url}")
    post_resp = requests.post(post_url, data=post_params)
    result = post_resp.json()
    
    print(f"[DEBUG] Photo post response: {result}")
    
    # If photo post failed, try text-only post
    if 'error' in result:
        print(f"[DEBUG] Photo post failed, trying text-only post: {result['error']}")
        
        # Fallback to regular text post
        text_post_url = f"https://graph.facebook.com/v23.0/{page_id}/feed"
        text_post_params = {
            'message': message + f"\n\n[Image could not be posted: {image_url}]",
            'access_token': page_token,
            'published': True
        }
        
        text_post_resp = requests.post(text_post_url, data=text_post_params)
        result = text_post_resp.json()
        print(f"[DEBUG] Fallback text post response: {result}")
    
    return result

def post_to_facebook_with_file_upload(page_id, message, file_path, page_token):
    """
    Posts a message with an image file directly uploaded to Facebook.
    """
    print(f"[DEBUG] post_to_facebook_with_file_upload called")
    print(f"[DEBUG] file_path: {file_path}")
    print(f"[DEBUG] file_path exists: {os.path.exists(file_path)}")
    
    if os.path.exists(file_path):
        print(f"[DEBUG] file size: {os.path.getsize(file_path)} bytes")
        
        # Check if file is actually an image
        try:
            from PIL import Image
            with Image.open(file_path) as img:
                print(f"[DEBUG] Image format: {img.format}, size: {img.size}, mode: {img.mode}")
        except Exception as img_err:
            print(f"[DEBUG] Not a valid image file: {img_err}")
    
    try:
        # Step 1: Upload photo to Facebook
        post_url = f"https://graph.facebook.com/v23.0/{page_id}/photos"
        print(f"[DEBUG] Facebook API URL: {post_url}")
        print(f"[DEBUG] Page ID: {page_id}")
        print(f"[DEBUG] Message length: {len(message)} chars")
        print(f"[DEBUG] Page token length: {len(page_token)} chars")
        
        with open(file_path, 'rb') as image_file:
            files = {
                'source': image_file
            }
            data = {
                'message': message,
                'access_token': page_token,
                'published': True
            }
            
            print(f"[DEBUG] Request data keys: {list(data.keys())}")
            print(f"[DEBUG] Request files keys: {list(files.keys())}")
            print(f"[DEBUG] Uploading file to Facebook...")
            
            response = requests.post(post_url, files=files, data=data)
            print(f"[DEBUG] Response status code: {response.status_code}")
            print(f"[DEBUG] Response headers: {dict(response.headers)}")
            
            result = response.json()
            print(f"[DEBUG] File upload response: {result}")
            
            # Check for specific error types
            if 'error' in result:
                error_info = result['error']
                print(f"[DEBUG] Facebook API Error Code: {error_info.get('code', 'Unknown')}")
                print(f"[DEBUG] Facebook API Error Type: {error_info.get('type', 'Unknown')}")
                print(f"[DEBUG] Facebook API Error Message: {error_info.get('message', 'Unknown')}")
                print(f"[DEBUG] Facebook API Error Subcode: {error_info.get('error_subcode', 'None')}")
            
            return result
            
    except FileNotFoundError:
        print(f"[DEBUG] File not found: {file_path}")
        return {'error': {'message': f'File not found: {file_path}'}}
    except Exception as e:
        print(f"[DEBUG] File upload exception: {str(e)}")
        print(f"[DEBUG] Exception type: {type(e).__name__}")
        import traceback
        traceback.print_exc()
        # Fallback to text-only post
        return post_to_facebook_page(page_id, message)

from flask import request, jsonify

@app.route('/api/facebook/debug', methods=['GET'])
def api_facebook_debug():
    """Debug endpoint to check Facebook configuration"""
    print(f"[DEBUG] Facebook Debug endpoint called at {time.strftime('%Y-%m-%d %H:%M:%S')}")
    
    try:
        # Check environment variables
        user_token = os.getenv('FACEBOOK_USER_ACCESS_TOKEN')
        page_id = os.getenv('FACEBOOK_PAGE_ID')
        
        debug_info = {
            'env_variables': {
                'FACEBOOK_USER_ACCESS_TOKEN': 'Present' if user_token else 'Missing',
                'FACEBOOK_USER_ACCESS_TOKEN_length': len(user_token) if user_token else 0,
                'FACEBOOK_PAGE_ID': page_id if page_id else 'Missing'
            }
        }
        
        if not user_token:
            return jsonify({'error': 'FACEBOOK_USER_ACCESS_TOKEN not found', 'debug': debug_info}), 400
        
        if not page_id:
            return jsonify({'error': 'FACEBOOK_PAGE_ID not found', 'debug': debug_info}), 400
        
        # Test Facebook API access
        print(f"[DEBUG] Testing Facebook API with token and page ID: {page_id}")
        
        # Test 1: Get user info
        user_url = f"https://graph.facebook.com/v23.0/me"
        user_params = {'access_token': user_token, 'fields': 'id,name'}
        user_resp = requests.get(user_url, params=user_params)
        
        debug_info['user_api_test'] = {
            'status_code': user_resp.status_code,
            'response': user_resp.json() if user_resp.status_code == 200 else user_resp.text
        }
        
        # Test 2: Get pages
        accounts_url = f"https://graph.facebook.com/v23.0/me/accounts"
        accounts_params = {'access_token': user_token}
        accounts_resp = requests.get(accounts_url, params=accounts_params)
        
        debug_info['pages_api_test'] = {
            'status_code': accounts_resp.status_code,
            'response': accounts_resp.json() if accounts_resp.status_code == 200 else accounts_resp.text
        }
        
        # Test 3: Check if our page ID is in the list
        if accounts_resp.status_code == 200:
            accounts_data = accounts_resp.json()
            page_found = False
            page_access_token = None
            
            if 'data' in accounts_data:
                for page in accounts_data['data']:
                    if page.get('id') == page_id:
                        page_found = True
                        page_access_token = page.get('access_token')
                        break
                
                debug_info['page_validation'] = {
                    'page_id_provided': page_id,
                    'page_found_in_accounts': page_found,
                    'total_pages_available': len(accounts_data['data']),
                    'available_pages': [{'id': p.get('id'), 'name': p.get('name')} for p in accounts_data['data']],
                    'page_has_access_token': bool(page_access_token)
                }
        
        return jsonify({
            'success': True,
            'message': 'Facebook configuration debug completed',
            'debug': debug_info
        })
        
    except Exception as e:
        print(f"[DEBUG] Exception in Facebook debug endpoint: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'Debug error: {str(e)}'}), 500

@app.route('/api/facebook/pages', methods=['GET'])
def api_facebook_pages():
    """Debug endpoint to list all pages accessible with current token"""
    try:
        user_token = os.getenv('FACEBOOK_USER_ACCESS_TOKEN')
        
        if not user_token:
            raise Exception("FACEBOOK_USER_ACCESS_TOKEN not found in environment. Please add your Facebook token to the .env file.")
        
        accounts_url = f"https://graph.facebook.com/v23.0/me/accounts"
        params = {'access_token': user_token}
        resp = requests.get(accounts_url, params=params)
        data = resp.json()
        
        print(f"[DEBUG] Pages API response: {data}")
        return jsonify(data)
    except Exception as e:
        print(f"[DEBUG] Error fetching pages: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/facebook/post', methods=['POST'])
def api_facebook_post():
    print(f"[DEBUG] Facebook API endpoint called at {time.strftime('%Y-%m-%d %H:%M:%S')}")
    print(f"[DEBUG] Request method: {request.method}")
    print(f"[DEBUG] Request headers: {dict(request.headers)}")
    
    try:
        data = request.json
        print(f"[DEBUG] Request JSON data: {data}")
    except Exception as e:
        print(f"[DEBUG] Error parsing JSON: {str(e)}")
        return jsonify({'error': 'Invalid JSON data'}), 400
    
    page_id = data.get('page_id')
    message = data.get('message')
    image_url = data.get('image_url')  # Optional image URL
    
    # Get page ID from environment if not provided
    if not page_id:
        page_id = os.getenv('FACEBOOK_PAGE_ID')
    
    print(f"[DEBUG] Extracted page_id: {page_id}")
    print(f"[DEBUG] Extracted message length: {len(message) if message else 0}")
    print(f"[DEBUG] Message preview: {message[:100] if message else 'None'}...")
    print(f"[DEBUG] Image URL: {image_url}")
    
    if not page_id or not message:
        print(f"[DEBUG] Missing required fields - page_id: {bool(page_id)}, message: {bool(message)}")
        return jsonify({'error': 'Missing page_id or message. Please add FACEBOOK_PAGE_ID to your .env file or provide page_id in request.'}), 400
    
    try:
        print(f"[DEBUG] Calling post_to_facebook_page function...")
        if image_url:
            result = post_to_facebook_page_with_image(page_id, message, image_url)
        else:
            result = post_to_facebook_page(page_id, message)
        print(f"[DEBUG] Facebook API response: {result}")
        return jsonify(result)
    except Exception as e:
        print(f"[DEBUG] Exception in post_to_facebook_page: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

def post_to_instagram_business(user_id, caption, image_url):
    """
    Posts content to Instagram Business Account using Facebook Graph API.
    Instagram posting requires a 2-step process:
    1. Create media container
    2. Publish the container
    """
    print(f"[DEBUG] post_to_instagram_business called")
    print(f"[DEBUG] user_id: {user_id}")
    print(f"[DEBUG] caption length: {len(caption)}")
    print(f"[DEBUG] image_url: {image_url}")
    
    try:
        # Get Instagram access token from environment
        access_token = os.getenv('INSTAGRAM_ACCESS_TOKEN')
        if not access_token:
            print("[ERROR] INSTAGRAM_ACCESS_TOKEN not found in environment")
            return {
                'success': False,
                'error': 'Instagram access token not configured. Please add INSTAGRAM_ACCESS_TOKEN to your .env file.',
                'details': 'Instagram posting requires a valid access token from Instagram Business API.'
            }
        
        print(f"[DEBUG] Access token length: {len(access_token)} chars")
        
        # Step 1: Create Media Container
        print("[DEBUG] Creating Instagram media container...")
        container_url = f"https://graph.facebook.com/v18.0/{user_id}/media"
        container_params = {
            'image_url': image_url,
            'caption': caption,
            'access_token': access_token
        }
        
        print(f"[DEBUG] Container creation URL: {container_url}")
        print(f"[DEBUG] Container params: {container_params}")
        
        container_resp = requests.post(container_url, data=container_params)
        container_result = container_resp.json()
        
        print(f"[DEBUG] Container creation response: {container_result}")
        
        if 'error' in container_result:
            error_info = container_result['error']
            print(f"[ERROR] Container creation failed: {error_info}")
            return {
                'success': False,
                'error': f"Failed to create Instagram media container: {error_info.get('message', 'Unknown error')}",
                'details': error_info
            }
        
        # Get the container ID
        container_id = container_result.get('id')
        if not container_id:
            print("[ERROR] No container ID in response")
            return {
                'success': False,
                'error': 'Failed to get container ID from Instagram API',
                'details': container_result
            }
        
        print(f"[DEBUG] Container created with ID: {container_id}")
        
        # Step 2: Publish the Container
        print("[DEBUG] Publishing Instagram media container...")
        publish_url = f"https://graph.facebook.com/v18.0/{user_id}/media_publish"
        publish_params = {
            'creation_id': container_id,
            'access_token': access_token
        }
        
        print(f"[DEBUG] Publish URL: {publish_url}")
        print(f"[DEBUG] Publish params: {publish_params}")
        
        publish_resp = requests.post(publish_url, data=publish_params)
        publish_result = publish_resp.json()
        
        print(f"[DEBUG] Publish response: {publish_result}")
        
        if 'error' in publish_result:
            error_info = publish_result['error']
            print(f"[ERROR] Publishing failed: {error_info}")
            return {
                'success': False,
                'error': f"Failed to publish Instagram post: {error_info.get('message', 'Unknown error')}",
                'details': error_info
            }
        
        # Success!
        post_id = publish_result.get('id')
        print(f"[DEBUG] Instagram post published successfully with ID: {post_id}")
        
        return {
            'success': True,
            'post_id': post_id,
            'container_id': container_id,
            'message': 'Successfully posted to Instagram'
        }
        
    except Exception as e:
        print(f"[ERROR] Exception in post_to_instagram_business: {str(e)}")
        import traceback
        traceback.print_exc()
        return {
            'success': False,
            'error': f'Instagram posting failed: {str(e)}',
            'details': str(e)
        }

@app.route('/api/instagram/post', methods=['POST'])
def api_instagram_post():
    """
    API endpoint to post content to Instagram Business Account
    """
    print(f"[DEBUG] Instagram API endpoint called at {time.strftime('%Y-%m-%d %H:%M:%S')}")
    print(f"[DEBUG] Request method: {request.method}")
    print(f"[DEBUG] Request headers: {dict(request.headers)}")
    
    try:
        data = request.json
        print(f"[DEBUG] Request JSON data: {data}")
    except Exception as e:
        print(f"[DEBUG] Error parsing JSON: {str(e)}")
        return jsonify({'error': 'Invalid JSON data'}), 400
    
    user_id = data.get('user_id')
    caption = data.get('caption', '')
    image_url = data.get('image_url')
    
    # Get user ID from environment if not provided
    if not user_id:
        user_id = os.getenv('INSTAGRAM_USER_ID')
    
    print(f"[DEBUG] Extracted user_id: {user_id}")
    print(f"[DEBUG] Extracted caption length: {len(caption) if caption else 0}")
    print(f"[DEBUG] Caption preview: {caption[:100] if caption else 'None'}...")
    print(f"[DEBUG] Image URL: {image_url}")
    
    if not user_id:
        print(f"[DEBUG] Missing user_id")
        return jsonify({'error': 'Instagram User ID is required. Please add INSTAGRAM_USER_ID to your .env file or provide user_id in request.'}), 400
    
    if not image_url:
        print(f"[DEBUG] Missing image_url")
        return jsonify({'error': 'Image URL is required for Instagram posts'}), 400
    
    try:
        print(f"[DEBUG] Calling post_to_instagram_business function...")
        result = post_to_instagram_business(user_id, caption, image_url)
        print(f"[DEBUG] Instagram API response: {result}")
        
        if result.get('success'):
            return jsonify({
                'success': True,
                'id': result.get('post_id', 'unknown'),
                'message': 'Posted successfully to Instagram'
            })
        else:
            return jsonify({
                'error': result.get('error', 'Unknown error occurred'),
                'details': result.get('details', '')
            }), 400
            
    except Exception as e:
        print(f"[DEBUG] Exception in Instagram API endpoint: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'Server error: {str(e)}'}), 500

def post_to_linkedin(user_id, text, image_url=None, access_token=None):
    """
    Posts content to LinkedIn using LinkedIn API v2
    """
    print(f"[DEBUG] post_to_linkedin called")
    print(f"[DEBUG] user_id: {user_id}")
    print(f"[DEBUG] text length: {len(text)}")
    print(f"[DEBUG] image_url: {image_url}")
    
    try:
        # Get LinkedIn access token from parameter or environment
        if not access_token:
            access_token = os.getenv('LINKEDIN_ACCESS_TOKEN')
        
        if not access_token:
            print("[ERROR] LINKEDIN_ACCESS_TOKEN not found in environment or parameter")
            return {
                'success': False,
                'error': 'LinkedIn access token not configured. Please add LINKEDIN_ACCESS_TOKEN to your .env file or pass as parameter.',
                'details': 'LinkedIn posting requires a valid access token from LinkedIn API.'
            }
        
        print(f"[DEBUG] Access token length: {len(access_token)} chars")
        print(f"[DEBUG] Access token preview: {access_token[:20]}...")
        
        # LinkedIn Posts API endpoint (current API)
        url = 'https://api.linkedin.com/rest/posts'
        
        headers = {
            'Authorization': f'Bearer {access_token}',
            'Content-Type': 'application/json',
            'X-Restli-Protocol-Version': '2.0.0',
            'LinkedIn-Version': '202506'  # Updated to current active version
        }
        
        print(f"[DEBUG] LinkedIn API URL: {url}")
        print(f"[DEBUG] Headers: {headers}")
        
        # LinkedIn Posts API format
        post_data = {
            "author": f"urn:li:person:{user_id}",
            "commentary": text,
            "visibility": "PUBLIC",
            "distribution": {
                "feedDistribution": "MAIN_FEED",
                "targetEntities": [],
                "thirdPartyDistributionChannels": []
            },
            "lifecycleState": "PUBLISHED",
            "isReshareDisabledByAuthor": False
        }
        
        # Handle image posting if image URL is provided
        if image_url:
            print(f"[DEBUG] Processing image for LinkedIn post: {image_url}")
            print(f"[DEBUG] Original image_url: {image_url}")
            print(f"[DEBUG] UPLOAD_FOLDER: {UPLOAD_FOLDER}")
            
            # Handle image URL conversion (similar to Facebook implementation)
            if image_url.startswith('/uploads/'):
                # Local file - use direct file upload (similar to Facebook approach)
                filename = os.path.basename(image_url)
                local_file_path = os.path.join(UPLOAD_FOLDER, filename)
                print(f"[DEBUG] Detected local image file: {filename}")
                print(f"[DEBUG] Local file path: {local_file_path}")
                print(f"[DEBUG] Local file path absolute: {os.path.abspath(local_file_path)}")
                
                # List all files in upload directory for debugging
                if os.path.exists(UPLOAD_FOLDER):
                    files_in_upload = os.listdir(UPLOAD_FOLDER)
                    print(f"[DEBUG] Files in upload directory: {files_in_upload}")
                else:
                    print(f"[DEBUG] Upload directory does not exist: {UPLOAD_FOLDER}")
                
                if os.path.exists(local_file_path):
                    print(f"[DEBUG] File exists, will upload directly to LinkedIn")
                    # Use direct file upload instead of URL-based upload
                    return post_to_linkedin_with_file_upload(user_id, text, local_file_path, access_token)
                else:
                    print(f"[DEBUG] File not found: {local_file_path}")
                    # Try alternative paths (similar to Facebook implementation)
                    alt_paths = [
                        os.path.join(os.getcwd(), 'uploads', filename),
                        os.path.join(os.path.dirname(__file__), 'uploads', filename),
                        os.path.join('static', 'social_uploads', filename)
                    ]
                    for alt_path in alt_paths:
                        print(f"[DEBUG] Trying alternative path: {alt_path}")
                        if os.path.exists(alt_path):
                            print(f"[DEBUG] Found file at alternative path: {alt_path}")
                            return post_to_linkedin_with_file_upload(user_id, text, alt_path, access_token)
                    
                    print(f"[DEBUG] File not found in any location, falling back to text-only post")
                    # Fall back to text-only post
                    pass  # Continue with text-only post below
            else:
                # External URL - use URL-based upload
                processed_image_url = image_url
                if not image_url.startswith('http'):
                    processed_image_url = f"http://localhost:5001{image_url}"
                    print(f"[DEBUG] Converted to absolute URL: {processed_image_url}")
                
                try:
                    print(f"[DEBUG] Attempting LinkedIn image upload with URL: {processed_image_url}")
                    
                    # Step 1: Upload image to LinkedIn Images API to get image URN
                    image_urn = upload_image_to_linkedin(processed_image_url, user_id, access_token)
                    
                    if image_urn:
                        print(f"[DEBUG] Successfully uploaded image to LinkedIn, URN: {image_urn}")
                        
                        # Step 2: Add image content to post
                        post_data["content"] = {
                            "media": {
                                "id": image_urn,
                                "altText": "Image shared via blog generation"
                            }
                        }
                        print(f"[DEBUG] Added image content to LinkedIn post data")
                    else:
                        print(f"[WARNING] Failed to upload image to LinkedIn, posting text-only")
                        print(f"[WARNING] Original image URL: {image_url}")
                        print(f"[WARNING] Processed image URL: {processed_image_url}")
                        
                except Exception as img_error:
                    print(f"[ERROR] Image upload exception: {str(img_error)}")
                    print(f"[ERROR] Exception type: {type(img_error).__name__}")
                    import traceback
                    traceback.print_exc()
                    print(f"[WARNING] Falling back to text-only post due to image upload error")
        
        print(f"[DEBUG] LinkedIn post data: {post_data}")
        
        response = requests.post(url, headers=headers, json=post_data)
        
        print(f"[DEBUG] LinkedIn API response status: {response.status_code}")
        print(f"[DEBUG] LinkedIn API response headers: {dict(response.headers)}")
        
        if response.status_code == 201:
            # Success - get post ID from response header
            post_id = response.headers.get('x-restli-id', 'unknown')
            print(f"[DEBUG] LinkedIn post created successfully with ID: {post_id}")
            return {
                'success': True,
                'post_id': post_id,
                'message': 'Posted successfully to LinkedIn'
            }
        else:
            try:
                result = response.json()
            except:
                result = {'error': 'Failed to parse response', 'text': response.text}
            
            error_message = result.get('message', f'HTTP {response.status_code}: {response.text}')
            print(f"[ERROR] LinkedIn posting failed: {error_message}")
            return {
                'success': False,
                'error': f"LinkedIn posting failed: {error_message}",
                'details': result
            }
            
    except Exception as e:
        print(f"[ERROR] Exception in post_to_linkedin: {str(e)}")
        import traceback
        traceback.print_exc()
        return {
            'success': False,
            'error': f'LinkedIn posting failed: {str(e)}',
            'details': str(e)
        }

def post_to_linkedin_with_file_upload(user_id, text, file_path, access_token):
    """
    Posts a message with an image file directly uploaded to LinkedIn.
    Similar to Facebook's implementation but for LinkedIn API.
    """
    print(f"[DEBUG] post_to_linkedin_with_file_upload called")
    print(f"[DEBUG] file_path: {file_path}")
    print(f"[DEBUG] file_path exists: {os.path.exists(file_path)}")
    print(f"[DEBUG] user_id: {user_id}")
    
    if os.path.exists(file_path):
        print(f"[DEBUG] file size: {os.path.getsize(file_path)} bytes")
        
        # Check if file is actually an image
        try:
            from PIL import Image
            with Image.open(file_path) as img:
                print(f"[DEBUG] Image format: {img.format}, size: {img.size}, mode: {img.mode}")
        except Exception as img_err:
            print(f"[DEBUG] Not a valid image file: {img_err}")
            # Fall back to text-only post
            return post_to_linkedin(user_id, text, None, access_token)
    
    try:
        # Read the image file
        with open(file_path, 'rb') as image_file:
            image_data = image_file.read()
        
        print(f"[DEBUG] Successfully read file, size: {len(image_data)} bytes")
        
        # Upload image using the existing upload function by creating a pseudo URL
        # We'll modify upload_image_to_linkedin to handle direct binary data
        image_urn = upload_image_to_linkedin_direct(image_data, user_id, access_token)
        
        if image_urn:
            print(f"[DEBUG] Image uploaded successfully, creating post with image")
            
            # Create LinkedIn post with image
            url = 'https://api.linkedin.com/rest/posts'
            
            headers = {
                'Authorization': f'Bearer {access_token}',
                'Content-Type': 'application/json',
                'X-Restli-Protocol-Version': '2.0.0',
                'LinkedIn-Version': '202506'
            }
            
            post_data = {
                "author": f"urn:li:person:{user_id}",
                "commentary": text,
                "visibility": "PUBLIC",
                "distribution": {
                    "feedDistribution": "MAIN_FEED",
                    "targetEntities": [],
                    "thirdPartyDistributionChannels": []
                },
                "lifecycleState": "PUBLISHED",
                "isReshareDisabledByAuthor": False,
                "content": {
                    "media": {
                        "id": image_urn,
                        "altText": "Image shared via blog generation"
                    }
                }
            }
            
            response = requests.post(url, headers=headers, json=post_data)
            
            if response.status_code == 201:
                post_id = response.headers.get('x-restli-id', 'unknown')
                print(f"[DEBUG] LinkedIn post with image created successfully: {post_id}")
                return {
                    'success': True,
                    'post_id': post_id,
                    'message': 'Posted successfully to LinkedIn with image'
                }
            else:
                try:
                    result = response.json()
                except:
                    result = {'error': 'Failed to parse response', 'text': response.text}
                
                print(f"[ERROR] LinkedIn post creation failed: {result}")
                return {
                    'success': False,
                    'error': f"LinkedIn posting failed: {result}",
                    'details': result
                }
        else:
            print(f"[WARNING] Image upload failed, falling back to text-only post")
            return post_to_linkedin(user_id, text, None, access_token)
            
    except FileNotFoundError:
        print(f"[DEBUG] File not found: {file_path}")
        return {'success': False, 'error': f'File not found: {file_path}'}
    except Exception as e:
        print(f"[DEBUG] File upload exception: {str(e)}")
        print(f"[DEBUG] Exception type: {type(e).__name__}")
        import traceback
        traceback.print_exc()
        # Fallback to text-only post
        return post_to_linkedin(user_id, text, None, access_token)

def upload_image_to_linkedin_direct(image_data, user_id, access_token):
    """
    Upload image binary data directly to LinkedIn Images API and return the image URN
    """
    print(f"[DEBUG] upload_image_to_linkedin_direct called")
    print(f"[DEBUG] Image data size: {len(image_data)} bytes")
    print(f"[DEBUG] user_id: {user_id}")
    print(f"[DEBUG] access_token length: {len(access_token) if access_token else 0}")
    
    try:
        # Step 1: Initialize upload with LinkedIn Images API
        init_url = 'https://api.linkedin.com/rest/images?action=initializeUpload'
        
        init_headers = {
            'Authorization': f'Bearer {access_token}',
            'Content-Type': 'application/json',
            'X-Restli-Protocol-Version': '2.0.0',
            'LinkedIn-Version': '202506'
        }
        
        init_data = {
            "initializeUploadRequest": {
                "owner": f"urn:li:person:{user_id}"
            }
        }
        
        print(f"[DEBUG] Initializing direct image upload to LinkedIn...")
        init_response = requests.post(init_url, headers=init_headers, json=init_data)
        
        print(f"[DEBUG] Initialize response status: {init_response.status_code}")
        
        if init_response.status_code != 200:
            print(f"[ERROR] Failed to initialize direct image upload: {init_response.status_code} - {init_response.text}")
            return None
        
        init_result = init_response.json()
        upload_url = init_result['value']['uploadUrl']
        image_urn = init_result['value']['image']
        
        print(f"[DEBUG] Direct upload URL obtained: {upload_url[:50]}...")
        print(f"[DEBUG] Direct image URN: {image_urn}")
        
        # Step 2: Upload the image binary data
        upload_headers = {
            'Authorization': f'Bearer {access_token}',
            'Content-Type': 'application/octet-stream'
        }
        
        upload_response = requests.put(upload_url, headers=upload_headers, data=image_data)
        
        if upload_response.status_code in [200, 201]:
            print(f"[DEBUG] Direct image upload successful")
            return image_urn
        else:
            print(f"[ERROR] Failed direct image upload: {upload_response.status_code} - {upload_response.text}")
            return None
            
    except Exception as e:
        print(f"[ERROR] Exception in upload_image_to_linkedin_direct: {str(e)}")
        import traceback
        traceback.print_exc()
        return None

def upload_image_to_linkedin(image_url, user_id, access_token):
    """
    Upload an image to LinkedIn Images API and return the image URN
    """
    print(f"[DEBUG] upload_image_to_linkedin called")
    print(f"[DEBUG] Original image_url: {image_url}")
    print(f"[DEBUG] user_id: {user_id}")
    print(f"[DEBUG] access_token length: {len(access_token) if access_token else 0}")
    print(f"[DEBUG] UPLOAD_FOLDER: {UPLOAD_FOLDER}")
    
    try:
        image_data = None
        
        # Handle local file paths (similar to Facebook implementation)
        if image_url.startswith('/uploads/'):
            # Local file - read directly from file system
            filename = os.path.basename(image_url)
            local_file_path = os.path.join(UPLOAD_FOLDER, filename)
            print(f"[DEBUG] Detected local file upload")
            print(f"[DEBUG] Extracted filename: {filename}")
            print(f"[DEBUG] Local file path: {local_file_path}")
            print(f"[DEBUG] Local file path absolute: {os.path.abspath(local_file_path)}")
            
            # List all files in upload directory for debugging
            if os.path.exists(UPLOAD_FOLDER):
                files_in_upload = os.listdir(UPLOAD_FOLDER)
                print(f"[DEBUG] Files in upload directory: {files_in_upload}")
            else:
                print(f"[DEBUG] Upload directory does not exist: {UPLOAD_FOLDER}")
            
            if os.path.exists(local_file_path):
                print(f"[DEBUG] File exists, reading directly from filesystem")
                print(f"[DEBUG] File size: {os.path.getsize(local_file_path)} bytes")
                
                # Check if file is actually an image
                try:
                    from PIL import Image
                    with Image.open(local_file_path) as img:
                        print(f"[DEBUG] Image format: {img.format}, size: {img.size}, mode: {img.mode}")
                except Exception as img_err:
                    print(f"[DEBUG] Not a valid image file: {img_err}")
                    return None
                
                with open(local_file_path, 'rb') as f:
                    image_data = f.read()
                print(f"[DEBUG] Successfully read local file, size: {len(image_data)} bytes")
            else:
                print(f"[DEBUG] File not found: {local_file_path}")
                # Try alternative paths (similar to Facebook implementation)
                alt_paths = [
                    os.path.join(os.getcwd(), 'uploads', filename),
                    os.path.join(os.path.dirname(__file__), 'uploads', filename),
                    os.path.join('static', 'social_uploads', filename)
                ]
                for alt_path in alt_paths:
                    print(f"[DEBUG] Trying alternative path: {alt_path}")
                    if os.path.exists(alt_path):
                        print(f"[DEBUG] Found file at alternative path: {alt_path}")
                        with open(alt_path, 'rb') as f:
                            image_data = f.read()
                        print(f"[DEBUG] Successfully read from alternative path, size: {len(image_data)} bytes")
                        break
                
                if not image_data:
                    print(f"[DEBUG] File not found in any location")
                    return None
        else:
            # External URL - download the image
            print(f"[DEBUG] Detected external URL, downloading...")
            # Convert to absolute URL if needed
            if not image_url.startswith('http'):
                image_url = f"http://localhost:5001{image_url}"
                print(f"[DEBUG] Converted to absolute URL: {image_url}")
            
            img_response = requests.get(image_url, timeout=30)
            print(f"[DEBUG] Download response status: {img_response.status_code}")
            
            if img_response.status_code != 200:
                print(f"[ERROR] Failed to download image: {img_response.status_code}")
                print(f"[ERROR] Response text: {img_response.text}")
                return None
            
            image_data = img_response.content
            print(f"[DEBUG] Downloaded image from URL, size: {len(image_data)} bytes")
        
        if not image_data:
            print(f"[ERROR] No image data obtained")
            return None
        
        # Step 2: Initialize upload with LinkedIn Images API
        init_url = 'https://api.linkedin.com/rest/images?action=initializeUpload'
        
        init_headers = {
            'Authorization': f'Bearer {access_token}',
            'Content-Type': 'application/json',
            'X-Restli-Protocol-Version': '2.0.0',
            'LinkedIn-Version': '202506'  # Updated to current active version
        }
        
        init_data = {
            "initializeUploadRequest": {
                "owner": f"urn:li:person:{user_id}"
            }
        }
        
        print(f"[DEBUG] Step 2: Initializing image upload to LinkedIn...")
        print(f"[DEBUG] Init URL: {init_url}")
        print(f"[DEBUG] Init headers: {init_headers}")
        print(f"[DEBUG] Init data: {init_data}")
        
        init_response = requests.post(init_url, headers=init_headers, json=init_data)
        
        print(f"[DEBUG] Initialize response status: {init_response.status_code}")
        print(f"[DEBUG] Initialize response headers: {dict(init_response.headers)}")
        
        if init_response.status_code != 200:
            print(f"[ERROR] Failed to initialize image upload")
            print(f"[ERROR] Status code: {init_response.status_code}")
            print(f"[ERROR] Response text: {init_response.text}")
            try:
                error_json = init_response.json()
                print(f"[ERROR] Response JSON: {error_json}")
            except:
                print(f"[ERROR] Could not parse response as JSON")
            return None
        
        try:
            init_result = init_response.json()
            print(f"[DEBUG] Initialize response JSON: {init_result}")
        except Exception as json_err:
            print(f"[ERROR] Failed to parse initialize response JSON: {json_err}")
            return None
        
        if 'value' not in init_result:
            print(f"[ERROR] 'value' key not found in initialize response")
            return None
            
        upload_url = init_result['value'].get('uploadUrl')
        image_urn = init_result['value'].get('image')
        
        if not upload_url or not image_urn:
            print(f"[ERROR] Missing uploadUrl or image URN from initialize response")
            print(f"[ERROR] uploadUrl: {upload_url}")
            print(f"[ERROR] image_urn: {image_urn}")
            return None
        
        print(f"[DEBUG] Upload URL obtained: {upload_url[:50]}...")
        print(f"[DEBUG] Image URN: {image_urn}")
        
        # Step 3: Upload the image binary data
        upload_headers = {
            'Authorization': f'Bearer {access_token}',
            'Content-Type': 'application/octet-stream'
        }
        
        print(f"[DEBUG] Step 3: Uploading image binary data...")
        print(f"[DEBUG] Upload URL: {upload_url}")
        print(f"[DEBUG] Upload headers: {upload_headers}")
        print(f"[DEBUG] Image data size: {len(image_data)} bytes")
        
        upload_response = requests.put(upload_url, headers=upload_headers, data=image_data)
        
        print(f"[DEBUG] Upload response status: {upload_response.status_code}")
        print(f"[DEBUG] Upload response headers: {dict(upload_response.headers)}")
        
        if upload_response.status_code in [200, 201]:
            print(f"[DEBUG] Image uploaded successfully to LinkedIn")
            print(f"[DEBUG] Final image URN: {image_urn}")
            return image_urn
        else:
            print(f"[ERROR] Failed to upload image binary to LinkedIn")
            print(f"[ERROR] Status code: {upload_response.status_code}")
            print(f"[ERROR] Response text: {upload_response.text}")
            try:
                error_json = upload_response.json()
                print(f"[ERROR] Response JSON: {error_json}")
            except:
                print(f"[ERROR] Could not parse upload response as JSON")
            return None
            
    except Exception as e:
        print(f"[ERROR] Exception in upload_image_to_linkedin: {str(e)}")
        import traceback
        traceback.print_exc()
        return None

@app.route('/api/linkedin/post', methods=['POST'])
def api_linkedin_post():
    """
    API endpoint to post content to LinkedIn
    """
    print(f"[DEBUG] LinkedIn API endpoint called at {time.strftime('%Y-%m-%d %H:%M:%S')}")
    print(f"[DEBUG] Request method: {request.method}")
    print(f"[DEBUG] Request headers: {dict(request.headers)}")
    
    try:
        data = request.json
        print(f"[DEBUG] Request JSON data: {data}")
    except Exception as e:
        print(f"[DEBUG] Error parsing JSON: {str(e)}")
        return jsonify({'error': 'Invalid JSON data'}), 400
    
    user_id = data.get('user_id')
    text = data.get('text', '')
    image_url = data.get('image_url')
    access_token = data.get('access_token')  # Get access token from frontend
    
    # Get user ID from environment if not provided
    if not user_id:
        user_id = os.getenv('LINKEDIN_USER_ID')
    
    print(f"[DEBUG] Extracted user_id: {user_id}")
    print(f"[DEBUG] Extracted text length: {len(text) if text else 0}")
    print(f"[DEBUG] Text preview: {text[:100] if text else 'None'}...")
    print(f"[DEBUG] Image URL: {image_url}")
    print(f"[DEBUG] Access token provided: {bool(access_token)}")
    
    if not user_id:
        print(f"[DEBUG] Missing user_id")
        return jsonify({'error': 'LinkedIn User ID is required. Please add LINKEDIN_USER_ID to your .env file or provide user_id in request.'}), 400
    
    if not text.strip():
        print(f"[DEBUG] Missing text content")
        return jsonify({'error': 'Text content is required for LinkedIn posts'}), 400
    
    try:
        print(f"[DEBUG] Calling post_to_linkedin function...")
        result = post_to_linkedin(user_id, text, image_url, access_token)
        print(f"[DEBUG] LinkedIn API response: {result}")
        
        if result.get('success'):
            return jsonify({
                'success': True,
                'id': result.get('post_id', 'unknown'),
                'message': 'Posted successfully to LinkedIn'
            })
        else:
            return jsonify({
                'error': result.get('error', 'Unknown error occurred'),
                'details': result.get('details', '')
            }), 400
            
    except Exception as e:
        print(f"[DEBUG] Exception in LinkedIn API endpoint: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'Server error: {str(e)}'}), 500

# Add LinkedIn OAuth callback route
@app.route('/auth/linkedin/callback')
def linkedin_callback():
    """
    Handle LinkedIn OAuth callback
    """
    # Get authorization code from LinkedIn
    code = request.args.get('code')
    error = request.args.get('error')
    
    if error:
        return f"LinkedIn OAuth Error: {error}", 400
    
    if not code:
        return "No authorization code received", 400
    
    # Exchange code for access token
    token_url = 'https://www.linkedin.com/oauth/v2/accessToken'
    token_data = {
        'grant_type': 'authorization_code',
        'code': code,
        'client_id': '86kzeueky3660a',
        'client_secret': os.getenv('LINKEDIN_CLIENT_SECRET'),  # You'll need to add this to your .env
        'redirect_uri': 'http://localhost:5001/auth/linkedin/callback'
    }
    
    try:
        token_response = requests.post(token_url, data=token_data)
        token_result = token_response.json()
        
        if 'access_token' in token_result:
            access_token = token_result['access_token']
            
            # Try multiple LinkedIn API endpoints to get user profile
            user_id = None
            first_name = ''
            last_name = ''
            email = ''
            
            # Try OpenID Connect userinfo endpoint first
            try:
                profile_url = 'https://api.linkedin.com/v2/userinfo'
                headers = {'Authorization': f'Bearer {access_token}'}
                profile_response = requests.get(profile_url, headers=headers)
                profile_data = profile_response.json()
                
                if profile_response.status_code == 200:
                    user_id = profile_data.get('sub')
                    first_name = profile_data.get('given_name', '')
                    last_name = profile_data.get('family_name', '')
                    email = profile_data.get('email', '')
                    print(f"OpenID Connect Profile API Response: {profile_data}")
                else:
                    print(f"OpenID Connect failed: {profile_data}")
                    raise Exception("OpenID Connect failed")
            except:
                # Fallback to legacy LinkedIn API v2 people endpoint
                try:
                    profile_url = 'https://api.linkedin.com/v2/people/~?projection=(id,firstName,lastName)'
                    headers = {'Authorization': f'Bearer {access_token}'}
                    profile_response = requests.get(profile_url, headers=headers)
                    profile_data = profile_response.json()
                    
                    print(f"Legacy Profile API Response: {profile_data}")
                    
                    if profile_response.status_code == 200:
                        user_id = profile_data.get('id')
                        first_name = profile_data.get('firstName', {}).get('localized', {}).get('en_US', '')
                        last_name = profile_data.get('lastName', {}).get('localized', {}).get('en_US', '')
                        
                        # Try to get email separately
                        try:
                            email_url = 'https://api.linkedin.com/v2/emailAddress?q=members&projection=(elements*(handle~))'
                            email_response = requests.get(email_url, headers=headers)
                            email_data = email_response.json()
                            print(f"Email API Response: {email_data}")
                            
                            if email_response.status_code == 200:
                                elements = email_data.get('elements', [])
                                if elements:
                                    email = elements[0].get('handle~', {}).get('emailAddress', '')
                        except:
                            print("Email retrieval failed")
                            pass
                    else:
                        print(f"Legacy API also failed: {profile_data}")
                except Exception as e:
                    print(f"All profile retrieval attempts failed: {str(e)}")
                    # Continue with empty values - at least we have the access token
            
            # Last resort: Try to get user ID from UGC API (works with w_member_social scope)
            if not user_id:
                try:
                    # This endpoint sometimes works with just w_member_social scope
                    ugc_url = 'https://api.linkedin.com/v2/people/~:(id)'
                    headers = {'Authorization': f'Bearer {access_token}'}
                    ugc_response = requests.get(ugc_url, headers=headers)
                    ugc_data = ugc_response.json()
                    
                    print(f"UGC API Response: {ugc_data}")
                    
                    if ugc_response.status_code == 200:
                        user_id = ugc_data.get('id')
                        print(f"Successfully got User ID from UGC API: {user_id}")
                except Exception as e:
                    print(f"UGC API also failed: {str(e)}")
                    # If all else fails, we can still use the access token without user ID
                    pass
            
            return f"""
            <h2>LinkedIn Authentication Successful!</h2>
            <p><strong>Access Token:</strong> {access_token[:20]}...</p>
            <p><strong>User ID:</strong> {user_id}</p>
            <p><strong>Name:</strong> {first_name} {last_name}</p>
            <p><strong>Email:</strong> {email}</p>
            <p>You can now close this window and use LinkedIn posting features.</p>
            <script>
                // Store tokens in localStorage for use by the main app
                localStorage.setItem('linkedin_access_token', '{access_token}');
                localStorage.setItem('linkedin_user_id', '{user_id}');
                localStorage.setItem('linkedin_user_email', '{email}');
                alert('LinkedIn authentication successful! You can now close this window.');
            </script>
            """
        else:
            return f"Token exchange failed: {token_result}", 400
            
    except Exception as e:
        return f"Error during token exchange: {str(e)}", 500

@app.route('/test_linkedin_token', methods=['GET'])
def test_linkedin_token():
    """
    Test endpoint to validate LinkedIn access token
    """
    access_token = os.getenv('LINKEDIN_ACCESS_TOKEN')
    user_id = os.getenv('LINKEDIN_USER_ID')
    
    if not access_token:
        return jsonify({
            'error': 'LINKEDIN_ACCESS_TOKEN not found in .env file',
            'token_preview': 'None'
        }), 400
    
    print(f"[DEBUG] Testing LinkedIn token: {access_token[:20]}...")
    
    try:
        # Test token by calling LinkedIn profile API
        headers = {'Authorization': f'Bearer {access_token}'}
        
        # Try userinfo endpoint (OpenID Connect)
        response = requests.get('https://api.linkedin.com/v2/userinfo', headers=headers)
        
        return jsonify({
            'token_preview': f"{access_token[:20]}...",
            'token_length': len(access_token),
            'user_id_from_env': user_id,
            'api_response_status': response.status_code,
            'api_response': response.json() if response.status_code == 200 else response.text,
            'is_valid': response.status_code == 200
        })
        
    except Exception as e:
        return jsonify({
            'error': f'Token test failed: {str(e)}',
            'token_preview': f"{access_token[:20]}...",
            'token_length': len(access_token)
        }), 500

# --- Begin merged endpoints for new frontend features ---
import base64
import requests
from werkzeug.utils import secure_filename

ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif'}
GROQ_API_URL = 'https://api.groq.com/openai/v1/chat/completions'
GROQ_MODEL = 'llama3-70b-8192'

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def call_groq(system_message, user_prompt):
    headers = {
        "Authorization": f"Bearer {GROQ_API_KEY}",
        "Content-Type": "application/json"
    }
    payload = {
        "model": GROQ_MODEL,
        "messages": [
            {"role": "system", "content": system_message},
            {"role": "user", "content": user_prompt}
        ]
    }
    print(f"[DEBUG] Calling Groq with payload: {payload}")
    response = requests.post(GROQ_API_URL, headers=headers, json=payload)
    print(f"[DEBUG] Groq response status: {response.status_code}, body: {response.text}")
    response.raise_for_status()
    return response.json()['choices'][0]['message']['content'].strip()

@app.route('/upload-image', methods=['POST'])
def upload_image():
    if 'images' not in request.files:
        return jsonify({'error': 'No image part'}), 400
    files = request.files.getlist('images')
    file_urls = []
    for file in files:
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(path)
            file_urls.append(f"/{path}")
    return jsonify({'urls': file_urls})

@app.route('/refine-description', methods=['POST'])
def refine_description():
    data = request.get_json()
    user_description = data.get('description', '')
    if not user_description:
        return jsonify({'error': 'Description is required'}), 400
    prompt = (
        "Refine the following user-written description to make it more polished and engaging "
        "for a social media post. Preserve the original meaning but improve grammar, flow, and style. "
        "only give the description do not give extra answers\n\n"
        f"{user_description}"
    )
    system_msg = "You are a professional social media content writer who enhances user-generated descriptions."
    try:
        refined = call_groq(system_msg, prompt)
        return jsonify({'refined_description': refined})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/generate-title', methods=['POST'])
def generate_title():
    data = request.get_json()
    description = data.get('description', '')
    if not description:
        return jsonify({'error': 'Description required'}), 400
    prompt = f"Generate a short, catchy, and creative title for a social media post based on the following description, give only one most relevant title, do not need to give options, only the title:\n\n\"{description}\""
    system_msg = "You are an expert content writer for social media."
    try:
        title = call_groq(system_msg, prompt)
        return jsonify({'title': title})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/generate-description', methods=['POST'])
def generate_description():
    data = request.get_json()
    topic = data.get('topic', '')
    if not topic:
        return jsonify({'error': 'Topic required'}), 400
    
    prompt = (
        f"Generate a 4-5 line engaging description about '{topic}' for social media content.\n\n"
        f"Requirements:\n"
        f"- Write 4-5 lines of engaging content\n"
        f"- Make it informative and interesting\n"
        f"- Include relevant details and insights\n"
        f"- Use conversational tone\n"
        f"- Keep it suitable for social media sharing\n\n"
        f"Topic: {topic}"
    )
    system_msg = "You are a professional content writer who creates engaging social media descriptions."
    try:
        description = call_groq(system_msg, prompt)
        return jsonify({'description': description})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/generate-content', methods=['POST'])
def generate_content():
    data = request.get_json()
    title = data.get('title', '')
    description = data.get('description', '')
    platform = data.get('platform', 'Instagram')
    if not description:
        return jsonify({'error': 'Description required'}), 400
    prompt = (
        f"Generate a social media post for the platform '{platform}'.\n\n"
        f"Title: {title}\n"
        f"Description: {description}\n\n"
        f"The post should:\n"
        f"- Be tailored for {platform} audience.\n"
        f"- Include emojis (for casual platforms like Instagram, Twitter, Facebook).\n"
        f"- Be professional if it's for LinkedIn.\n"
        f"- Include a short caption and call-to-action if relevant.\n"
        f"- Suggest 5-7 relevant hashtags."
    )
    system_msg = "You are a professional social media marketer who writes engaging posts for various platforms."
    try:
        content = call_groq(system_msg, prompt)
        hashtags = [word for word in content.split() if word.startswith('#')]
        return jsonify({
            'content': content,
            'hashtags': hashtags
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500
# --- End merged endpoints for new frontend features ---

@app.route('/api/templates/list', methods=['GET'])
def list_templates():
    templates_dir = os.path.join('static', 'templates')
    # List only directories (folders), ignore files like .json
    template_folders = [
        name for name in os.listdir(templates_dir)
        if os.path.isdir(os.path.join(templates_dir, name))
    ]
    
    # Hide specific templates: health, finance, tech, travel
    hidden_templates = ['health', 'finance', 'tech', 'travel']
    visible_templates = [template for template in template_folders if template not in hidden_templates]
    
    return jsonify({'templates': visible_templates})

if __name__ == "__main__":
    app.run(debug=True, host="0.0.0.0", port=5001)
